import random

import pandas as pd
import telegram
from sqlalchemy.dialects.postgresql import psycopg2
from telegram.constants import ParseMode
from telegram import Update, InlineKeyboardButton, InlineKeyboardMarkup, ReplyKeyboardMarkup, KeyboardButton, \
    InputMediaPhoto
from telegram.ext import Application, CommandHandler, CallbackQueryHandler, MessageHandler, filters, ContextTypes
import logging
import json
import requests
from openpyxl import load_workbook
from openpyxl.styles import Font
import os
from datetime import datetime, timedelta
import aiohttp
from typing import List
from dateutil import parser
import requests
from docx2pdf import convert
import aiohttp
import os
import subprocess
import platform
from docx import Document
from docx.shared import Pt
from datetime import datetime
import asyncio
from collections import defaultdict
import psycopg2
import pandas as pd
from pandas import ExcelWriter
from docx import Document
from docx.shared import Pt, Inches
import subprocess
import os
import random


# Включаем логирование
logging.basicConfig(
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    level=logging.INFO
)

logger = logging.getLogger(__name__)

WEBHOOK_URL = "https://script.google.com/macros/s/AKfycbxg_xgttLRvUXc9nbpeqLakDoeh84zWDkHSWryDxNihwU6REtyXGntrhEOe5zi94JM/exec"

API_BASE_URL = "http://127.0.0.1:8000"

# URL вашего Django API
DJANGO_API_URL = 'http://127.0.0.1:8000/'
DJANGO_MEDIA_URL = 'http://localhost:8000/api'

# Директория для хранения фотографий
PHOTO_DIR = 'photos'

if not os.path.exists(PHOTO_DIR):
    os.makedirs(PHOTO_DIR)

reply_keyboard_main = [
    [KeyboardButton("/info")],
    [KeyboardButton("/start")],
    [KeyboardButton("/choice")],
]
reply_markup_kb_main = ReplyKeyboardMarkup(reply_keyboard_main, resize_keyboard=True, one_time_keyboard=False)


async def welcome_message(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    reply_keyboard = [
        [KeyboardButton("/info")],
        [KeyboardButton("/start")],
        [KeyboardButton("/choice")]
    ]
    reply_markup_kb = ReplyKeyboardMarkup(reply_keyboard, resize_keyboard=True, one_time_keyboard=False)

    await update.message.reply_text(
        "Добро пожаловать! Я бот-помощник по передаче фронт работ. Вот что я умею: \n"
        "/info — Для вызова информации всех команд\n"
        "/start — Для начала работы\n"
        "/choice — Для смены организации\n",
        reply_markup=reply_markup_kb
    )


async def choose_organization(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    user_id = update.message.from_user.id

    # Удаление кнопок у предыдущего сообщения бота
    if 'main_menu_message_id' in context.user_data:
        try:

            await context.bot.delete_message(
                chat_id=update.message.chat.id,
                message_id=context.user_data['main_menu_message_id']
            )
        except:
            pass

    # Получаем список организаций
    response = requests.get(f'{DJANGO_API_URL}organizations/')
    if response.status_code == 200:
        organizations = response.json()
        # Исключаем организацию с id = 3
        filtered_organizations = [org for org in organizations if org['organization'] != "БОС"]
        # Сортируем организации по алфавиту
        filtered_organizations.sort(key=lambda org: org['organization'])
        # Создание кнопок в две колонки
        keyboard = []
        for i in range(0, len(filtered_organizations), 2):
            row = [
                InlineKeyboardButton(filtered_organizations[i]['organization'], callback_data=f'org_{filtered_organizations[i]["id"]}')
            ]
            if i + 1 < len(filtered_organizations):
                row.append(InlineKeyboardButton(filtered_organizations[i + 1]['organization'], callback_data=f'org_{filtered_organizations[i + 1]["id"]}'))
            keyboard.append(row)
        keyboard.append([InlineKeyboardButton("Назад", callback_data='main_menu')])
        reply_markup = InlineKeyboardMarkup(keyboard)

        await update.message.reply_text('Выберите вашу организацию:', reply_markup=reply_markup)
        context.user_data['stage'] = 'choose_organization'
    else:
        await update.message.reply_text('Ошибка при получении списка организаций. Попробуйте снова.')


# Стартовая команда
async def start(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    user_id = update.message.from_user.id
    password = context.args[0] if context.args else None

    # Удаление кнопок у предыдущего сообщения бота
    if 'main_menu_message_id' in context.user_data:
        try:

            await context.bot.delete_message(
                chat_id=update.message.chat.id,
                message_id=context.user_data['main_menu_message_id']
            )
        except:
            pass  # Если возникнет ошибка при редактировании сообщения, игнорируем её

    context.user_data['stage'] = None
    reset_user_states(context)
    response = requests.get(f'{DJANGO_API_URL}users/chat/{user_id}/')

    if response.status_code == 404:
        if str(password).lower() == 'secret_password':
            context.user_data['is_authorized'] = True
            await update.message.reply_text('Пожалуйста, представьтесь. Введите ваше ФИО:')
            context.user_data['stage'] = 'get_full_name'

        elif str(password).lower() == 'secret_password_boss_12345':
            context.user_data['is_authorized'] = True
            context.user_data['organization_id'] = 3  # Устанавливаем организацию Босу
            await update.message.reply_text('Пожалуйста, представьтесь. Введите ваше ФИО:')
            context.user_data['stage'] = 'get_full_name_boss'

        elif str(password).startswith('baseinfo_'):
            context.user_data['params'] = password.split('_')[1:]  # Сохраняем параметры для дальнейшего использования
            await handle_baseinfo(update, context)

        else:
            await update.message.reply_text('Пожалуйста, введите пароль для авторизации командой /start [пароль]:')
            context.user_data['stage'] = 'get_password'
    else:
        user_data = response.json()
        if user_data['is_authorized']:
            if str(password).startswith('baseinfo_'):
                context.user_data['params'] = password.split('_')[1:]  # Сохраняем параметры для дальнейшего использования
                await handle_baseinfo(update, context)

            elif user_data['organization_id']:
                await send_main_menu(update.message.chat.id, context, user_data['full_name'], user_data['organization_id'])
            else:
                await update.message.reply_text('Пожалуйста, выберите организацию командой /choice.')
        else:
            if str(password).lower() == 'secret_password':
                user_data['is_authorized'] = True
                requests.put(f'{DJANGO_API_URL}users/{user_id}/', json=user_data)
                await update.message.reply_text('Пожалуйста, представьтесь. Введите ваше ФИО:')
                context.user_data['stage'] = 'get_full_name'

            elif str(password).lower() == 'secret_password_boss_12345':
                user_data['is_authorized'] = True
                user_data['organization_id'] = 3  # Устанавливаем организацию Босу
                requests.put(f'{DJANGO_API_URL}users/{user_id}/', json=user_data)
                await update.message.reply_text('Пожалуйста, представьтесь. Введите ваше ФИО:')
                context.user_data['stage'] = 'get_full_name_boss'

            else:
                await update.message.reply_text('Пожалуйста, введите пароль для авторизации:')
                context.user_data['stage'] = 'get_password'


async def handle_tech_command(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    user_id = update.message.from_user.id


    # Получаем список открытых запросов техподдержки
    response = requests.get(f'{DJANGO_API_URL}support_tickets/?status=open')
    if response.status_code != 200:
        await update.message.reply_text('Ошибка при получении списка запросов техподдержки. Попробуйте снова позже.')
        return

    support_tickets = response.json()
    if not support_tickets:
        await update.message.reply_text('Нет открытых запросов техподдержки.')
        return

    keyboard = []

    for ticket in support_tickets:
        # Получаем информацию о пользователе
        user_response = requests.get(f'{DJANGO_API_URL}users/{ticket["sender_id"]}/')
        if user_response.status_code != 200:
            continue
        user_data = user_response.json()

        # Получаем информацию об организации
        org_response = requests.get(f'{DJANGO_API_URL}organizations/{user_data["organization_id"]}/')
        if org_response.status_code != 200:
            continue
        organization_name = org_response.json().get('organization', 'Неизвестная организация')

        # Форматируем ФИО
        full_name = user_data['full_name']
        if full_name:
            parts = full_name.split()
            formatted_name = f"{parts[0]} {parts[1][0]}. {parts[2][0]}." if len(parts) > 2 else full_name
        else:
            formatted_name = "Неизвестный пользователь"

        # Форматируем дату
        created_at = ticket['created_at'].split('T')[0]

        button_text = f"{organization_name}, {formatted_name} - {created_at}"
        keyboard.append([InlineKeyboardButton(button_text, callback_data=f'ticket_{ticket["id"]}')])

    reply_markup = InlineKeyboardMarkup(keyboard)
    await update.message.reply_text('Выберите запрос:', reply_markup=reply_markup)


async def handle_baseinfo(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    def get_organization_by_id(organization_id):
        response = requests.get(f'{DJANGO_API_URL}organizations/{organization_id}')
        if response.status_code == 200:
            return response.json()
        return {'organization': 'неизвестно'}

    def get_work_type_by_id(work_type_id):
        response = requests.get(f'{DJANGO_API_URL}worktypes/{work_type_id}')
        if response.status_code == 200:
            return response.json()
        return {'name': 'неизвестно'}

    try:
        id_object, block_section_id, floor = context.user_data['params']
        workforce_response = requests.get(f'{DJANGO_API_URL}frontworkforces/')
        volume_response = requests.get(f'{DJANGO_API_URL}volumes/')

        # Получаем информацию об объекте и секции
        object_response = requests.get(f'{DJANGO_API_URL}objects/{id_object}/')
        block_section_response = requests.get(f'{DJANGO_API_URL}blocksections/{block_section_id}/')

        if (workforce_response.status_code == 200 and volume_response.status_code == 200 and
                object_response.status_code == 200 and block_section_response.status_code == 200):

            workforces = workforce_response.json()
            volumes = volume_response.json()
            object_name = object_response.json().get('name', 'неизвестный объект')
            block_section_name = block_section_response.json().get('name', 'неизвестная секция')

            filtered_workforces = [
                wf for wf in workforces if wf['object_id'] == int(id_object) and
                                           wf['block_section_id'] == int(block_section_id) and
                                           wf['floor'] == floor
            ]

            filtered_volumes = [
                vol for vol in volumes if vol['object_id'] == int(id_object) and
                                          vol['block_section_id'] == int(block_section_id) and
                                          vol['floor'] == floor
            ]

            message = f"🏗️ *{object_name} - {block_section_name} - этаж {floor}*\n\n"

            if filtered_workforces:
                # Группируем данные по дате для workforces
                grouped_workforces = defaultdict(list)
                for wf in filtered_workforces:
                    date = datetime.fromisoformat(wf['date']).date()
                    grouped_workforces[date].append(wf)

                # Сортируем даты по убыванию
                sorted_workforce_dates = sorted(grouped_workforces.keys(), reverse=True)

                message += "\U0001F477 *Численность:*\n"
                for date in sorted_workforce_dates:
                    message += f"Дата: {date.strftime('%d.%m.%Y')}\n"
                    for wf in grouped_workforces[date]:
                        organization = get_organization_by_id(wf['organization_id'])
                        work_type = get_work_type_by_id(wf['work_type_id'])
                        message += (
                            f"{organization['organization']} - {work_type['name']} - {wf['workforce_count']} ч.\n")
                    message += "\n"

            if filtered_volumes:
                # Группируем данные по дате для volumes
                grouped_volumes = defaultdict(list)
                for vol in filtered_volumes:
                    date = datetime.fromisoformat(vol['date']).date()
                    grouped_volumes[date].append(vol)

                # Сортируем даты по убыванию
                sorted_volume_dates = sorted(grouped_volumes.keys(), reverse=True)

                message += "📐 *Объемы:*\n"
                for date in sorted_volume_dates:
                    message += f"Дата: {date.strftime('%d.%m.%Y')}\n"
                    for vol in grouped_volumes[date]:
                        organization = get_organization_by_id(vol['organization_id'])
                        work_type = get_work_type_by_id(vol['work_type_id'])
                        message += (f"{organization['organization']} - {work_type['name']} - {vol['volume']} м³\n")
                    message += "\n"

            await update.message.reply_text(message, parse_mode=ParseMode.MARKDOWN,)

        else:
            await update.message.reply_text(
                "Ошибка при получении данных о численности, объемах, объекте или секции. Попробуйте позже.")

    except Exception as e:
        await update.message.reply_text(f"Произошла ошибка: {str(e)}")





async def handle_message(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    user_id = update.message.from_user.id
    text = update.message.text

    stage = context.user_data.get('stage')

    reset_user_states(context)


    if stage == 'get_full_name':
        full_name = text
        context.user_data['full_name'] = full_name
        organization_id = context.user_data.get('organization_id')

        # Создаем пользователя в базе данных
        user_data = {
            'chat_id': user_id,
            'full_name': full_name,
            'is_authorized': context.user_data.get('is_authorized', False),
            'organization_id': organization_id,  # Передаем organization как None
        }
        response = requests.post(f'{DJANGO_API_URL}users/', json=user_data)
        if response.status_code == 201:
            response = requests.get(f'{DJANGO_API_URL}organizations/')
            if response.status_code == 200:
                organizations = response.json()

                # Исключаем организацию с id = 3
                filtered_organizations = [org for org in organizations if org['organization'] != "БОС"]
                # Сортируем организации по алфавиту
                filtered_organizations.sort(key=lambda org: org['organization'])
                # Создание кнопок в две колонки
                keyboard = []
                for i in range(0, len(filtered_organizations), 2):
                    row = [
                        InlineKeyboardButton(filtered_organizations[i]['organization'],
                                             callback_data=f'org_{filtered_organizations[i]["id"]}')
                    ]
                    if i + 1 < len(filtered_organizations):
                        row.append(InlineKeyboardButton(filtered_organizations[i + 1]['organization'],
                                                        callback_data=f'org_{filtered_organizations[i + 1]["id"]}'))
                    keyboard.append(row)

                reply_markup = InlineKeyboardMarkup(keyboard)

                await update.message.reply_text('Выберите вашу организацию:', reply_markup=reply_markup)
                context.user_data['stage'] = 'choose_organization'
            else:
                await update.message.reply_text('Ошибка при получении списка организаций. Попробуйте снова.')
        else:
            await update.message.reply_text('Ошибка при создании пользователя. Попробуйте снова.')

    elif stage == 'get_full_name_boss':
        full_name = text
        context.user_data['full_name'] = full_name
        organization_id = context.user_data.get('organization_id')

        # Создаем пользователя в базе данных
        user_data = {
            'chat_id': user_id,
            'full_name': full_name,
            'is_authorized': context.user_data.get('is_authorized', False),
            'organization_id': organization_id,  # Передаем organization_id для босса
        }
        response = requests.post(f'{DJANGO_API_URL}users/', json=user_data)
        if response.status_code == 201:
            response = requests.get(f'{DJANGO_API_URL}objects/')
            if response.status_code == 200:
                objects = response.json()

                if objects:
                    # Сортируем объекты по имени в алфавитном порядке
                    objects.sort(key=lambda obj: obj['name'])
                    # Создаем клавиатуру с кнопками в две колонки
                    keyboard = []
                    for i in range(0, len(objects), 2):
                        row = [
                            InlineKeyboardButton(objects[i]['name'], callback_data=f'object_{objects[i]["id"]}')
                        ]
                        if i + 1 < len(objects):
                            row.append(InlineKeyboardButton(objects[i + 1]['name'],
                                                            callback_data=f'object_{objects[i + 1]["id"]}'))
                        keyboard.append(row)

                    reply_markup = InlineKeyboardMarkup(keyboard)



                reply_markup = InlineKeyboardMarkup(keyboard)
                await update.message.reply_text('Выберите ваш объект:', reply_markup=reply_markup)
                context.user_data['stage'] = 'choose_object'
            else:
                await update.message.reply_text('Ошибка при получении списка объектов. Попробуйте снова.')
        else:
            await update.message.reply_text('Ошибка при создании пользователя. Попробуйте снова.')


    elif stage == 'get_password':
        if text == 'secret_password':
            response = requests.get(f'{DJANGO_API_URL}users/{user_id}/')
            if response.status_code == 200:
                user_data = response.json()
                user_data['is_authorized'] = True
                response = requests.put(f'{DJANGO_API_URL}users/{user_id}/', json=user_data)
                if response.status_code == 200:
                    await update.message.reply_text(f'Вы успешно авторизованы, {user_data["full_name"]}!')
                    await update.message.reply_text('Пожалуйста, представьтесь. Введите ваше ФИО:')
                    context.user_data['stage'] = 'get_full_name'
                else:
                    await update.message.reply_text('Ошибка при обновлении данных. Попробуйте снова.')
            else:
                # Создание нового пользователя если не существует
                context.user_data['is_authorized'] = True
                await update.message.reply_text('Пожалуйста, представьтесь. Введите ваше ФИО:')
                context.user_data['stage'] = 'get_full_name'

        elif text == 'secret_password_boss_12345':
            response = requests.get(f'{DJANGO_API_URL}users/{user_id}/')
            if response.status_code == 200:
                user_data = response.json()
                user_data['is_authorized'] = True
                user_data['organization_id'] = 3  # Устанавливаем организацию Босу
                response = requests.put(f'{DJANGO_API_URL}users/{user_id}/', json=user_data)
                if response.status_code == 200:
                    await update.message.reply_text(f'Вы успешно авторизованы, {user_data["full_name"]}!')
                    await update.message.reply_text('Пожалуйста, представьтесь. Введите ваше ФИО:')
                    context.user_data['stage'] = 'get_full_name_boss'
                else:
                    await update.message.reply_text('Ошибка при обновлении данных. Попробуйте снова.')
            else:
                context.user_data['is_authorized'] = True
                context.user_data['organization_id'] = 3  # Устанавливаем организацию Босу
                await update.message.reply_text('Пожалуйста, представьтесь. Введите ваше ФИО:')
                context.user_data['stage'] = 'get_full_name_boss'

        elif text == 'test_front_section':
            await update.message.reply_text('привет из Get_password')

        else:
            await update.message.reply_text('Неверный пароль, попробуйте еще раз:')



    elif stage and stage.startswith('rework_'):
        front_id = int(stage.split('_')[1])
        boss_response = requests.get(f'{DJANGO_API_URL}users/chat/{user_id}/')
        if boss_response.status_code == 200:
            boss_data = boss_response.json()
            boss_id = boss_data['id']
            boss_name = boss_data['full_name']
            front_response = requests.get(f'{DJANGO_API_URL}fronttransfers/{front_id}/')
            if front_response.status_code == 200:
                front_data = front_response.json()

                # Обновляем необходимые поля, оставляя остальные без изменений
                updated_data = front_data
                updated_data.update({
                    'status': 'in_process',
                    'remarks': text,
                    'boss': boss_id
                })

                logger.info(

                    f"Отправка данных в API для обновления записи FrontTransfer: {json.dumps(updated_data, indent=2)}")

                response = requests.put(f'{DJANGO_API_URL}fronttransfers/{front_id}/', json=updated_data)
                logger.info(
                    f"Ответ от API при обновлении записи FrontTransfer: {response.status_code}, {response.text}")
                if response.status_code == 200:
                    sender_chat_id = front_data['sender_chat_id']

                    # Получаем названия объекта, вида работ и блока/секции
                    object_name = "неизвестно"
                    block_section_name = "неизвестно"
                    work_type_name = "неизвестно"
                    object_response = requests.get(f'{DJANGO_API_URL}objects/{front_data["object_id"]}/')

                    if object_response.status_code == 200:
                        object_name = object_response.json().get('name', "неизвестно")

                    block_section_response = requests.get(
                        f'{DJANGO_API_URL}blocksections/{front_data["block_section_id"]}/')

                    if block_section_response.status_code == 200:
                        block_section_name = block_section_response.json().get('name', "неизвестно")

                    work_type_response = requests.get(f'{DJANGO_API_URL}worktypes/{front_data["work_type_id"]}/')

                    if work_type_response.status_code == 200:
                        work_type_name = work_type_response.json().get('name', "неизвестно")

                    notification_text = (
                        f"\U0000274C Генеральный подрядчик *{boss_name}* отклонил передачу фронт работ:\n"
                        f"\n\n*Объект:* {object_name}\n"
                        f"*Секция/Блок:* {block_section_name}\n"
                        f"*Этаж:* {front_data.get('floor', 'неизвестно')}\n"
                        f"*Вид работ:* {work_type_name}\n"
                        f"*Причина:* {text}\n"
                        "\n\n_Вы можете повторно передать фронт с учетом замечаний._"
                    )

                    keyboard = [
                        [InlineKeyboardButton("Передать фронт заново", callback_data='transfer')],

                    ]

                    reply_markup = InlineKeyboardMarkup(keyboard)

                    await context.bot.send_message(
                        chat_id=sender_chat_id,
                        text=notification_text,
                        parse_mode=ParseMode.MARKDOWN,
                        reply_markup=reply_markup

                    )

                    keyboard2 = [

                        [InlineKeyboardButton("Просмотр переданных фронтов", callback_data='view_fronts')],

                    ]

                    reply_markup2 = InlineKeyboardMarkup(keyboard2)
                    await update.message.reply_text('Комментарий отправлен подрядчику. Фронт отправлен на доработку.',
                                                    reply_markup=reply_markup2)

                else:
                    await update.message.reply_text(f'Ошибка при обновлении статуса фронта: {response.text}')
            else:
                await update.message.reply_text('Ошибка при получении деталей фронта. Попробуйте снова.')
        else:
            await update.message.reply_text('Ошибка при получении данных пользователя. Попробуйте снова.')
        context.user_data['stage'] = None

    elif stage and stage.startswith('decline_'):
        front_id = int(stage.split('_')[1])
        user_chat_id = update.message.from_user.id
        user_name = requests.get(f'{DJANGO_API_URL}users/chat/{user_chat_id}/').json()['full_name']
        decline_reason = text

        # Обновляем статус фронта на "отклонен"
        front_response = requests.get(f'{DJANGO_API_URL}fronttransfers/{front_id}/')
        if front_response.status_code == 200:
            front_data = front_response.json()
            boss_id = front_data['boss_id']
            sender_chat_id = front_data['sender_chat_id']

            # Получаем chat_id босса по его id
            boss_response = requests.get(f'{DJANGO_API_URL}users/{boss_id}/')
            if boss_response.status_code == 200:
                boss_chat_id = boss_response.json()['chat_id']
            else:
                await update.message.reply_text('Ошибка при получении chat_id ген подрядчика.')
                return

            # Получаем названия объекта, вида работ и блока/секции
            object_response = requests.get(f'{DJANGO_API_URL}objects/{front_data["object_id"]}/')
            work_type_response = requests.get(f'{DJANGO_API_URL}worktypes/{front_data["work_type_id"]}/')
            block_section_response = requests.get(
                f'{DJANGO_API_URL}blocksections/{front_data["block_section_id"]}/')

            if object_response.status_code == 200:
                object_name = object_response.json().get('name', 'Неизвестный объект')
            else:
                object_name = 'Неизвестный объект'

            if work_type_response.status_code == 200:
                work_type_name = work_type_response.json().get('name', 'Неизвестный вид работ')
            else:
                work_type_name = 'Неизвестный вид работ'

            if block_section_response.status_code == 200:
                block_section_name = block_section_response.json().get('name', 'Неизвестный блок/секция')
            else:
                block_section_name = 'Неизвестный блок/секция'

            # Получаем данные об организации отправителя
            sender_response = requests.get(f'{DJANGO_API_URL}users/{front_data["sender_id"]}/')
            if sender_response.status_code == 200:
                sender_data = sender_response.json()
                organization_id = sender_data['organization_id']
                organization_response = requests.get(f'{DJANGO_API_URL}organizations/{organization_id}/')

                if organization_response.status_code == 200:
                    organization_name = organization_response.json().get('organization', 'неизвестно')
                else:
                    organization_name = 'неизвестно'
            # Получаем данные о новом виде работ
            next_work_type_id = front_data['next_work_type_id']
            if next_work_type_id:
                next_work_type_response = requests.get(f'{DJANGO_API_URL}worktypes/{next_work_type_id}/')
                if next_work_type_response.status_code == 200:
                    next_work_type_name = next_work_type_response.json().get('name', 'Не указан')
                else:
                    next_work_type_name = 'Не указан'
            else:
                next_work_type_name = 'Не указан'

            updated_data = {
                'sender_id': front_data['sender_id'],
                'sender_chat_id': front_data['sender_chat_id'],
                'object_id': front_data['object_id'],
                'work_type_id': front_data['work_type_id'],
                'block_section_id': front_data['block_section_id'],
                'floor': front_data['floor'],
                'status': 'in_process',
                'remarks': decline_reason,
                'created_at': front_data['created_at'],
                'approval_at': front_data['approval_at'],
                'boss_id': boss_id,
                'receiver_id': front_data['receiver_id'],
                'next_work_type_id': front_data['next_work_type_id'],
                'photo_ids': front_data['photo_ids'],
            }

            response = requests.put(f'{DJANGO_API_URL}fronttransfers/{front_id}/', json=updated_data)
            if response.status_code == 200:
                # Уведомление ген подрядчика и создателя
                notification_text = (
                    f"\U0000274C Фронт работ отклонен *{user_name}*:\n"
                    f"\n\n*Объект:* {object_name}\n"
                    f"*Секция/Блок:* {block_section_name}\n"
                    f"*Этаж:* {front_data['floor']}\n\n"
                    f"*Вид работ:* {work_type_name}\n"
                    f"*Новый вид работ:* {next_work_type_name}\n"
                    f"*Причина отклонения:* {decline_reason}\n"

                )
                notification_text2 = (
                    f"\U0000274C Фронт работ отклонен *{user_name}*:\n"
                    f"\n\n*Объект:* {object_name}\n"
                    f"*Секция/Блок:* {block_section_name}\n"
                    f"*Этаж:* {front_data['floor']}\n\n"
                    f"*Вид работ:* {work_type_name}\n"
                    f"*Новый вид работ:* {next_work_type_name}\n"
                    f"*Причина отклонения:* {decline_reason}\n"
                    "\n\n_Исправьте замечания и повторно передайте фронт через команду /start_"
                )

                await context.bot.send_message(
                    chat_id=boss_chat_id,
                    text=notification_text,
                    parse_mode=ParseMode.MARKDOWN
                )
                await context.bot.send_message(
                    chat_id=sender_chat_id,
                    text=notification_text2,
                    parse_mode=ParseMode.MARKDOWN
                )
                await update.message.reply_text('Фронт успешно отклонен и уведомления отправлены.')

            else:
                await update.message.reply_text(f'Ошибка при обновлении статуса фронта: {response.text}')
        else:
            await update.message.reply_text('Ошибка при получении данных фронта. Попробуйте снова.')
        context.user_data['stage'] = None


    elif stage and stage.startswith('delete_error_'):
        front_id = int(stage.split('_')[2])
        delete_reason = text  # Получаем комментарий

        # Обновляем статус фронта на "deleted"
        front_response = requests.get(f'{DJANGO_API_URL}fronttransfers/{front_id}/')
        if front_response.status_code == 200:
            front_data = front_response.json()
            updated_data = front_data
            updated_data.update({
                'status': 'deleted',
                'remarks': delete_reason
            })

            response = requests.put(f'{DJANGO_API_URL}fronttransfers/{front_id}/', json=updated_data)
            if response.status_code == 200:
                sender_chat_id = front_data['sender_chat_id']

                # Уведомление отправителя
                notification_text = (
                    f"\U0000274C Генеральный подрядчик удалил ваш фронт работ по причине:\n\n"
                    f"*Комментарий:* _{delete_reason}_"
                )
                await context.bot.send_message(
                    chat_id=sender_chat_id,
                    text=notification_text,
                    parse_mode=ParseMode.MARKDOWN
                )
                keyboard = [
                    [InlineKeyboardButton("Просмотр переданных фронтов", callback_data='view_fronts')],
                ]
                reply_markup = InlineKeyboardMarkup(keyboard)

                await update.message.reply_text('Фронт успешно удален и уведомление отправлено.', reply_markup=reply_markup)
            else:
                await update.message.reply_text(f'Ошибка при удалении фронта: {response.text}')
        else:
            await update.message.reply_text('Ошибка при получении данных фронта. Попробуйте снова.')
        context.user_data['stage'] = None

    else:
        response = requests.get(f'{DJANGO_API_URL}users/chat/{user_id}/')
        if response.status_code == 404:
            await update.message.reply_text('Пожалуйста, представьтесь. Введите ваше ФИО:')
            context.user_data['stage'] = 'get_full_name'
        else:
            user_data = response.json()
            if user_data['is_authorized']:
                if user_data['organization_id']:
                    await welcome_message(update, context)
                else:
                    await update.message.reply_text('Пожалуйста, выберите организацию командой /choice.')
            else:
                await update.message.reply_text('Пожалуйста, введите пароль для авторизации:')
                context.user_data['stage'] = 'get_password'


def reset_user_states(context):
    context.user_data['expecting_workforce_count'] = False
    context.user_data['expecting_new_workforce_count'] = False
    context.user_data['expecting_volume_count'] = False
    context.user_data['expecting_new_volume_count'] = False
    context.user_data['expecting_prefab_quantity'] = False
    context.user_data['expecting_shipment_quantity'] = False
    context.user_data['expecting_sgp_quantity'] = False
    context.user_data['expecting_new_status_quantity'] = False
    context.user_data['refactor_prefab_count'] = False
    context.user_data['expecting_new_status_prefab'] = False
    context.user_data['stage'] = None
    context.user_data['expecting_montage_quantity'] = False
    context.user_data['expecting_stock_quantity'] = False


async def send_main_menu(chat_id, context: ContextTypes.DEFAULT_TYPE, full_name: str, organization_id: int) -> None:
    if not organization_id:
        await context.bot.send_message(
            chat_id=chat_id,
            text="Вы не выбрали организацию. Пожалуйста, выберите её командой /choice."
        )
        return

    reset_user_states(context)
    response = requests.get(f'{DJANGO_API_URL}organizations/{organization_id}/')
    if response.status_code == 200:
        organization_data = response.json()
        organization_name = organization_data['organization']
        is_general_contractor = organization_data.get('is_general_contractor', False)
        is_factory = organization_data.get('factory', False)
    else:
        organization_name = "Неизвестная организация"
        is_general_contractor = False
        is_factory = False

    if is_general_contractor:
        user_response = requests.get(f'{DJANGO_API_URL}users/chat/{chat_id}/')
        if user_response.status_code == 200:
            user_data = user_response.json()
            object_id = user_data.get('object_id')
            if object_id:
                object_response = requests.get(f'{DJANGO_API_URL}objects/{object_id}/')
                if object_response.status_code == 200:
                    object_data = object_response.json()
                    object_name = object_data.get('name', 'Неизвестный объект')
                else:
                    object_name = 'Неизвестный объект'
            else:
                object_name = 'Объект не указан'
        else:
            object_name = 'Ошибка при получении данных пользователя'
    else:
        object_name = ''

    if is_general_contractor:
        keyboard = [
            [InlineKeyboardButton("\U0001F4C4 Фронт работ", callback_data='frontbutton')],
            [InlineKeyboardButton("\U0001F477 Просмотреть численность", callback_data='view_workforce')],
            [InlineKeyboardButton("📐 Просмотреть объем", callback_data='view_volume')],
            [InlineKeyboardButton("⚒️ Префабы", callback_data='prefabsoptionlist')],
            [InlineKeyboardButton("🔄 Сменить объект", callback_data='changeobject')],
            [InlineKeyboardButton("📞 Тех. поддержка", callback_data='support')]
        ]
        text = f'Здравствуйте, {full_name} из организации "{organization_name}"! Вы привязаны к объекту "{object_name}". Выберите действие:'
    elif is_factory:
        keyboard = [
            [InlineKeyboardButton("🏭 Факт на производство", callback_data='fact_production')],
            [InlineKeyboardButton("📋 СГП", callback_data='sgp')],
            [InlineKeyboardButton("🚚 Отгрузка", callback_data='shipment')],
            [InlineKeyboardButton("📝 Замечания", callback_data='remarks')],
            [InlineKeyboardButton("✏️ Редактирование префаба", callback_data='edit_prefab')],
            [InlineKeyboardButton("📊 Сводка по объекту", callback_data='summary_by_object')],
            [InlineKeyboardButton("📞 Тех. поддержка", callback_data='support')]
        ]
        text = f'Здравствуйте, {full_name} с завода "{organization_name}"! Выберите действие:'
    else:
        keyboard = [
            [InlineKeyboardButton("\U0001F4C4 Фронт", callback_data='front_menu')],
            [InlineKeyboardButton("\U0001F477 Численность", callback_data='workforce_menu')],
            [InlineKeyboardButton("📐 Объем", callback_data='volume_menu')],
            [InlineKeyboardButton("📞 Тех. поддержка", callback_data='support')]
        ]
        text = f'Здравствуйте, {full_name} из организации "{organization_name}"! Выберите действие:'

    reply_markup = InlineKeyboardMarkup(keyboard)

    message = await context.bot.send_message(
        chat_id=chat_id,
        text=text,
        reply_markup=reply_markup
    )
    context.user_data['main_menu_message_id'] = message.message_id

async def show_front_details(query: Update, context: ContextTypes.DEFAULT_TYPE, front_id: int) -> None:
    await query.message.delete()
    response = requests.get(f'{DJANGO_API_URL}fronttransfers/{front_id}/')
    if response.status_code == 200:
        front = response.json()
        sender_chat_id = front['sender_chat_id']

        sender_response = requests.get(f'{DJANGO_API_URL}users/chat/{sender_chat_id}/')
        if sender_response.status_code == 200:
            sender_full_name = sender_response.json()["full_name"]

            object_name = requests.get(f'{DJANGO_API_URL}objects/{front["object_id"]}/').json().get('name', 'неизвестно')
            work_type_name = requests.get(f'{DJANGO_API_URL}worktypes/{front["work_type_id"]}/').json().get('name', 'неизвестно')
            block_section_name = requests.get(f'{DJANGO_API_URL}blocksections/{front["block_section_id"]}/').json().get('name', 'неизвестно')
            work_type_new_name = requests.get(f'{DJANGO_API_URL}worktypes/{front["next_work_type_id"]}/').json().get('name', 'неизвестно')

            if front['next_work_type_id']:

                message_text = (
                    f"*Отправитель:* {sender_full_name}\n\n"
                    f"*Объект:* {object_name}\n"
                    f"*Вид работ:* {work_type_name}\n"
                    f"*Блок/Секция:* {block_section_name}\n"
                    f"*Этаж:* {front['floor']}\n\n"
                    f"*Новый вид работ:* {work_type_new_name}\n"
                    f"*Дата передачи (МСК):* {datetime.fromisoformat(front['created_at']).strftime('%d.%m.%Y')}"
                )
            else:
                message_text = (
                    f"*Кому:* {sender_full_name}\n\n"
                    f"*Объект:* {object_name}\n"
                    f"*Вид работ:* {work_type_name}\n"
                    f"*Блок/Секция:* {block_section_name}\n"
                    f"*Этаж:* {front['floor']}\n\n"
                    # f"*Новый вид работ:* {work_type_new_name}\n"
                    f"*Дата передачи (МСК):* {datetime.fromisoformat(front['created_at']).strftime('%d.%m.%Y')}"
                )

            media_group = []
            photo_ids = front.get('photo_ids', [])
            for idx, photo_id in enumerate(photo_ids):
                if photo_id:
                    if idx == 0:
                        media_group.append(InputMediaPhoto(media=photo_id, caption=message_text, parse_mode=ParseMode.MARKDOWN))
                    else:
                        media_group.append(InputMediaPhoto(media=photo_id))

            keyboard = [
                [InlineKeyboardButton("\U00002705 Принять", callback_data=f"accept_front_{front_id}")],
                 [InlineKeyboardButton("\U0000274C Отклонить", callback_data=f"decline_front_{front_id}")]
            ]
            reply_markup = InlineKeyboardMarkup(keyboard)

            if media_group:
                await context.bot.send_media_group(chat_id=query.message.chat.id, media=media_group)
                await context.bot.send_message(
                    chat_id=query.message.chat.id,
                    text="Выберите действие:",
                    reply_markup=reply_markup
                )
            else:
                await context.bot.send_message(
                    chat_id=query.message.chat.id,
                    text=message_text,
                    parse_mode=ParseMode.MARKDOWN
                )
                await context.bot.send_message(
                    chat_id=query.message.chat.id,
                    text="Выберите действие:",
                    reply_markup=reply_markup
                )
        else:
            await query.message.reply_text("Ошибка при получении данных отправителя.")
    else:
        await query.message.reply_text("Ошибка при получении данных фронта.")

async def list_accept_fronts(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    await update.callback_query.message.delete()
    user_id = update.callback_query.from_user.id

    # Получаем информацию о пользователе по chat_id
    response = requests.get(f'{DJANGO_API_URL}users/chat/{user_id}/')
    if response.status_code == 200:
        user_data = response.json()
        receiver_id = user_data['id']

        # Получаем все фронты с нужными параметрами
        response = requests.get(f'{DJANGO_API_URL}fronttransfers/?status=on_consideration')
        if response.status_code == 200:
            fronts = response.json()
            # Фильтруем фронты по receiver_id
            filtered_fronts = [front for front in fronts if front['receiver_id'] == receiver_id]

            if filtered_fronts:
                keyboard = []
                for front in filtered_fronts:
                    # Получаем имена объектов, видов работ и блоков/секции
                    object_name = requests.get(f'{DJANGO_API_URL}objects/{front["object_id"]}/').json().get('name', 'неизвестно')
                    work_type_name = requests.get(f'{DJANGO_API_URL}worktypes/{front["work_type_id"]}/').json().get('name', 'неизвестно')
                    block_section_name = requests.get(f'{DJANGO_API_URL}blocksections/{front["block_section_id"]}/').json().get('name', 'неизвестно')

                    button_text = f"{object_name} - {work_type_name} - {block_section_name} - {front['floor']}"
                    callback_data = f"accept_{front['id']}"
                    keyboard.append([InlineKeyboardButton(button_text, callback_data=callback_data)])

                keyboard.append([InlineKeyboardButton("Назад", callback_data='main_menu')])
                reply_markup = InlineKeyboardMarkup(keyboard)

                await update.callback_query.message.reply_text("Выберите фронт для принятия:", reply_markup=reply_markup)
            else:
                await update.callback_query.message.reply_text("Нет фронтов для принятия.")
                await send_main_menu(update.callback_query.message.chat.id, context, user_data['full_name'], user_data['organization_id'])

        else:
            await update.callback_query.message.reply_text("Ошибка при получении фронтов.")
    else:
        await update.callback_query.message.reply_text("Ошибка при получении данных пользователя.")

async def choose_work_type(query: Update, context: ContextTypes.DEFAULT_TYPE, object_id: int) -> None:
    await query.message.delete()  # Удаление предыдущего сообщения

    user_chat_id = query.message.chat.id
    user_response = requests.get(f'{DJANGO_API_URL}users/chat/{user_chat_id}')

    if user_response.status_code == 200:
        user_data = user_response.json()
        organization_id = user_data['organization_id']

        if organization_id is None:
            await query.message.reply_text('Ошибка: пользователь не принадлежит ни одной организации.')
            return

        common_work_types_ids = await get_common_work_types(object_id, organization_id)

        if common_work_types_ids:
            ids_query = "&".join([f"ids={id}" for id in common_work_types_ids])
            async with aiohttp.ClientSession() as session:
                async with session.get(f'{DJANGO_API_URL}worktypes/?{ids_query}') as response:
                    if response.status == 200:
                        work_types = await response.json()
                        keyboard = [
                            [InlineKeyboardButton(work['name'], callback_data=f'work_{work["id"]}')] for work in work_types
                        ]
                        keyboard.append([InlineKeyboardButton("Назад", callback_data='front_menu')])
                        reply_markup = InlineKeyboardMarkup(keyboard)
                        await query.message.reply_text('Выберите вид работ:', reply_markup=reply_markup)
                        context.user_data['object_id'] = object_id
                        context.user_data['stage'] = 'choose_work_type'
                    else:
                        await query.message.reply_text('Ошибка при получении списка типов работ. Попробуйте снова.')
        else:
            await query.message.reply_text('Нет доступных видов работ для выбранного объекта и вашей организации.')
    else:
        await query.message.reply_text('Ошибка при получении данных пользователя. Попробуйте снова.')


async def choose_block_section(query: Update, context: ContextTypes.DEFAULT_TYPE, work_type_id: int) -> None:
    await query.message.delete()  # Удаление предыдущего сообщения
    object_id = context.user_data['object_id']
    response = requests.get(f'{DJANGO_API_URL}objects/{object_id}/blocksections/')
    if response.status_code == 200:
        block_sections = response.json()
        keyboard = [
            [InlineKeyboardButton(block['name'], callback_data=f'block_{block["id"]}')] for block in block_sections
        ]
        keyboard.append([InlineKeyboardButton("Назад", callback_data='front_menu')])
        reply_markup = InlineKeyboardMarkup(keyboard)
        await query.message.reply_text('Выберите блок или секцию:', reply_markup=reply_markup)
        context.user_data['work_type_id'] = work_type_id
        context.user_data['stage'] = 'choose_block_section'
    else:
        await query.message.reply_text('Ошибка при получении списка блоков или секций. Попробуйте снова.')


async def choose_floor(query: Update, context: ContextTypes.DEFAULT_TYPE, block_section_id: int) -> None:
    await query.message.delete()  # Удаление предыдущего сообщения
    response = requests.get(f'{DJANGO_API_URL}blocksections/{block_section_id}/')
    if response.status_code == 200:
        block_section = response.json()
        number_of_floors_bottom = block_section['number_of_floors_bottom']
        number_of_floors = block_section['number_of_floors']

        # Генерация кнопок этажей в две колонки, исключая 0
        keyboard = []
        for i in range(number_of_floors_bottom, number_of_floors + 1):
            if i == 0:
                continue
            if len(keyboard) == 0 or len(keyboard[-1]) == 2:
                keyboard.append([InlineKeyboardButton(f'{i} этаж', callback_data=f'floor_{i}')])
            else:
                keyboard[-1].append(InlineKeyboardButton(f'{i} этаж', callback_data=f'floor_{i}'))

        # Добавление кнопки кровли
        keyboard.append([InlineKeyboardButton('Кровля', callback_data='floor_roof')])
        keyboard.append([InlineKeyboardButton("Назад", callback_data='front_menu')])
        reply_markup = InlineKeyboardMarkup(keyboard)
        await query.message.reply_text('Выберите этаж:', reply_markup=reply_markup)
        context.user_data['block_section_id'] = block_section_id
        context.user_data['stage'] = 'choose_floor'
    else:
        await query.message.reply_text('Ошибка при получении информации о блоке или секции. Попробуйте снова.')


async def confirm_transfer_data(query: Update, context: ContextTypes.DEFAULT_TYPE, floor: str) -> None:
    await query.message.delete()  # Удаление предыдущего сообщения
    floor_number = floor if floor != 'roof' else 'Кровля'
    context.user_data['floor'] = floor_number

    # Получаем названия объектов, видов работ и блоков/секций
    object_id = context.user_data['object_id']
    work_type_id = context.user_data['work_type_id']
    block_section_id = context.user_data['block_section_id']

    object_name = requests.get(f'{DJANGO_API_URL}objects/{object_id}/').json()['name']
    work_type_name = requests.get(f'{DJANGO_API_URL}worktypes/{work_type_id}/').json()['name']
    block_section_name= requests.get(f'{DJANGO_API_URL}blocksections/{block_section_id}/').json()['name']

    transfer_info = (
        f"Объект: {object_name}\n"
        f"Вид работ: {work_type_name}\n"
        f"Блок/Секция: {block_section_name}\n"
        f"Этаж: {floor_number}\n"
    )

    keyboard = [
        [InlineKeyboardButton("\U00002705 Да", callback_data='confirm_yes')],
        [InlineKeyboardButton("\U0000274C Нет", callback_data='confirm_no')]
    ]
    reply_markup = InlineKeyboardMarkup(keyboard)
    await query.message.reply_text(f'Подтвердите данные:\n{transfer_info}', reply_markup=reply_markup)

async def change_object_id(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    query = update.callback_query
    await query.answer()

    data = query.data
    user_id = query.from_user.id

    if data.startswith('object_'):
        object_id = int(data.split('_')[1])

        # Получаем текущие данные пользователя
        response = requests.get(f'{DJANGO_API_URL}users/chat/{user_id}/')
        if response.status_code == 200:
            user_data = response.json()
            user_data['object_id'] = object_id  # Обновляем поле object_id

            logger.info(f"Отправка данных в API: {json.dumps(user_data, indent=2)}")
            response = requests.put(f'{DJANGO_API_URL}users/{user_data["id"]}/', json=user_data)
            logger.info(f"Ответ от API: {response.status_code}, {response.text}")

            if response.status_code == 200:
                reply_keyboard = [
                    [KeyboardButton("/info")],
                    [KeyboardButton("/start")],
                    [KeyboardButton("/choice")]
                ]
                reply_markup_kb = ReplyKeyboardMarkup(reply_keyboard, resize_keyboard=True, one_time_keyboard=False)
                await query.message.delete()
                await query.message.reply_text(
                    'Объект успешно выбран!',
                    reply_markup=reply_markup_kb
                )
                user_data = response.json()
                await send_main_menu(query.message.chat.id, context, user_data['full_name'], user_data['organization_id'])
            else:
                await query.message.reply_text('Ошибка при сохранении данных. Попробуйте снова.')
        else:
            await query.message.reply_text('Ошибка при получении данных пользователя. Попробуйте снова.')


async def handle_transfer_confirmation(query: Update, context: ContextTypes.DEFAULT_TYPE, confirmed: bool) -> None:
    if confirmed:
        user_chat_id = str(query.from_user.id)  # Преобразуем в строку

        # Получаем данные пользователя по chat_id
        response = requests.get(f'{DJANGO_API_URL}users/chat/{user_chat_id}/')
        if response.status_code == 200:
            user_data = response.json()
            user_id = user_data['id']
        else:
            await query.message.reply_text('Ошибка при получении данных пользователя. Попробуйте снова.')
            return

        transfer_data = {
            'sender_id': user_id,  # Изменено на sender_id
            'sender_chat_id': user_chat_id,
            'object_id': context.user_data['object_id'],  # Изменено на object_id
            'work_type_id': context.user_data['work_type_id'],  # Изменено на work_type_id
            'block_section_id': context.user_data['block_section_id'],  # Изменено на block_section_id
            'floor': context.user_data['floor'],
            'created_at': datetime.now().isoformat(),
            'approval_at': datetime.now().isoformat(),
            'status': 'transferred',
            'photos': [],
            'is_finish': True
        }
        logger.info(f"Отправка данных в API для создания передачи фронта: {json.dumps(transfer_data, indent=2)}")
        response = requests.post(f'{DJANGO_API_URL}fronttransfers/', json=transfer_data)
        logger.info(f"Ответ от API при создании передачи фронта: {response.status_code}, {response.text}")
        if response.status_code == 200:
            transfer = response.json()
            context.user_data['transfer_id'] = transfer['id']
            context.user_data['photos'] = []  # Сброс списка фотографий
            context.user_data.pop('last_photo_message_id', None)  # Сброс идентификатора последнего сообщения

            # Удаляем только кнопки

            await query.message.reply_text(
                'Этаж успешно выбран! Пожалуйста, прикрепите фотографии (до 10 штук) или нажмите /done:')
            context.user_data['stage'] = 'attach_photos'

            # Оповещение главных подрядчиков
            transfer_data.update({
                'object_name': requests.get(f'{DJANGO_API_URL}objects/{context.user_data["object_id"]}/').json()['name'],
                'work_type_name': requests.get(f'{DJANGO_API_URL}worktypes/{context.user_data["work_type_id"]}/').json()['name'],
                'block_section_name': requests.get(f'{DJANGO_API_URL}blocksections/{context.user_data["block_section_id"]}/').json()['name'],
            })
            logger.info(f"Передача данных в notify_general_contractors: {json.dumps(transfer_data, indent=2)}")
            # await notify_general_contractors(context, transfer_data)
        else:
            await query.message.reply_text('Ошибка при создании передачи фронта. Попробуйте снова.')
    else:
        await query.message.delete()
        user_id = query.from_user.id
        response = requests.get(f'{DJANGO_API_URL}users/chat/{user_id}/')
        if response.status_code == 200:
            user_data = response.json()
            full_name = user_data.get('full_name', 'Пользователь')
            organization_id = user_data.get('organization_id', None)
            if organization_id:
                await query.message.reply_text('Передача отменена.', reply_markup=reply_markup_kb_main)
                await send_main_menu(query.message.chat.id, context, full_name, organization_id)
            else:
                await query.message.reply_text(
                    'Ошибка: у вас не выбрана организация. Пожалуйста, выберите организацию командой /choice.')
        else:
            await query.message.reply_text('Ошибка при получении данных пользователя. Попробуйте снова.')


async def handle_prefab_comment(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    comment = update.message.text
    updated_prefabs = context.user_data.get('updated_prefabs', [])
    print(updated_prefabs)
    for prefab_in_work_id in updated_prefabs:
        update_data = {
            'comments': comment
        }
        response = requests.patch(f'{DJANGO_API_URL}prefabs_in_work/{prefab_in_work_id}', json=update_data)
        if response.status_code != 200:
            await update.message.reply_text(f'Ошибка при добавлении комментария для префаба ID {prefab_in_work_id}. Попробуйте снова.')
            return

    await update.message.reply_text("Комментарий успешно добавлен ко всем префабам. Теперь добавьте фотографии (до 10) и нажмите /done для завершения.")
    context.user_data['stage'] = 'attach_photos_prefab_in_work'


async def handle_photo_upload(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    if context.user_data.get('stage') == 'attach_photos':
        transfer_id = context.user_data.get('transfer_id')
        if transfer_id is None:
            await update.message.reply_text('Ошибка: идентификатор передачи фронта не найден. Попробуйте снова.')
            return

        photos = context.user_data.get('photos', [])
        file_id = update.message.photo[-1].file_id
        photos.append(file_id)
        context.user_data['photos'] = photos

        # Удаляем предыдущее сообщение об успешной загрузке, если оно существует
        if 'last_photo_message_id' in context.user_data:
            try:
                await context.bot.delete_message(
                    chat_id=update.message.chat.id,
                    message_id=context.user_data['last_photo_message_id']
                )
            except telegram.error.BadRequest as e:
                if str(e) != "Message to delete not found":
                    raise

        if len(photos) < 10:
            reply_keyboard = [
                [KeyboardButton("/done")]
            ]
            reply_markup = ReplyKeyboardMarkup(reply_keyboard, resize_keyboard=True, one_time_keyboard=False)

            message = await update.message.reply_text(
                f'Фотография {len(photos)} из 10 успешно загружена. Прикрепите ещё или нажмите /done для завершения.',
                reply_markup=reply_markup
            )
            context.user_data['last_photo_message_id'] = message.message_id
        else:
            await finalize_photo_upload(update, context)

    elif context.user_data.get('stage') == 'attach_photos_prefab_in_work':
        updated_prefabs = context.user_data.get('updated_prefabs', [])
        if not updated_prefabs:
            await update.message.reply_text('Ошибка: идентификаторы префабов не найдены. Попробуйте снова.')
            return

        photos = context.user_data.get('photos', [])
        file_id = update.message.photo[-1].file_id
        photos.append(file_id)
        context.user_data['photos'] = photos

        if 'last_photo_message_id' in context.user_data:
            try:
                await context.bot.delete_message(
                    chat_id=update.message.chat.id,
                    message_id=context.user_data['last_photo_message_id']
                )
            except telegram.error.BadRequest as e:
                if str(e) != "Message to delete not found":
                    raise

        if len(photos) < 10:
            reply_keyboard = [
                [KeyboardButton("/done")]
            ]
            reply_markup = ReplyKeyboardMarkup(reply_keyboard, resize_keyboard=True, one_time_keyboard=False)

            message = await update.message.reply_text(
                f'Фотография {len(photos)} из 10 успешно загружена. Прикрепите ещё или нажмите /done для завершения.',
                reply_markup=reply_markup
            )
            context.user_data['last_photo_message_id'] = message.message_id
        else:
            await finalize_photo_upload_prefab_in_work(update, context)

    elif context.user_data.get('stage') == 'attach_photos_prefab_in_montage':
        await handle_prefab_photo_upload(update, context)

    elif context.user_data.get('stage') == 'attach_photos_support_ticket':
        ticket_id = context.user_data.get('ticket_id')
        if ticket_id is None:
            await update.message.reply_text('Ошибка: идентификатор заявки тех. поддержки не найден. Попробуйте снова.')
            return

        photos = context.user_data.get('photos', [])
        file_id = update.message.photo[-1].file_id
        photos.append(file_id)
        context.user_data['photos'] = photos

        if 'last_photo_message_id' in context.user_data:
            try:
                await context.bot.delete_message(
                    chat_id=update.message.chat.id,
                    message_id=context.user_data['last_photo_message_id']
                )
            except telegram.error.BadRequest as e:
                if str(e) != "Message to delete not found":
                    raise

        if len(photos) < 10:
            reply_keyboard = [
                [KeyboardButton("/done")]
            ]
            reply_markup = ReplyKeyboardMarkup(reply_keyboard, resize_keyboard=True, one_time_keyboard=False)

            message = await update.message.reply_text(
                f'Фотография {len(photos)} из 10 успешно загружена. Прикрепите ещё или нажмите /done для завершения.',
                reply_markup=reply_markup
            )
            context.user_data['last_photo_message_id'] = message.message_id
        else:
            await finalize_photo_upload_support_ticket(update, context)

async def finalize_photo_upload_prefab_in_work(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    updated_prefabs = context.user_data.get('updated_prefabs', [])
    photos = context.user_data.get('photos', [])

    # Отправка фотографий на сервер для каждого префаба
    for prefab_in_work_id in updated_prefabs:
        update_data = {
            'photos': photos
        }
        response = requests.patch(f'{DJANGO_API_URL}prefabs_in_work/{prefab_in_work_id}', json=update_data)
        if response.status_code != 200:
            await update.message.reply_text(f'Ошибка при загрузке фотографий для префаба ID {prefab_in_work_id}. Попробуйте снова.')
            return

    await update.message.reply_text("Фотографии успешно загружены ко всем префабам.")
    # Сброс стадии и вызов основного меню
    context.user_data['stage'] = None
    user_data = requests.get(f'{DJANGO_API_URL}users/chat/{update.message.chat.id}/').json()
    full_name = user_data.get('full_name', 'Пользователь')
    organization_id = user_data.get('organization_id', None)
    await send_main_menu(update.message.chat.id, context, full_name, organization_id)


    context.user_data['photos'] = []

    # Отправка уведомлений после загрузки фотографий
    if updated_prefabs:
        prefab_in_work_response = requests.get(f'{DJANGO_API_URL}prefabs_in_work/{updated_prefabs[0]}')
        if prefab_in_work_response.status_code == 200:
            prefab_in_work_data = prefab_in_work_response.json()
            prefab_id = prefab_in_work_data['prefab_id']

            prefab_response = requests.get(f'{DJANGO_API_URL}prefabs/{prefab_id}')
            if prefab_response.status_code == 200:
                prefab_data = prefab_response.json()
                organization_id = prefab_data['organization_id']

                # Получаем пользователей с этим organization_id
                users_response = requests.get(f'{DJANGO_API_URL}users/')
                if users_response.status_code == 200:
                    users_data = users_response.json()
                    target_users = [user for user in users_data if user['organization_id'] == organization_id]

                    media = [InputMediaPhoto(photo) for photo in photos]

                    # Получаем статус префаба
                    status = prefab_in_work_data.get('status', '')

                    # Определяем тип уведомления на основе статуса префаба
                    comment = prefab_in_work_data.get('comments', '')
                    if status == 'remark':
                        notification_text = (f"Инженер по качеству во время отгрузки обнаружил брак. Комментарий: {comment}\n\n"
                                             f"Перейдите во вкладку *Замечания* для изменения статуса.")
                    elif status == 'stock':
                        return

                    else:
                        notification_text = f"Генеральный подрядчик оставил комментарий: {comment}"

                    # Отправляем уведомления пользователям
                    for user in target_users:
                        chat_id = user['chat_id']
                        if media:
                            await context.bot.send_media_group(
                                chat_id=chat_id,
                                media=media,
                                parse_mode=ParseMode.MARKDOWN
                            )
                        await context.bot.send_message(
                            chat_id=chat_id,
                            text=notification_text,
                            parse_mode=ParseMode.MARKDOWN
                        )




async def handle_done_command(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    stage = context.user_data.get('stage')
    if stage == 'attach_photos':
        await finalize_photo_upload(update, context)
    elif stage == 'attach_photos_prefab_in_work':
        await finalize_photo_upload_prefab_in_work(update, context)
    elif context.user_data.get('stage') == 'attach_photos_prefab_in_montage':
        await finalize_photo_montage(update, context)
    elif stage == 'attach_photos_support_ticket':
        await finalize_photo_upload_support_ticket(update, context)

async def finalize_photo_upload(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    transfer_id = context.user_data.get('transfer_id')
    photos = context.user_data.get('photos', [])

    # Получаем текущие данные о фронте
    front_response = requests.get(f'{DJANGO_API_URL}fronttransfers/{transfer_id}/')
    if front_response.status_code == 200:
        front_data = front_response.json()
        # Обновляем поля, включая фото и статус
        update_data = {
            'sender_id': front_data.get('sender_id'),
            'object_id': front_data.get('object_id'),
            'work_type_id': front_data.get('work_type_id'),
            'block_section_id': front_data.get('block_section_id'),
            'floor': front_data.get('floor'),
            'status': 'transferred',
            'created_at': front_data.get('created_at'),
            'approval_at': front_data.get('approval_at'),
            'sender_chat_id': front_data.get('sender_chat_id'),
            'is_finish': front_data.get('is_finish'),
            'photo_ids': photos,
        }

        response = requests.put(f'{DJANGO_API_URL}fronttransfers/{transfer_id}/', json=update_data)
        logger.info(f"Ответ от API при обновлении передачи фронта: {response.status_code}, {response.text}")
        if response.status_code == 200:
            await update.message.reply_text('Фотографии успешно загружены. Передача фронта завершена!', reply_markup=reply_markup_kb_main)
            context.user_data['stage'] = None


            object_response = requests.get(f'{DJANGO_API_URL}objects/{front_data["object_id"]}/').json()
            # Уведомляем ген подрядчиков
            transfer_data = {
                'object_name': object_response['name'],
                'work_type_name': requests.get(f'{DJANGO_API_URL}worktypes/{front_data["work_type_id"]}/').json()['name'],
                'block_section_name': requests.get(f'{DJANGO_API_URL}blocksections/{front_data["block_section_id"]}/').json()['name'],
                'floor': front_data['floor'],
                'sender_chat_id': front_data['sender_chat_id'],
                'object_id': object_response['id']
            }
            await notify_general_contractors(context, transfer_data)

            # Отправляем сообщение с главным меню
            user_id = str(update.message.from_user.id)
            response = requests.get(f'{DJANGO_API_URL}users/chat/{user_id}/')
            if response.status_code == 200:
                user_data = response.json()
                full_name = user_data.get('full_name', 'Пользователь')
                organization_id = user_data.get('organization_id', None)
                if organization_id:
                    await send_main_menu(update.message.chat.id, context, full_name, organization_id)
                else:
                    await update.message.reply_text(
                        'Ошибка: у вас не выбрана организация. Пожалуйста, выберите организацию командой /choice.')
            else:
                await update.message.reply_text('Ошибка при получении данных пользователя. Попробуйте снова.')
        else:
            await update.message.reply_text(f'Ошибка при загрузке фотографий: {response.text}')
    else:
        await update.message.reply_text('Ошибка при получении данных фронта. Попробуйте снова.')

async def finalize_photo_upload_support_ticket(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    ticket_id = context.user_data.get('ticket_id')
    photo_ids = context.user_data.get('photos', [])

    # Отправка фотографий на сервер
    update_data = {
        'photo_ids': photo_ids
    }
    response = requests.patch(f'{DJANGO_API_URL}support_tickets/{ticket_id}', json=update_data)
    if response.status_code == 200:
        await update.message.reply_text("\U00002705 Ваш вопрос отправлен в тех. поддержку. Ожидайте ответа.", reply_markup=reply_markup_kb_main)

        # Получаем данные пользователя
        user_response = requests.get(f'{DJANGO_API_URL}users/chat/{update.message.chat.id}/')
        if user_response.status_code == 200:
            user_data = user_response.json()
            full_name = user_data.get('full_name', 'Пользователь')
            organization_id = user_data.get('organization_id')

            # Вызываем send_main_menu
            await send_main_menu(update.message.chat.id, context, full_name, organization_id)
        else:
            await update.message.reply_text('Ошибка при получении данных пользователя. Пожалуйста, попробуйте снова.')

        context.user_data['stage'] = None
        context.user_data['photos'] = []
        context.user_data['ticket_id'] = None
    else:
        logger.error(f"Error creating support ticket: {response.status_code} {response.text}")
        await update.message.reply_text('Ошибка при загрузке фотографий. Попробуйте снова.')

async def view_fronts(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    await update.callback_query.message.delete()

    user_chat_id = update.callback_query.from_user.id
    user_response = requests.get(f'{DJANGO_API_URL}users/chat/{user_chat_id}/')

    if user_response.status_code == 200:
        user_data = user_response.json()
        user_object_id = user_data.get('object_id')

        if user_object_id is None:
            await update.callback_query.message.reply_text("У вас нет назначенного объекта.")
            return

        response = requests.get(f'{DJANGO_API_URL}fronttransfers/?status=transferred')
        if response.status_code == 200:
            fronts = response.json()
            filtered_fronts = [front for front in fronts if front['object_id'] == user_object_id]

            if filtered_fronts:
                keyboard = []
                for front in filtered_fronts:
                    # Получаем имена объектов, видов работ и блоков/секций
                    object_response = requests.get(f'{DJANGO_API_URL}objects/{front["object_id"]}/')
                    work_type_response = requests.get(f'{DJANGO_API_URL}worktypes/{front["work_type_id"]}/')
                    block_section_response = requests.get(f'{DJANGO_API_URL}blocksections/{front["block_section_id"]}/')

                    if object_response.status_code == 200 and work_type_response.status_code == 200 and block_section_response.status_code == 200:
                        object_name = object_response.json()['name']
                        work_type_name = work_type_response.json()['name']
                        block_section_name = block_section_response.json()['name']

                        button_text = f"{object_name} - {work_type_name} - {block_section_name} - {front['floor']}"
                        callback_data = f"front_{front['id']}"
                        keyboard.append([InlineKeyboardButton(button_text, callback_data=callback_data)])
                    else:
                        # Обработка ошибок при получении данных
                        await update.callback_query.message.reply_text("Ошибка при получении данных. Попробуйте снова.")
                        return

                keyboard.append([InlineKeyboardButton("↻ Обновить", callback_data='view_fronts')])
                keyboard.append([InlineKeyboardButton("Назад", callback_data='frontbutton')])

                reply_markup = InlineKeyboardMarkup(keyboard)
                await update.callback_query.message.reply_text("Список текущих фронтов работ:", reply_markup=reply_markup)
            else:
                keyboard = []
                keyboard.append([InlineKeyboardButton("↻ Обновить", callback_data='view_fronts')])
                keyboard.append([InlineKeyboardButton("Назад", callback_data='frontbutton')])

                reply_markup = InlineKeyboardMarkup(keyboard)
                await update.callback_query.message.reply_text("Нет доступных фронтов работ со статусом 'передано'.", reply_markup=reply_markup)
        else:
            await update.callback_query.message.reply_text("Ошибка при получении списка фронтов работ. Попробуйте снова.")
    else:
        await update.callback_query.message.reply_text("Ошибка при получении данных пользователя. Попробуйте снова.")



async def view_front_details(query: Update, context: ContextTypes.DEFAULT_TYPE, front_id: int) -> None:
    await query.message.delete()  # Удаление предыдущего сообщения
    response = requests.get(f'{DJANGO_API_URL}fronttransfers/{front_id}/')
    if response.status_code == 200:
        front = response.json()
        sender_chat_id = front['sender_chat_id']  # Используем sender_chat_id

        sender_response = requests.get(f'{DJANGO_API_URL}users/chat/{sender_chat_id}/')
        if sender_response.status_code == 200:

            sender_response_id_org =  sender_response.json()["organization_id"]
            org_response = requests.get(f'{DJANGO_API_URL}organizations/{sender_response_id_org}/').json()["organization"]

            sender_full_name = sender_response.json()["full_name"]

            # Получаем имена объекта, вида работ и блока/секции
            object_response = requests.get(f'{DJANGO_API_URL}objects/{front["object_id"]}/')
            work_type_response = requests.get(f'{DJANGO_API_URL}worktypes/{front["work_type_id"]}/')
            block_section_response = requests.get(f'{DJANGO_API_URL}blocksections/{front["block_section_id"]}/')

            if object_response.status_code == 200 and work_type_response.status_code == 200 and block_section_response.status_code == 200:
                object_name = object_response.json()['name']
                work_type_name = work_type_response.json()['name']
                block_section_name = block_section_response.json()['name']
            else:
                await query.message.reply_text("Ошибка при получении данных. Попробуйте снова.")
                return

            message_text = (
                f"*Отправитель:* {sender_full_name}\n"
                f"*Организация:* {org_response}\n\n"
                f"*Объект:* {object_name}\n"
                f"*Вид работ:* {work_type_name}\n"
                f"*Блок/Секция:* {block_section_name}\n"
                f"*Этаж:* {front['floor']}\n\n"
                f"*Дата передачи (МСК):* {datetime.fromisoformat(front['created_at']).strftime('%d.%m.%Y')}"
            )

            # Список InputMediaPhoto для отправки группой
            media_group = []

            # Кнопка "На доработку"
            keyboard = [
                [InlineKeyboardButton("\U0000274C Доработка", callback_data=f"rework_{front_id}"),
                 InlineKeyboardButton("👥 Передать", callback_data=f"transfer_{front_id}"),
                 InlineKeyboardButton("\U00002705 Принять", callback_data=f"approve_{front_id}")],
                [InlineKeyboardButton("\U0001F6AB Удалить/Ошибка", callback_data=f"delete_error_{front_id}")],
                [InlineKeyboardButton("К списку фронтов", callback_data='view_fronts')]
            ]
            reply_markup = InlineKeyboardMarkup(keyboard)

            # Отправка фотографий
            photo_ids = front.get('photo_ids', [])
            for idx, photo_id in enumerate(photo_ids):
                if photo_id:
                    if idx == 0:
                        media_group.append(
                            InputMediaPhoto(media=photo_id, caption=message_text, parse_mode=ParseMode.MARKDOWN))
                    else:
                        media_group.append(InputMediaPhoto(media=photo_id))

            # Если есть фото, отправляем группу медиа
            if media_group:
                await context.bot.send_media_group(chat_id=query.message.chat.id, media=media_group)

                # Отправляем кнопки после попытки отправки медиа группы
                await context.bot.send_message(
                    chat_id=query.message.chat.id,
                    text="Выберите действие:",
                    reply_markup=reply_markup
                )
            else:
                await context.bot.send_message(
                    chat_id=query.message.chat.id,
                    text=message_text,
                    parse_mode=ParseMode.MARKDOWN
                )

                # Отправляем кнопки после попытки отправки медиа группы
                await context.bot.send_message(
                    chat_id=query.message.chat.id,
                    text="Выберите действие:",
                    reply_markup=reply_markup
                )
        else:
            await query.message.reply_text("Ошибка при получении данных отправителя.")
    else:
        await query.message.reply_text("Ошибка при получении деталей фронта работ.")


async def view_fronts_in_process(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    await update.callback_query.message.delete()

    user_chat_id = update.callback_query.from_user.id
    user_response = requests.get(f'{DJANGO_API_URL}users/chat/{user_chat_id}/')

    if user_response.status_code == 200:
        user_data = user_response.json()
        user_object_id = user_data.get('object_id')

        if user_object_id is None:
            await update.callback_query.message.reply_text("У вас нет назначенного объекта.")
            return

        response = requests.get(f'{DJANGO_API_URL}fronttransfers/?status=in_process')
        if response.status_code == 200:
            fronts = response.json()
            filtered_fronts = [front for front in fronts if front['object_id'] == user_object_id]

            if filtered_fronts:
                keyboard = []
                for front in filtered_fronts:
                    # Получаем имена объектов, видов работ, блоков/секций и организаций через sender_id
                    object_response = requests.get(f'{DJANGO_API_URL}objects/{front["object_id"]}/')
                    work_type_response = requests.get(f'{DJANGO_API_URL}worktypes/{front["work_type_id"]}/')
                    block_section_response = requests.get(f'{DJANGO_API_URL}blocksections/{front["block_section_id"]}/')
                    user_sender_response = requests.get(f'{DJANGO_API_URL}users/{front["sender_id"]}/')

                    if (object_response.status_code == 200 and
                            work_type_response.status_code == 200 and
                            block_section_response.status_code == 200 and
                            user_sender_response.status_code == 200):

                        object_name = object_response.json()['name']
                        work_type_name = work_type_response.json()['name']
                        block_section_name = block_section_response.json()['name']

                        sender_user_data = user_sender_response.json()
                        sender_organization_id = sender_user_data['organization_id']

                        organization_response = requests.get(f'{DJANGO_API_URL}organizations/{sender_organization_id}/')

                        if organization_response.status_code == 200:
                            organization_name = organization_response.json()['organization']

                            button_text = f"{organization_name} - {work_type_name} - {block_section_name} - этаж {front['floor']}"
                            callback_data = f"fronts_info_{front['id']}"
                            keyboard.append([InlineKeyboardButton(button_text, callback_data=callback_data)])
                        else:
                            await update.callback_query.message.reply_text("Ошибка при получении данных организации.")
                            return
                    else:
                        await update.callback_query.message.reply_text("Ошибка при получении данных. Попробуйте снова.")
                        return

                keyboard.append([InlineKeyboardButton("↻ Обновить", callback_data='fronts_in_process')])
                keyboard.append([InlineKeyboardButton("Назад", callback_data='frontbutton')])

                reply_markup = InlineKeyboardMarkup(keyboard)
                await update.callback_query.message.reply_text("Список текущих фронтов в работе:",
                                                               reply_markup=reply_markup)
            else:
                keyboard = []
                keyboard.append([InlineKeyboardButton("↻ Обновить", callback_data='fronts_in_process')])
                keyboard.append([InlineKeyboardButton("Назад", callback_data='frontbutton')])
                reply_markup = InlineKeyboardMarkup(keyboard)
                await update.callback_query.message.reply_text("Нет доступных фронтов работ со статусом 'в работе'.", reply_markup=reply_markup)
        else:
            await update.callback_query.message.reply_text(
                "Ошибка при получении списка фронтов работ. Попробуйте снова.")
    else:
        await update.callback_query.message.reply_text("Ошибка при получении данных пользователя. Попробуйте снова.")



async def show_front_info(update: Update, context: ContextTypes.DEFAULT_TYPE, front_id: int) -> None:
    query = update.callback_query
    await query.answer()

    # Удаление предыдущего сообщения, если оно существует
    if query.message:
        await query.message.delete()

    response = requests.get(f'{DJANGO_API_URL}fronttransfers/{front_id}/')
    if response.status_code == 200:
        front = response.json()
        sender_chat_id = front['sender_chat_id']  # Используем sender_chat_id

        sender_response = requests.get(f'{DJANGO_API_URL}users/chat/{sender_chat_id}/')
        if sender_response.status_code == 200:

            sender_response_id_org = sender_response.json()["organization_id"]
            org_response = requests.get(f'{DJANGO_API_URL}organizations/{sender_response_id_org}/').json()[
                "organization"]

            sender_full_name = sender_response.json()["full_name"]

            # Получаем имена объекта, вида работ и блока/секции
            object_response = requests.get(f'{DJANGO_API_URL}objects/{front["object_id"]}/')
            work_type_response = requests.get(f'{DJANGO_API_URL}worktypes/{front["work_type_id"]}/')
            block_section_response = requests.get(f'{DJANGO_API_URL}blocksections/{front["block_section_id"]}/')

            if object_response.status_code == 200 and work_type_response.status_code == 200 and block_section_response.status_code == 200:
                object_name = object_response.json()['name']
                work_type_name = work_type_response.json()['name']
                block_section_name = block_section_response.json()['name']
            else:
                await query.message.reply_text("Ошибка при получении данных. Попробуйте снова.")
                return

            message_text = (
                f"*Отправитель:* {sender_full_name}\n"
                f"*Организация:* {org_response}\n\n"
                f"*Объект:* {object_name}\n"
                f"*Вид работ:* {work_type_name}\n"
                f"*Блок/Секция:* {block_section_name}\n"
                f"*Этаж:* {front['floor']}\n\n"
                f"*Дата начала работ (МСК):* {datetime.fromisoformat(front['created_at']).strftime('%d.%m.%Y')}"
            )
            keyboard = []
            keyboard.append([InlineKeyboardButton("Назад", callback_data='fronts_in_process')])

            reply_markup = InlineKeyboardMarkup(keyboard)


            await context.bot.send_message(
                chat_id=query.message.chat.id,
                text=message_text,
                parse_mode=ParseMode.MARKDOWN,
                reply_markup=reply_markup
            )


        else:
            await query.message.reply_text("Ошибка при получении данных отправителя.")
    else:
        await query.message.reply_text("Ошибка при получении деталей фронта работ.")

async def handle_rework(query: Update, context: ContextTypes.DEFAULT_TYPE, front_id: int) -> None:
    await query.message.delete()
    await query.message.reply_text('Пожалуйста, введите причину доработки:')
    context.user_data['stage'] = f'rework_{front_id}'



#Сделать async
async def fetch(session, url):
    async with session.get(url) as response:
        return await response.json()


async def generate_pdf(front_id: int) -> str:
    API_URL = 'http://127.0.0.1:8000'
    async with aiohttp.ClientSession() as session:
        # Получение данных фронта
        front = await fetch(session, f'{API_URL}/fronttransfers/{front_id}')
        if not front:
            raise Exception(f'Ошибка при получении данных фронта: {front.status}')

        # Получение необходимых данных для замены
        object_name = (await fetch(session, f'{API_URL}/objects/{front["object_id"]}')).get('name', 'неизвестно')
        block_section_name = (await fetch(session, f'{API_URL}/blocksections/{front["block_section_id"]}')).get('name', 'неизвестно')
        boss_name = (await fetch(session, f'{API_URL}/users/{front["boss_id"]}')).get('full_name', 'неизвестно')
        receiver = await fetch(session, f'{API_URL}/users/{front["sender_id"]}')
        receiver_name = receiver.get('full_name', 'неизвестно')
        organization_name = (await fetch(session, f'{API_URL}/organizations/{receiver["organization_id"]}')).get('organization', 'неизвестно')
        work_type = (await fetch(session, f'{API_URL}/worktypes/{front["work_type_id"]}')).get('name', 'неизвестно')

        # Извлечение дня и месяца из поля approval_at
        approval_at = datetime.fromisoformat(front['approval_at'])
        day = approval_at.day
        months = ["января", "февраля", "марта", "апреля", "мая", "июня",
                  "июля", "августа", "сентября", "октября", "ноября", "декабря"]
        month = months[approval_at.month - 1]

        # Открытие DOCX-файла
        doc_path = os.path.abspath('PDF_акты/Акт_приема_передачи_фронта_работ_двухсторонний.docx')
        doc = Document(doc_path)

        # Замена плейсхолдеров в документе и установка размера шрифта
        def replace_placeholder(doc, placeholder, replacement):
            for paragraph in doc.paragraphs:
                if placeholder in paragraph.text:
                    for run in paragraph.runs:
                        if placeholder in run.text:
                            run.text = run.text.replace(placeholder, replacement)
                            run.font.size = Pt(10)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            if placeholder in paragraph.text:
                                for run in paragraph.runs:
                                    if placeholder in run.text:
                                        run.text = run.text.replace(placeholder, replacement)
                                        run.font.size = Pt(10)

        replace_placeholder(doc, 'objectname', object_name)
        replace_placeholder(doc, 'blocksectionid', block_section_name)
        replace_placeholder(doc, 'bossname', boss_name)
        replace_placeholder(doc, 'sendername', receiver_name)
        replace_placeholder(doc, 'floor', str(front['floor']))
        replace_placeholder(doc, 'orgname', organization_name)
        replace_placeholder(doc, 'day', str(day))
        replace_placeholder(doc, 'month', month)
        replace_placeholder(doc, 'worktype', work_type)

        # Сохранение обновленного документа в буфер
        temp_docx_path = os.path.abspath('PDF_акты/temp_updated_document.docx')
        doc.save(temp_docx_path)
        # print("Документ сохранен: ", temp_docx_path)

        random_number = random.randint(10000000, 99999999)
        # Конвертация DOCX в PDF с использованием LibreOffice
        pdf_output_name = f'{object_name}_{work_type}_{boss_name}_{receiver_name}_двухсторонний_{random_number}.pdf'
        pdf_output_path = os.path.abspath(f'PDF_акты/{pdf_output_name}')

        # Определяем команду для конвертации в зависимости от операционной системы
        if platform.system() == "Windows":
            libreoffice_path = "C:\\Program Files\\LibreOffice\\program\\soffice.exe"
        else:
            libreoffice_path = "libreoffice"  # для Linux предполагается, что LibreOffice доступен в PATH

        libreoffice_command = [
            libreoffice_path, '--headless', '--convert-to', 'pdf', '--outdir', os.path.dirname(pdf_output_path), temp_docx_path
        ]

        # Выполняем команду
        subprocess.run(libreoffice_command, check=True)

        # Переименовываем файл
        temp_pdf_path = os.path.join(os.path.dirname(pdf_output_path), 'temp_updated_document.pdf')
        if os.path.exists(temp_pdf_path):
            os.rename(temp_pdf_path, pdf_output_path)

        # print(f'Документ успешно обновлен и конвертирован в PDF и сохранен как {pdf_output_path}')
        return pdf_output_path

async def generate_pdf_reverse(front_id: int) -> str:
    API_URL = 'http://127.0.0.1:8000'
    async with aiohttp.ClientSession() as session:
        # Получение данных фронта
        front = await fetch(session, f'{API_URL}/fronttransfers/{front_id}')
        if not front:
            raise Exception(f'Ошибка при получении данных фронта: {front.status}')

        # Получение необходимых данных для замены
        object_name = (await fetch(session, f'{API_URL}/objects/{front["object_id"]}')).get('name', 'неизвестно')
        block_section_name = (await fetch(session, f'{API_URL}/blocksections/{front["block_section_id"]}')).get('name', 'неизвестно')
        boss_name = (await fetch(session, f'{API_URL}/users/{front["boss_id"]}')).get('full_name', 'неизвестно')
        receiver = await fetch(session, f'{API_URL}/users/{front["sender_id"]}')
        receiver_name = receiver.get('full_name', 'неизвестно')
        organization_name = (await fetch(session, f'{API_URL}/organizations/{receiver["organization_id"]}')).get('organization', 'неизвестно')
        work_type = (await fetch(session, f'{API_URL}/worktypes/{front["work_type_id"]}')).get('name', 'неизвестно')

        # Извлечение дня и месяца из поля approval_at
        approval_at = datetime.fromisoformat(front['approval_at'])
        day = approval_at.day
        months = ["января", "февраля", "марта", "апреля", "мая", "июня",
                  "июля", "августа", "сентября", "октября", "ноября", "декабря"]
        month = months[approval_at.month - 1]

        # Открытие DOCX-файла
        doc_path = os.path.abspath('PDF_акты/Акт_приема_передачи_фронта_работ_двухсторонний_reverse.docx')
        doc = Document(doc_path)

        # Замена плейсхолдеров в документе и установка размера шрифта
        def replace_placeholder(doc, placeholder, replacement):
            for paragraph in doc.paragraphs:
                if placeholder in paragraph.text:
                    for run in paragraph.runs:
                        if placeholder in run.text:
                            run.text = run.text.replace(placeholder, replacement)
                            run.font.size = Pt(10)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            if placeholder in paragraph.text:
                                for run in paragraph.runs:
                                    if placeholder in run.text:
                                        run.text = run.text.replace(placeholder, replacement)
                                        run.font.size = Pt(10)

        replace_placeholder(doc, 'objectname', object_name)
        replace_placeholder(doc, 'blocksectionid', block_section_name)
        replace_placeholder(doc, 'bossname', boss_name)
        replace_placeholder(doc, 'sendername', receiver_name)
        replace_placeholder(doc, 'floor', str(front['floor']))
        replace_placeholder(doc, 'orgname', organization_name)
        replace_placeholder(doc, 'day', str(day))
        replace_placeholder(doc, 'month', month)
        replace_placeholder(doc, 'worktype', work_type)

        # Сохранение обновленного документа в буфер
        temp_docx_path = os.path.abspath('PDF_акты/temp_updated_document_reverse.docx')
        doc.save(temp_docx_path)
        # print("Документ сохранен: ", temp_docx_path)

        random_number = random.randint(10000000, 99999999)
        # Конвертация DOCX в PDF с использованием LibreOffice
        pdf_output_name = f'{object_name}_{work_type}_{boss_name}_двусторонний_reverse_{random_number}.pdf'
        pdf_output_path = os.path.abspath(f'PDF_акты/{pdf_output_name}')

        # Определяем команду для конвертации в зависимости от операционной системы
        if platform.system() == "Windows":
            libreoffice_path = "C:\\Program Files\\LibreOffice\\program\\soffice.exe"
        else:
            libreoffice_path = "libreoffice"  # для Linux предполагается, что LibreOffice доступен в PATH

        libreoffice_command = [
            libreoffice_path, '--headless', '--convert-to', 'pdf', '--outdir', os.path.dirname(pdf_output_path), temp_docx_path
        ]

        # Выполняем команду
        subprocess.run(libreoffice_command, check=True)

        # Переименовываем файл
        temp_pdf_path = os.path.join(os.path.dirname(pdf_output_path), 'temp_updated_document_reverse.pdf')
        if os.path.exists(temp_pdf_path):
            os.rename(temp_pdf_path, pdf_output_path)

        # print(f'Документ успешно обновлен и конвертирован в PDF и сохранен как {pdf_output_path}')
        return pdf_output_path

async def approve_front(query: Update, context: ContextTypes.DEFAULT_TYPE, front_id: int) -> None:
    try:
        await query.message.delete()
        await query.message.reply_text(text="Создание документа, подождите...")
        await query.edit_message_reply_markup(reply_markup=None)
    except telegram.error.BadRequest:
        logger.warning("Message to delete not found")

    user_id = query.message.chat.id

    # Получаем данные фронта
    response = requests.get(f'{DJANGO_API_URL}fronttransfers/{front_id}/')
    if response.status_code == 200:
        front = response.json()

        # Получаем данные генерального подрядчика
        boss_response = requests.get(f'{DJANGO_API_URL}users/chat/{user_id}/')
        if boss_response.status_code == 200:
            boss_data = boss_response.json()
            boss_name = boss_data['full_name']
            boss_id = boss_data['id']

            # Обновляем статус фронта на "approved"
            update_data = {
                'status': 'approved',
                'approval_at': datetime.now().isoformat(),
                'boss_id': boss_id,
                'object_id': front['object_id'],
                'work_type_id': front['work_type_id'],
                'block_section_id': front['block_section_id'],
                'floor': front['floor'],
                'sender_id': front['sender_id'],
                'created_at': front['created_at'],
                'sender_chat_id': front['sender_chat_id'],
                'photo_ids': front['photo_ids'],
                'is_finish': front['is_finish']
            }
            response = requests.put(f'{DJANGO_API_URL}fronttransfers/{front_id}/', json=update_data)
            if response.status_code == 200:
                # Уведомляем отправителя
                sender_chat_id = front['sender_chat_id']

                # Получаем названия объекта, вида работ и блока/секции
                object_name = "неизвестно"
                block_section_name = "неизвестно"
                work_type_name = "неизвестно"

                if 'object_name' in front:
                    object_name = front['object_name']
                else:
                    object_response = requests.get(f'{DJANGO_API_URL}objects/{front["object_id"]}/')
                    if object_response.status_code == 200:
                        object_name = object_response.json().get('name', "неизвестно")

                if 'block_section_name' in front:
                    block_section_name = front['block_section_name']
                else:
                    block_section_response = requests.get(
                        f'{DJANGO_API_URL}blocksections/{front["block_section_id"]}/')
                    if block_section_response.status_code == 200:
                        block_section_name = block_section_response.json().get('name', "неизвестно")

                if 'work_type_name' in front:
                    work_type_name = front['work_type_name']
                else:
                    work_type_response = requests.get(f'{DJANGO_API_URL}worktypes/{front["work_type_id"]}/')
                    if work_type_response.status_code == 200:
                        work_type_name = work_type_response.json().get('name', "неизвестно")

                notification_text = (
                    f"\U00002705 Ваш фронт работ был принят генеральным подрядчиком *{boss_name}*:\n"
                    f"\n\n*Объект:* {object_name}\n"
                    f"*Секция/Блок:* {block_section_name}\n"
                    f"*Этаж:* {front.get('floor', 'неизвестно')}\n"
                    f"*Вид работ:* {work_type_name}\n"
                )
                await send_to_google_sheets(front_id, action='update')
                pdf_path = await generate_pdf(front_id) # Генерация PDF

                # Отправка PDF файла ген. директору
                await context.bot.send_document(
                    chat_id=user_id,
                    document=open(pdf_path, 'rb'),
                    caption='Фронт работ успешно принят.',
                    reply_markup=reply_markup_kb_main
                )



                await context.bot.send_document(
                    chat_id=sender_chat_id,
                    caption=notification_text,
                    parse_mode=ParseMode.MARKDOWN,
                    document=open(pdf_path, 'rb'),

                )
                # await query.message.reply_text('Фронт успешно принят.',reply_markup=reply_markup_kb_main)

                # Получаем данные пользователя для отображения главного меню
                response = requests.get(f'{DJANGO_API_URL}users/chat/{user_id}/')
                if response.status_code == 200:
                    user_data = response.json()
                    full_name = user_data.get('full_name', 'Пользователь')
                    organization_id = user_data.get('organization_id', None)
                    await send_main_menu(user_id, context, full_name, organization_id)
                else:
                    await query.message.reply_text('Ошибка при получении данных пользователя. Попробуйте снова.')
            else:
                await query.message.reply_text('Ошибка при обновлении статуса фронта. Попробуйте снова.')
        else:
            await query.message.reply_text('Ошибка при получении данных генерального подрядчика. Попробуйте снова.')
    else:
        await query.message.reply_text('Ошибка при получении данных фронта. Попробуйте снова.')


async def handle_transfer(query: Update, context: ContextTypes.DEFAULT_TYPE, front_id: int) -> None:
    await query.message.delete()
    context.user_data['transfer_front_id'] = front_id
    response = requests.get(f'{DJANGO_API_URL}organizations/')
    if response.status_code == 200:
        organizations = response.json()
        keyboard = [
            [InlineKeyboardButton(org['organization'], callback_data=f'transfer_org_{org["id"]}')] for org in organizations if org['organization'] != "БОС"
        ]
        reply_markup = InlineKeyboardMarkup(keyboard)
        await query.message.reply_text('Выберите организацию:', reply_markup=reply_markup)
        context.user_data['stage'] = 'choose_transfer_organization'
    else:
        await query.message.reply_text('Ошибка при получении списка организаций. Попробуйте снова.')


async def choose_transfer_user(query: Update, context: ContextTypes.DEFAULT_TYPE, org_id: int) -> None:
    await query.message.delete()
    context.user_data['transfer_org_id'] = org_id

    # Получаем ID пользователя, которому принадлежит фронт
    front_id = context.user_data['transfer_front_id']
    front_response = requests.get(f'{DJANGO_API_URL}fronttransfers/{front_id}/')
    if front_response.status_code == 200:
        front_data = front_response.json()
        current_user_id = front_data['sender_id']  # Получаем ID отправителя фронта

        # Получаем список пользователей и фильтруем их по организации
        response = requests.get(f'{DJANGO_API_URL}users/')
        if response.status_code == 200:
            users = response.json()

            # Фильтрация пользователей по организации
            filtered_users = [user for user in users if user['organization_id'] == org_id]

            # Исключаем отправителя фронта из списка пользователей
            # filtered_users = [user for user in filtered_users if user['id'] != current_user_id]

            if filtered_users:
                keyboard = [
                    [InlineKeyboardButton(user['full_name'], callback_data=f'transfer_user_{user["chat_id"]}')] for user in filtered_users
                ]
                reply_markup = InlineKeyboardMarkup(keyboard)
                await query.message.reply_text('Выберите пользователя:', reply_markup=reply_markup)
                context.user_data['stage'] = 'choose_transfer_user'
            else:
                await query.message.reply_text('В этой организации нет доступных пользователей.',
                                               reply_markup=InlineKeyboardMarkup([
                                                   [InlineKeyboardButton("Назад", callback_data='transfer_')]
                                               ]))
        else:
            await query.message.reply_text('Ошибка при получении списка пользователей. Попробуйте снова.')
    else:
        await query.message.reply_text('Ошибка при получении данных фронта. Попробуйте снова.')


async def choose_transfer_work_type(query: Update, context: ContextTypes.DEFAULT_TYPE, user_id: int) -> None:
    await query.message.delete()
    context.user_data['transfer_user_id'] = user_id

    # Получаем ID организации из контекста
    org_id = context.user_data.get('transfer_org_id')

    # Получаем ID фронта из контекста и объект по его ID
    front_id = context.user_data.get('transfer_front_id')
    front_response = requests.get(f'{DJANGO_API_URL}fronttransfers/{front_id}/')

    if front_response.status_code == 200:
        front_data = front_response.json()
        object_id = front_data['object_id']

        # Получаем пересечения видов работ для объекта и организации
        common_work_types_ids = await get_common_work_types(object_id, org_id)

        if common_work_types_ids:
            ids_query = "&".join([f"ids={id}" for id in common_work_types_ids])
            async with aiohttp.ClientSession() as session:
                async with session.get(f'{DJANGO_API_URL}worktypes/?{ids_query}') as response:
                    if response.status == 200:
                        work_types = await response.json()
                        keyboard = [
                            [InlineKeyboardButton(work['name'], callback_data=f'transfer_work_{work["id"]}')] for work
                            in work_types
                        ]
                        reply_markup = InlineKeyboardMarkup(keyboard)
                        await query.message.reply_text('Выберите новый вид работ:', reply_markup=reply_markup)
                        context.user_data['stage'] = 'choose_transfer_work_type'
                    else:
                        await query.message.reply_text('Ошибка при получении списка типов работ. Попробуйте снова.')
        else:
            await query.message.reply_text('Нет доступных видов работ для выбранного объекта и вашей организации.')
    else:
        await query.message.reply_text('Ошибка при получении данных фронта. Попробуйте снова.')


async def confirm_transfer(query: Update, context: ContextTypes.DEFAULT_TYPE, work_type_id: int) -> None:
    await query.message.delete()
    front_id = context.user_data['transfer_front_id']
    org_id = context.user_data['transfer_org_id']
    user_chat_id = context.user_data['transfer_user_id']  # Это chat_id пользователя

    # Используем правильный маршрут для получения данных пользователя по chat_id
    response = requests.get(f'{DJANGO_API_URL}users/chat/{user_chat_id}/')
    if response.status_code == 200:
        user_data = response.json()
        user_id = user_data['id']  # Получаем id пользователя из данных

        boss_response = requests.get(f'{DJANGO_API_URL}users/chat/{query.from_user.id}/')
        if boss_response.status_code == 200:
            boss_data = boss_response.json()
            boss_id = boss_data['id']

            # Получаем данные о виде работ по его ID
            work_type_response = requests.get(f'{DJANGO_API_URL}worktypes/{work_type_id}/')
            if work_type_response.status_code == 200:
                work_type_data = work_type_response.json()
                work_type_name = work_type_data['name']
                work_type_id_base = work_type_data['id']
            else:
                await query.message.reply_text('Ошибка при получении данных вида работ. Попробуйте снова.')
                return

            # Получаем текущие данные о фронте
            front_response = requests.get(f'{DJANGO_API_URL}fronttransfers/{front_id}/')
            if front_response.status_code == 200:
                front_data = front_response.json()

                # Получаем дополнительные данные об объекте, блоке/секции и виде работ
                object_response = requests.get(f'{DJANGO_API_URL}objects/{front_data["object_id"]}/')
                block_section_response = requests.get(
                    f'{DJANGO_API_URL}blocksections/{front_data["block_section_id"]}/')
                current_work_type_response = requests.get(f'{DJANGO_API_URL}worktypes/{front_data["work_type_id"]}/')

                # Получаем данные об организации отправителя
                sender_response = requests.get(f'{DJANGO_API_URL}users/{front_data["sender_id"]}/')
                if sender_response.status_code == 200:
                    sender_data = sender_response.json()
                    organization_id = sender_data['organization_id']
                    organization_response = requests.get(f'{DJANGO_API_URL}organizations/{organization_id}/')

                    if organization_response.status_code == 200:
                        organization_name = organization_response.json().get('organization', 'неизвестно')
                    else:
                        organization_name = 'неизвестно'

                    if object_response.status_code == 200:
                        object_name = object_response.json().get('name', 'неизвестно')
                    else:
                        object_name = 'неизвестно'

                    if block_section_response.status_code == 200:
                        block_section_name = block_section_response.json().get('name', 'неизвестно')
                    else:
                        block_section_name = 'неизвестно'

                    if current_work_type_response.status_code == 200:
                        current_work_type_name = current_work_type_response.json().get('name', 'неизвестно')
                    else:
                        current_work_type_name = 'неизвестно'

                    # Обновляем статус фронта на "on_consideration"
                    front_data.update({
                        'status': 'on_consideration',
                        'receiver_id': user_id,  # Используем id пользователя
                        'next_work_type_id': work_type_id_base,
                        'boss_id': boss_id
                    })

                    logger.info(
                        f"Отправка данных в API для обновления записи FrontTransfer: {json.dumps(front_data, indent=2)}")
                    response = requests.put(f'{DJANGO_API_URL}fronttransfers/{front_id}/', json=front_data)
                    logger.info(
                        f"Ответ от API при обновлении записи FrontTransfer: {response.status_code}, {response.text}")

                    if response.status_code == 200:
                        await query.message.reply_text('Фронт успешно передан на рассмотрение.')

                        # Отправляем уведомление получателю
                        transfer_response = requests.get(f'{DJANGO_API_URL}fronttransfers/{front_id}/')
                        if transfer_response.status_code == 200:
                            transfer = transfer_response.json()

                            sender_response = requests.get(f'{DJANGO_API_URL}users/chat/{transfer["sender_chat_id"]}/')
                            if sender_response.status_code == 200:
                                sender_data = sender_response.json()
                                sender_full_name = sender_data.get('full_name', 'Неизвестный отправитель')

                                message_text = (
                                    f"*Вам передан фронт работ*\n\n"
                                    f"*Отправитель:* {sender_full_name}\n"
                                    f"*Организация:* {organization_name}\n\n"
                                    f"*Объект:* {object_name}\n"
                                    f"*Вид работ:* {current_work_type_name}\n"
                                    f"*Блок/Секция:* {block_section_name}\n"
                                    f"*Этаж:* {transfer['floor']}\n\n"
                                    f"*Дата передачи (МСК):* {datetime.fromisoformat(transfer['created_at']).strftime('%d.%m.%Y')}\n"
                                    f"*Новый вид работ:* {work_type_name}"
                                )

                                keyboard = [
                                    [InlineKeyboardButton("\U00002705 Принять",
                                                          callback_data=f"accept_front_{front_id}")],
                                    [InlineKeyboardButton("\U0000274C Отклонить",
                                                          callback_data=f"decline_front_{front_id}")]
                                ]
                                reply_markup = InlineKeyboardMarkup(keyboard)

                                # Отправка фотографий
                                media_group = []
                                photo_ids = transfer.get('photo_ids', [])
                                for idx, photo_id in enumerate(photo_ids):
                                    if photo_id:
                                        if idx == 0:
                                            media_group.append(InputMediaPhoto(media=photo_id, caption=message_text,
                                                                               parse_mode=ParseMode.MARKDOWN))
                                        else:
                                            media_group.append(InputMediaPhoto(media=photo_id))

                                # Если есть фото, отправляем группу медиа
                                if media_group:
                                    await context.bot.send_media_group(chat_id=user_chat_id, media=media_group)
                                    await context.bot.send_message(
                                        chat_id=user_chat_id,
                                        text="Выберите действие:",
                                        reply_markup=reply_markup
                                    )
                                else:
                                    await context.bot.send_message(
                                        chat_id=user_chat_id,
                                        text=message_text,
                                        reply_markup=reply_markup,
                                        parse_mode=ParseMode.MARKDOWN
                                    )

                            else:
                                await query.message.reply_text('Ошибка при получении данных отправителя.')
                        else:
                            await query.message.reply_text('Ошибка при получении данных передачи фронта.')
                    else:
                        await query.message.reply_text(f'Ошибка при обновлении статуса фронта: {response.text}')
                else:
                    await query.message.reply_text('Ошибка при получении данных отправителя.')
            else:
                await query.message.reply_text('Ошибка при получении данных фронта. Попробуйте снова.')
        else:
            await query.message.reply_text('Ошибка при получении данных пользователя. Попробуйте снова.')
    else:
        await query.message.reply_text('Ошибка при получении данных пользователя. Попробуйте снова.')


async def fetch_data(session, url):
    async with session.get(url) as response:
        if response.status != 200:
            raise Exception(f'Ошибка при получении данных: {response.status}')
        return await response.json()

async def generate_pdf_triple(front_id: int) -> str:
    API_URL = 'http://127.0.0.1:8000'
    async with aiohttp.ClientSession() as session:
        # Получение данных фронта
        front = await fetch_data(session, f'{API_URL}/fronttransfers/{front_id}')
        if not front:
            raise Exception(f'Ошибка при получении данных фронта: {front.status}')

        # Получение необходимых данных для замены
        object_name = (await fetch_data(session, f'{API_URL}/objects/{front["object_id"]}')).get('name', 'неизвестно')
        block_section_name = (await fetch_data(session, f'{API_URL}/blocksections/{front["block_section_id"]}')).get('name', 'неизвестно')
        boss_name = (await fetch_data(session, f'{API_URL}/users/{front["boss_id"]}')).get('full_name', 'неизвестно')

        sender = await fetch_data(session, f'{API_URL}/users/{front["sender_id"]}')
        sender_name = sender.get('full_name', 'неизвестно')
        organization_id_sender = sender['organization_id']

        receiver = await fetch_data(session, f'{API_URL}/users/{front["receiver_id"]}')
        receiver_name = receiver.get('full_name', 'неизвестно')
        organization_id_receiver = receiver['organization_id']

        organization_name1 = (await fetch_data(session, f'{API_URL}/organizations/{organization_id_sender}')).get('organization', 'неизвестно')
        organization_name2 = (await fetch_data(session, f'{API_URL}/organizations/{organization_id_receiver}')).get('organization', 'неизвестно')
        work_type = (await fetch_data(session, f'{API_URL}/worktypes/{front["work_type_id"]}')).get('name', 'неизвестно')

        # Извлечение дня и месяца из поля approval_at
        approval_at = datetime.fromisoformat(front['approval_at'])
        day = approval_at.day

        # Русские названия месяцев
        months = ["января", "февраля", "марта", "апреля", "мая", "июня",
                  "июля", "августа", "сентября", "октября", "ноября", "декабря"]
        month = months[approval_at.month - 1]

        # Открытие DOCX-файла
        doc_path = os.path.abspath('PDF_акты/Акт_приема_передачи_фронта_работ_трехсторонний.docx')
        doc = Document(doc_path)

        # Замена плейсхолдеров в документе и установка размера шрифта
        def replace_placeholder(doc, placeholder, replacement):
            for paragraph in doc.paragraphs:
                if placeholder in paragraph.text:
                    for run in paragraph.runs:
                        if placeholder in run.text:
                            run.text = run.text.replace(placeholder, replacement)
                            run.font.size = Pt(10)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            if placeholder in paragraph.text:
                                for run in paragraph.runs:
                                    if placeholder in run.text:
                                        run.text = run.text.replace(placeholder, replacement)
                                        run.font.size = Pt(10)

        replace_placeholder(doc, 'objectname', object_name)
        replace_placeholder(doc, 'blocksectionid', block_section_name)
        replace_placeholder(doc, 'bossname', boss_name)
        replace_placeholder(doc, 'sendername', sender_name)
        replace_placeholder(doc, 'receivername', receiver_name)
        replace_placeholder(doc, 'floor', str(front['floor']))
        replace_placeholder(doc, 'orgname1', organization_name1)
        replace_placeholder(doc, 'orgname2', organization_name2)
        replace_placeholder(doc, 'day', str(day))
        replace_placeholder(doc, 'month', month)
        replace_placeholder(doc, 'worktype', work_type)

        # Сохранение обновленного документа в буфер
        temp_docx_path = os.path.abspath('PDF_акты/temp_updated_document_triple.docx')
        doc.save(temp_docx_path)
        print("Документ сохранен: ", temp_docx_path)

        random_number = random.randint(10000000, 99999999)
        # Конвертация DOCX в PDF с использованием LibreOffice
        pdf_output_name = f'{object_name}_{work_type}_{boss_name}_{sender_name}_{receiver_name}_трехсторонний_{random_number}.pdf'
        pdf_output_path = os.path.abspath(f'PDF_акты/{pdf_output_name}')

        # Определяем команду для конвертации в зависимости от операционной системы
        if platform.system() == "Windows":
            libreoffice_path = "C:\\Program Files\\LibreOffice\\program\\soffice.exe"
        else:
            libreoffice_path = "libreoffice"  # для Linux предполагается, что LibreOffice доступен в PATH

        libreoffice_command = [
            libreoffice_path, '--headless', '--convert-to', 'pdf', '--outdir', os.path.dirname(pdf_output_path), temp_docx_path
        ]

        # Выполняем команду
        subprocess.run(libreoffice_command, check=True)

        # Переименовываем файл
        temp_pdf_path = os.path.join(os.path.dirname(pdf_output_path), 'temp_updated_document_triple.pdf')
        if os.path.exists(temp_pdf_path):
            os.rename(temp_pdf_path, pdf_output_path)

        print(f'Документ успешно обновлен и конвертирован в PDF и сохранен как {pdf_output_path}')
        return pdf_output_path



async def get_data_from_api(session, endpoint):
    async with session.get(f"{API_BASE_URL}{endpoint}") as response:
        response.raise_for_status()
        return await response.json()

async def send_to_google_sheets(front_id, action='append'):
    async with aiohttp.ClientSession() as session:
        # Получение данных о фронте работ
        front_transfer = await get_data_from_api(session, f"/fronttransfers/{front_id}")

        # Получение дополнительных данных
        object_data = await get_data_from_api(session, f"/objects/{front_transfer['object_id']}")
        block_section_data = await get_data_from_api(session, f"/blocksections/{front_transfer['block_section_id']}")
        work_type_data = await get_data_from_api(session, f"/worktypes/{front_transfer['work_type_id']}")
        sender_data = await get_data_from_api(session, f"/users/{front_transfer['sender_id']}")
        organization_data = await get_data_from_api(session, f"/organizations/{sender_data['organization_id']}")

        # Формирование данных для отправки в Google Sheets
        created_at = parser.parse(front_transfer['created_at'])
        approval_at = parser.parse(front_transfer['approval_at']) if front_transfer['approval_at'] else datetime.now()
        days_diff = (approval_at - created_at).days
        current_time = datetime.now().strftime("%d.%m.%Y %H:%M:%S")

        data = {
            "action": action,  # Добавляем тип действия
            "A": front_transfer['id'],
            "B": object_data['name'],
            "C": block_section_data['name'],
            "D": front_transfer['floor'],
            "E": work_type_data['name'],
            "F": created_at.strftime("%d.%m.%Y %H:%M:%S"),
            "G": approval_at.strftime("%d.%m.%Y %H:%M:%S"),
            "H": days_diff,
            "I": created_at != approval_at,
            "J": organization_data['organization'],
            "K": front_transfer['remarks'],
            "L": current_time,
            "M": front_transfer['is_finish']
        }

        async with session.post(WEBHOOK_URL, json=data) as response:
            response.raise_for_status()
            return await response.json()

async def send_workforce_to_google_sheets(object_name, block_section_name, floor, work_type_name, organization_name, workforce_count, id_workforce):
    async with aiohttp.ClientSession() as session:
        data = {
            "action": "workforce",  # Действие для численности
            "A": object_name,
            "B": block_section_name,
            "C": floor,
            "D": work_type_name,
            "E": organization_name,
            "F": workforce_count,
            "G": datetime.now().strftime("%d.%m.%Y"),
            "H": id_workforce,
        }
        print("Отправляемый JSON:", data)
        async with session.post(WEBHOOK_URL, json=data) as response:
            response.raise_for_status()
            return await response.json()

async def accept_front(query: Update, context: ContextTypes.DEFAULT_TYPE, front_id: int) -> None:
    user_id = query.message.chat.id
    await query.edit_message_reply_markup(reply_markup=None)
    response = requests.get(f'{DJANGO_API_URL}fronttransfers/{front_id}/')
    if response.status_code == 200:
        front = response.json()
        user_chat_id = str(query.from_user.id)  # Преобразуем в строку

        # Проверка наличия ключа 'receiver_id' в данных фронта
        if 'receiver_id' not in front:
            await query.message.reply_text('Ошибка: получатель фронта не указан.')
            return

        # Если next_work_type_id пустой, изменяем статус принимаемого фронта на in_process
        if not front['next_work_type_id']:
            await query.edit_message_text(text="Создание документа, подождите...")
            # Обновляем данные фронта
            old_front_data = {
                'sender_id': front['sender_id'],
                'object_id': front['object_id'],
                'work_type_id': front['work_type_id'],
                'block_section_id': front['block_section_id'],
                'floor': front['floor'],
                'status': 'in_process',  # Обновляем статус на in_process
                'photo1': front['photo1'],
                'photo2': front['photo2'],
                'photo3': front['photo3'],
                'photo4': front['photo4'],
                'photo5': front['photo5'],
                'receiver_id': front['receiver_id'],
                'remarks': front['remarks'],
                'next_work_type_id': front['next_work_type_id'],
                'boss_id': front['boss_id'],
                'created_at': front['created_at'],
                'approval_at': front['approval_at'],
                'photo_ids': front['photo_ids'],
                'sender_chat_id': front['sender_chat_id'],
                'is_finish': front['is_finish']
            }

            response = requests.put(f'{DJANGO_API_URL}fronttransfers/{front_id}/', json=old_front_data)
            if response.status_code == 200:


                # Получаем chat_id босса по его id
                boss_id = front['boss_id']
                boss_response = requests.get(f'{DJANGO_API_URL}users/{boss_id}/')
                if boss_response.status_code == 200:
                    boss_chat_id = boss_response.json()['chat_id']
                else:
                    await query.message.reply_text('Ошибка при получении chat_id ген подрядчика.')
                    return

                # Получаем chat_id создателя по его id
                sender_id = front['sender_id']
                sender_response = requests.get(f'{DJANGO_API_URL}users/{sender_id}/')
                if sender_response.status_code == 200:
                    sender_chat_id = sender_response.json()['chat_id']
                else:
                    await query.message.reply_text('Ошибка при получении chat_id создателя фронта.')
                    return

                # Получение данных для уведомления
                object_response = requests.get(f'{DJANGO_API_URL}objects/{front["object_id"]}/')
                block_section_response = requests.get(f'{DJANGO_API_URL}blocksections/{front["block_section_id"]}/')
                work_type_response = requests.get(f'{DJANGO_API_URL}worktypes/{front["work_type_id"]}/')

                if object_response.status_code == 200:
                    object_name = object_response.json().get('name', 'Неизвестный объект')
                else:
                    object_name = 'Неизвестный объект'

                if block_section_response.status_code == 200:
                    block_section_name = block_section_response.json().get('name', 'Неизвестный блок/секция')
                else:
                    block_section_name = 'Неизвестный блок/секция'

                if work_type_response.status_code == 200:
                    work_type_name = work_type_response.json().get('name', 'Неизвестный вид работ')
                else:
                    work_type_name = 'Неизвестный вид работ'

                pdf_path = await generate_pdf_reverse(front_id)

                notification_text = (
                    f"\U00002705 Фронт работ успешно принят:"
                    f"\n\n*Объект:* {object_name}\n"
                    f"*Секция/Блок:* {block_section_name}\n"
                    f"*Этаж:* {front['floor']}\n"
                    f"*Вид работ:* {work_type_name}\n\n"

                )

                await context.bot.send_document(
                    chat_id=boss_chat_id,
                    caption=notification_text,
                    parse_mode=ParseMode.MARKDOWN,
                    document=open(pdf_path, 'rb')
                )
                await context.bot.send_document(
                    chat_id=user_id,
                    document=open(pdf_path, 'rb'),
                    caption='Фронт успешно принят. Нажмите /start для быстрой передачи фронт работ',
                    reply_markup=reply_markup_kb_main
                )
                await send_to_google_sheets(front_id)


            else:
                await query.message.reply_text(f'Ошибка при обновлении статуса фронта: {response.text}', reply_markup=reply_markup_kb_main)
            return

        # Создание нового фронта с новым статусом, если next_work_type_id не пустой
        new_front_data = {
            'sender_id': front['receiver_id'],  # Принимающий становится отправителем
            'sender_chat_id': user_chat_id,
            'object_id': front['object_id'],
            'work_type_id': front['next_work_type_id'],
            'block_section_id': front['block_section_id'],
            'floor': front['floor'],
            'status': 'in_process',
            'created_at': datetime.now().isoformat(),
            'approval_at': datetime.now().isoformat(),
            'photo1': None,
            'photo2': None,
            'photo3': None,
            'photo4': None,
            'photo5': None,
            'receiver_id': None,
            'remarks': None,
            'next_work_type_id': None,
            'boss_id': None,
            'photo_ids': [],
            'is_finish': False
        }
        await query.edit_message_text(text="Создание документа, подождите...")

        response = requests.post(f'{DJANGO_API_URL}fronttransfers/', json=new_front_data)
        if response.status_code == 200:
            new_front = response.json()  # Получаем данные нового фронта, включая его ID
            new_front_id = new_front['id']
            # Получаем все старые данные от фронта
            old_front_data = {
                'sender_id': front['sender_id'],
                'object_id': front['object_id'],
                'work_type_id': front['work_type_id'],
                'block_section_id': front['block_section_id'],
                'floor': front['floor'],
                'status': 'approved',  # Обновляем статус
                'photo1': front['photo1'],
                'photo2': front['photo2'],
                'photo3': front['photo3'],
                'photo4': front['photo4'],
                'photo5': front['photo5'],
                'receiver_id': front['receiver_id'],
                'remarks': front['remarks'],
                'next_work_type_id': front['next_work_type_id'],
                'boss_id': front['boss_id'],
                'created_at': front['created_at'],
                'approval_at': datetime.now().isoformat(),  # Обновляем дату
                'photo_ids': front['photo_ids'],
                'sender_chat_id': front['sender_chat_id'],
                'is_finish': front['is_finish']
            }

            response = requests.put(f'{DJANGO_API_URL}fronttransfers/{front_id}/', json=old_front_data)
            if response.status_code == 200:

                # Получаем chat_id босса по его id
                boss_id = front['boss_id']
                boss_response = requests.get(f'{DJANGO_API_URL}users/{boss_id}/')
                if boss_response.status_code == 200:
                    boss_chat_id = boss_response.json()['chat_id']
                else:
                    await query.message.reply_text('Ошибка при получении chat_id ген подрядчика.')
                    return

                # Получаем chat_id создателя по его id
                sender_id = front['sender_id']
                sender_response = requests.get(f'{DJANGO_API_URL}users/{sender_id}/')
                if sender_response.status_code == 200:
                    sender_chat_id = sender_response.json()['chat_id']
                else:
                    await query.message.reply_text('Ошибка при получении chat_id создателя фронта.')
                    return

                # Получение данных для уведомления
                object_response = requests.get(f'{DJANGO_API_URL}objects/{front["object_id"]}/')
                block_section_response = requests.get(f'{DJANGO_API_URL}blocksections/{front["block_section_id"]}/')
                work_type_response = requests.get(f'{DJANGO_API_URL}worktypes/{front["work_type_id"]}/')

                if object_response.status_code == 200:
                    object_name = object_response.json().get('name', 'Неизвестный объект')
                else:
                    object_name = 'Неизвестный объект'

                if block_section_response.status_code == 200:
                    block_section_name = block_section_response.json().get('name', 'Неизвестный блок/секция')
                else:
                    block_section_name = 'Неизвестный блок/секция'

                if work_type_response.status_code == 200:
                    work_type_name = work_type_response.json().get('name', 'Неизвестный вид работ')
                else:
                    work_type_name = 'Неизвестный вид работ'

                pdf_path = await generate_pdf_triple(front_id)

                notification_text = (
                    f"\U00002705 Фронт работ успешно принят:"
                    f"\n\n*Объект:* {object_name}\n"
                    f"*Секция/Блок:* {block_section_name}\n"
                    f"*Этаж:* {front['floor']}\n"
                    f"*Вид работ:* {work_type_name}\n"
                )

                # Отправка уведомлений в зависимости от условия
                if sender_id == front['receiver_id']:
                    await context.bot.send_document(
                        chat_id=boss_chat_id,
                        caption=notification_text,
                        parse_mode=ParseMode.MARKDOWN,
                        document=open(pdf_path, 'rb')
                    )
                    await context.bot.send_document(
                        chat_id=user_id,
                        document=open(pdf_path, 'rb'),
                        caption='Фронт успешно принят. Нажмите /start для быстрой передачи фронт работ',
                        reply_markup=reply_markup_kb_main
                    )
                else:
                    await context.bot.send_document(
                        chat_id=boss_chat_id,
                        caption=notification_text,
                        parse_mode=ParseMode.MARKDOWN,
                        document=open(pdf_path, 'rb')
                    )
                    await context.bot.send_document(
                        chat_id=user_id,
                        document=open(pdf_path, 'rb'),
                        caption='Фронт успешно принят. Нажмите /start для быстрой передачи фронт работ',
                        reply_markup=reply_markup_kb_main
                    )
                    await context.bot.send_document(
                        chat_id=sender_chat_id,
                        caption=notification_text,
                        parse_mode=ParseMode.MARKDOWN,
                        document=open(pdf_path, 'rb'),
                    )

                await send_to_google_sheets(front_id, action='update')  # Обновляем старый фронт

                await send_to_google_sheets(new_front_id, action='append')

            else:
                await query.message.reply_text(f'Ошибка при обновлении статуса старого фронта: {response.text}', reply_markup=reply_markup_kb_main)
        else:
            await query.message.reply_text(f'Ошибка при создании нового фронта: {response.text}', reply_markup=reply_markup_kb_main)
    else:
        await query.message.reply_text('Ошибка при получении данных фронта. Попробуйте снова.', reply_markup=reply_markup_kb_main)



async def decline_front(query: Update, context: ContextTypes.DEFAULT_TYPE, front_id: int) -> None:
    await query.edit_message_reply_markup(reply_markup=None)
    await query.message.reply_text('Пожалуйста, введите причину отклонения:')
    context.user_data['stage'] = f'decline_{front_id}'


async def notify_general_contractors(context: ContextTypes.DEFAULT_TYPE, transfer_data: dict) -> None:
    logger.info("Начало выполнения notify_general_contractors")
    response = requests.get(f'{DJANGO_API_URL}users/?organization=3')
    if response.status_code == 200:
        all_users = response.json()

        # Фильтруем пользователей с organization_id = 3 и совпадающим object_id
        print(transfer_data)
        front_object_id = transfer_data['object_id']
        general_contractors = [
            user for user in all_users
            if user.get('organization_id') == 3 and user.get('object_id') == front_object_id
        ]

        if not general_contractors:
            logger.info("Нет генеральных подрядчиков для уведомления.")
            return

        sender_chat_id = transfer_data["sender_chat_id"]
        sender_response = requests.get(f'{DJANGO_API_URL}users/chat/{sender_chat_id}/')



        if sender_response.status_code == 200:
            sender_data = sender_response.json()
            sender_full_name = sender_data.get('full_name', 'Неизвестный отправитель')
            org_response = requests.get(f'{DJANGO_API_URL}organizations/{sender_data["organization_id"]}/').json()
            sender_organization = org_response["organization"]


            logger.info(f"Уведомление подрядчиков, количество: {len(general_contractors)}")
            for contractor in general_contractors:
                chat_id = contractor['chat_id']
                keyboard = [
                    [InlineKeyboardButton("Просмотр переданных фронтов ", callback_data='view_fronts')]
                ]
                reply_markup = InlineKeyboardMarkup(keyboard)

                message_text = (
                    f"*Передан фронт работ*\n\n"
                    f"*Отправитель:* {sender_full_name}\n"
                    f"*Организация:* {sender_organization}\n\n"
                    f"*Объект:* {transfer_data['object_name']}\n"
                    f"*Вид работ:* {transfer_data['work_type_name']}\n"
                    f"*Блок/Секция:* {transfer_data['block_section_name']}\n"
                    f"*Этаж:* {transfer_data['floor']}"
                )
                logger.info(f"Отправка сообщения подрядчику с chat_id: {chat_id}")
                await context.bot.send_message(
                    chat_id=chat_id,
                    text=message_text,
                    reply_markup=reply_markup,
                    parse_mode=telegram.constants.ParseMode.MARKDOWN
                )
        else:
            logger.error(f"Ошибка при получении данных отправителя: {sender_response.status_code}")
    else:
        logger.error(f"Ошибка при получении списка пользователей: {response.status_code}")


async def choose_existing_front(query: Update, context: ContextTypes.DEFAULT_TYPE, fronts: list, object_id: int) -> None:
    await query.message.delete()  # Удаление предыдущего сообщения

    user_id = query.from_user.id
    filtered_fronts = [front for front in fronts if front['sender_chat_id'] == str(user_id) and front['object_id'] == object_id]

    if not filtered_fronts:
        await query.message.reply_text("Нет фронтов для продолжения передачи на выбранном объекте.")
        return

    keyboard = []
    for front in filtered_fronts:
        # Получаем названия объектов, видов работ и блоков/секции
        object_name = "неизвестно"
        work_type_name = "неизвестно"
        block_section_name = "неизвестно"

        # Если данные уже есть в ответе, используем их
        if 'object_name' in front:
            object_name = front['object_name']
        else:
            object_response = requests.get(f'{DJANGO_API_URL}objects/{front["object_id"]}/')
            if object_response.status_code == 200:
                object_name = object_response.json().get('name', "неизвестно")

        if 'work_type_name' in front:
            work_type_name = front['work_type_name']
        else:
            work_type_response = requests.get(f'{DJANGO_API_URL}worktypes/{front["work_type_id"]}/')
            if work_type_response.status_code == 200:
                work_type_name = work_type_response.json().get('name', "неизвестно")

        if 'block_section_name' in front:
            block_section_name = front['block_section_name']
        else:
            block_section_response = requests.get(f'{DJANGO_API_URL}blocksections/{front["block_section_id"]}/')
            if block_section_response.status_code == 200:
                block_section_name = block_section_response.json().get('name', "неизвестно")

        button_text = f"{object_name} - {work_type_name} - {block_section_name} - Этаж {front['floor']}"
        callback_data = f"existing_front_{front['id']}"
        keyboard.append([InlineKeyboardButton(button_text, callback_data=callback_data)])

    keyboard.append([InlineKeyboardButton("Назад", callback_data='front_menu')])

    reply_markup = InlineKeyboardMarkup(keyboard)
    await query.message.reply_text('Выберите фронт для продолжения передачи:', reply_markup=reply_markup)
    context.user_data['stage'] = 'choose_existing_front'
    context.user_data['object_id'] = object_id



async def handle_existing_front_selection(query: Update, context: ContextTypes.DEFAULT_TYPE, front_id: int) -> None:
    await query.message.delete()  # Удаление предыдущего сообщения
    context.user_data['transfer_id'] = front_id
    context.user_data['photos'] = []  # Сброс списка фотографий
    context.user_data.pop('last_photo_message_id', None)  # Сброс идентификатора последнего сообщения

    # Получаем текущие данные о фронте
    front_response = requests.get(f'{DJANGO_API_URL}fronttransfers/{front_id}/')
    if front_response.status_code == 200:
        front_data = front_response.json()
        # Обновляем статус фронта на "transferred"
        front_data['status'] = 'transferred'
        # front_data['is_finish'] = True  # Добавляем поле is_finish


        response = requests.put(f'{DJANGO_API_URL}fronttransfers/{front_id}/', json=front_data)
        logger.info(f"Ответ от API при обновлении статуса фронта: {response.status_code}, {response.text}")

        if response.status_code == 200:
            # Уведомление ген подрядчиков
            transfer_data = {
                'object_name': front_data.get('object_name', 'неизвестно'),
                'work_type_name': front_data.get('work_type_name', 'неизвестно'),
                'block_section_name': front_data.get('block_section_name', 'неизвестно'),
                'floor': front_data['floor'],
                'sender_chat_id': front_data['sender_chat_id']
            }
            # await notify_general_contractors(context, transfer_data)

            await query.message.reply_text(
                'Продолжите передачу фронта. Пожалуйста, прикрепите фотографии (до 10 штук) или нажмите /done:')
            context.user_data['stage'] = 'attach_photos'
        else:
            await query.message.reply_text('Ошибка при обновлении статуса фронта. Попробуйте снова.')
    else:
        await query.message.reply_text('Ошибка при получении данных фронта. Попробуйте снова.')




async def handle_delete_error(query: Update, context: ContextTypes.DEFAULT_TYPE, front_id: int) -> None:
    await query.message.reply_text('Пожалуйста, введите комментарий для удаления фронта:')
    context.user_data['stage'] = f'delete_error_{front_id}'

# Функция для получения пересечения work_types_ids
async def get_common_work_types(object_id: int, org_id: int) -> List[int]:
    async with aiohttp.ClientSession() as session:
        async with session.get(f'{DJANGO_API_URL}objects/{object_id}') as response_obj, \
                   session.get(f'{DJANGO_API_URL}organizations/{org_id}') as response_org:

            if response_obj.status == 200 and response_org.status == 200:
                object_data = await response_obj.json()
                organization_data = await response_org.json()
                object_work_types = set(object_data.get('work_types_ids', []))
                organization_work_types = set(organization_data.get('work_types_ids', []))
                common_work_types = list(object_work_types.intersection(organization_work_types))
                return common_work_types
            else:
                return []


async def handle_transfer_workforce(query: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    user_chat_id = query.from_user.id
    await query.message.delete()

    # Получаем данные организации пользователя
    user_response = requests.get(f'{DJANGO_API_URL}users/chat/{user_chat_id}/')
    if user_response.status_code == 200:
        user_data = user_response.json()
        organization_id = user_data['organization_id']
        context.user_data['organization_id'] = organization_id

        # Получаем данные организации, чтобы узнать object_ids
        organization_response = requests.get(f'{DJANGO_API_URL}organizations/{organization_id}/')
        if organization_response.status_code == 200:
            organization_data = organization_response.json()
            organization_object_ids = organization_data.get('object_ids', [])
        else:
            await query.message.reply_text('Ошибка при получении данных организации. Попробуйте снова.')
            return
    else:
        await query.message.reply_text('Ошибка при получении данных пользователя. Попробуйте снова.')
        return

    # Получаем объекты
    response = requests.get(f'{DJANGO_API_URL}objects/')
    if response.status_code == 200:
        objects = response.json()
        filtered_objects = [obj for obj in objects if obj['id'] in organization_object_ids]
        if filtered_objects:
            keyboard = [
                [InlineKeyboardButton(obj['name'], callback_data=f'workforce_obj_{obj["id"]}')] for obj in filtered_objects
            ]
            keyboard.append([InlineKeyboardButton("Назад", callback_data='workforce_menu')])
            reply_markup = InlineKeyboardMarkup(keyboard)
            await query.message.reply_text('Выберите объект:', reply_markup=reply_markup)
        else:
            await query.message.reply_text('Нет доступных объектов для данной организации.')
    else:
        await query.message.reply_text('Ошибка при получении списка объектов. Попробуйте снова.')



async def update_workforce_in_google_sheets(workforce_id, object_id, block_section_id, floor, work_type_id, organization_id, new_workforce_count):
    async with aiohttp.ClientSession() as session:
        # Получение дополнительных данных
        object_data = await get_data_from_api(session, f"/objects/{object_id}")
        block_section_data = await get_data_from_api(session, f"/blocksections/{block_section_id}")
        work_type_data = await get_data_from_api(session, f"/worktypes/{work_type_id}")
        organization_data = await get_data_from_api(session, f"/organizations/{organization_id}")

        data = {
            "action": "update_workforce",
            "A": object_data['name'],
            "B": block_section_data['name'],
            "C": floor,
            "D": work_type_data['name'],
            "E": organization_data['organization'],
            "F": new_workforce_count,
            "G": datetime.now().strftime("%d.%m.%Y"),
            "H": workforce_id,  # Используется для поиска записи в Google Sheets
        }

        async with session.post(WEBHOOK_URL, json=data) as response:
            response.raise_for_status()
            return await response.json()


async def handle_workforce_count(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    if context.user_data.get('expecting_workforce_count'):
        try:
            workforce_count = int(update.message.text)
            object_id = context.user_data['workforce_object_id']
            block_section_id = context.user_data['workforce_block_section_id']
            floor = context.user_data['workforce_floor']
            if floor == 'roof':
                floor = 'Кровля'
            work_type_id = context.user_data['workforce_work_type_id']
            organization_id = context.user_data['organization_id']
            user_id = update.message.from_user.id

            # Получаем id пользовтаеля в базе
            user_response = requests.get(f'{DJANGO_API_URL}users/chat/{user_id}/')
            if user_response.status_code == 200:
                user_name = user_response.json()['id']
            else:
                await update.message.reply_text('Ошибка при получении id пользователя')
                return

            # Получаем имя блока/секции
            block_section_response = requests.get(f'{DJANGO_API_URL}blocksections/{block_section_id}/')
            if block_section_response.status_code == 200:
                block_section_name = block_section_response.json()['name']
                context.user_data['workforce_block_section_name'] = block_section_name
            else:
                await update.message.reply_text('Ошибка при получении данных блока/секции. Попробуйте снова.')
                return

            # Получаем имя объекта
            object_response = requests.get(f'{DJANGO_API_URL}objects/{object_id}/')
            if object_response.status_code == 200:
                object_name = object_response.json()['name']
            else:
                await update.message.reply_text('Ошибка при получении данных объекта. Попробуйте снова.')
                return

            # Получаем имя вида работ
            work_type_response = requests.get(f'{DJANGO_API_URL}worktypes/{work_type_id}/')
            if work_type_response.status_code == 200:
                work_type_name = work_type_response.json()['name']
                context.user_data['workforce_work_type_name'] = work_type_name
            else:
                await update.message.reply_text('Ошибка при получении данных вида работ. Попробуйте снова.')
                return

            # Получаем имя организации
            organization_response = requests.get(f'{DJANGO_API_URL}organizations/{organization_id}/')
            if organization_response.status_code == 200:
                organization_name = organization_response.json()['organization']
            else:
                await update.message.reply_text('Ошибка при получении данных организации. Попробуйте снова.')
                return

            workforce_data = {
                'object_id': object_id,
                'block_section_id': block_section_id,
                'floor': floor,
                'work_type_id': work_type_id,
                'organization_id': organization_id,
                'workforce_count': workforce_count,
                'date': datetime.now().isoformat(),
                'user_id': user_name
            }

            response = requests.post(f'{DJANGO_API_URL}frontworkforces/', json=workforce_data)

            if response.status_code == 200:
                id_workforce = response.json()['id']
                await update.message.reply_text(
                    f"Численность успешно передана на вид работ: \n"
                    f"{context.user_data['workforce_work_type_name']}\n"
                    f"{context.user_data['workforce_block_section_name']} Этаж {floor}",
                    reply_markup=InlineKeyboardMarkup(
                        [[InlineKeyboardButton(f"Повторить для {context.user_data['workforce_work_type_name']}",
                                               callback_data='repeat_workforce')],
                         [InlineKeyboardButton("Назад", callback_data='main_menu')]]
                    )
                )
                # Отправка данных в Google Sheets
                asyncio.create_task(send_workforce_to_google_sheets(
                    object_name,
                    block_section_name,
                    floor,
                    work_type_name,
                    organization_name,
                    workforce_count,
                    id_workforce,
                ))
            else:
                await update.message.reply_text(f'Ошибка при передаче численности: {response.text}')
        except ValueError:
            await update.message.reply_text('Пожалуйста, введите корректное число.')

        context.user_data['expecting_workforce_count'] = False

    elif context.user_data.get('expecting_new_workforce_count'):
        try:
            new_workforce_count = int(update.message.text)
            workforce_id = context.user_data['workforce_id_to_refactor']

            # Получаем данные о текущей записи
            response = requests.get(f'{DJANGO_API_URL}frontworkforces/{workforce_id}/')
            if response.status_code == 200:
                workforce_data = response.json()

                # Обновляем данные в базе данных
                workforce_data['workforce_count'] = new_workforce_count
                update_response = requests.patch(f'{DJANGO_API_URL}frontworkforces/{workforce_id}/',
                                                 json=workforce_data)

                print(update_response.json())
                if update_response.status_code == 200:
                    if 'refactor_message_id' in context.user_data:
                        try:
                            await context.bot.delete_message(
                                chat_id=update.message.chat.id,
                                message_id=context.user_data['refactor_message_id']
                            )
                        except:
                            pass
                    await update.message.reply_text('Численность успешно обновлена.')

                    # Получение данных для вызова send_main_menu
                    user_data = requests.get(f'{DJANGO_API_URL}users/chat/{update.message.chat.id}/').json()
                    full_name = user_data.get('full_name', 'Пользователь')
                    organization_id = user_data.get('organization_id', None)
                    await send_main_menu(update.message.chat.id, context, full_name, organization_id)

                    # Отправка данных в Google Sheets
                    asyncio.create_task(update_workforce_in_google_sheets(
                        workforce_data['id'],
                        workforce_data['object_id'],
                        workforce_data['block_section_id'],
                        workforce_data['floor'],
                        workforce_data['work_type_id'],
                        workforce_data['organization_id'],
                        new_workforce_count
                    ))

                else:
                    await update.message.reply_text('Ошибка при обновлении численности в базе данных.')
            else:
                await update.message.reply_text('Ошибка при получении данных численности.')
        except ValueError:
            await update.message.reply_text('Пожалуйста, введите корректное число.')

        context.user_data['expecting_new_workforce_count'] = False

    elif context.user_data.get('expecting_volume_count'):
        await handle_volume_count(update, context)

    elif context.user_data.get('expecting_new_volume_count'):
        await handle_new_volume_count(update, context)

    elif context.user_data.get('expecting_prefab_quantity'):
        await handle_prefab_quantity(update, context)

    elif context.user_data.get('expecting_sgp_quantity'):
        await handle_prefab_sgp_quantity(update, context)

    elif context.user_data.get('expecting_shipment_quantity'):
        await handle_prefab_shipment_quantity(update, context)

    elif context.user_data.get('expecting_stock_quantity'):
        await handle_prefab_stock_quantity(update, context)

    elif context.user_data.get('stage') == 'attach_comments_prefab_in_work':
        await handle_prefab_comment(update, context)

    elif context.user_data.get('expecting_montage_quantity'):
        await handle_montage_quantity(update, context)

    elif context.user_data.get('stage') == 'support_question':
        await handle_support_question(update, context)

    elif context.user_data.get('stage') == 'support_answer':
        await handle_support_answer(update, context)

    elif context.user_data.get('expecting_remark_quantity'):
        await handle_remark_quantity(update, context)

    elif context.user_data.get('expecting_new_status_quantity'):
        await remarks_prefab_newstatus(update, context)

    elif context.user_data.get('expecting_new_status'):
        await handle_new_status(update, context)

    elif context.user_data.get('refactor_prefab_count'):
        await handle_refactor_prefab_quantity(update, context)

    else:
        await handle_message(update, context)



async def view_today_workforce(query: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    user_data = requests.get(f'{DJANGO_API_URL}users/chat/{query.message.chat.id}/').json()
    user_object_id = user_data.get('object_id')
    today = datetime.now().date()

    response = requests.get(f'{DJANGO_API_URL}frontworkforces/')
    if response.status_code == 200:
        workforces = response.json()
        today_workforces = [wf for wf in workforces if parser.parse(wf['date']).date() == today and wf['object_id'] == user_object_id]
        if today_workforces:
            message = "\U0001F477 *Численность за сегодня:*\n"
            for wf in today_workforces:
                block_section_name = requests.get(f'{DJANGO_API_URL}blocksections/{wf["block_section_id"]}/').json()['name']
                work_type_name = requests.get(f'{DJANGO_API_URL}worktypes/{wf["work_type_id"]}/').json()['name']
                organization_name = requests.get(f'{DJANGO_API_URL}organizations/{wf["organization_id"]}/').json()['organization']
                message += (
                    f"{wf['workforce_count']} чел. — {work_type_name} — {block_section_name} — Этаж {wf['floor']} — {organization_name}\n\n"
                )
            await query.message.reply_text(
                text=message,
                parse_mode=ParseMode.MARKDOWN
            )
        else:
            await query.message.reply_text("Сегодня численность не передавалась.")
    else:
        await query.message.reply_text('Ошибка при получении данных. Попробуйте снова.')

    full_name = user_data.get('full_name', 'Пользователь')
    organization_id = user_data.get('organization_id', None)
    await send_main_menu(query.message.chat.id, context, full_name, organization_id)


async def view_specific_day_workforce(query: Update, context: ContextTypes.DEFAULT_TYPE, day: int, month: int) -> None:
    user_data = requests.get(f'{DJANGO_API_URL}users/chat/{query.message.chat.id}/').json()
    user_object_id = user_data.get('object_id')
    year = datetime.now().year
    specific_date = datetime(year, month, day).date()

    response = requests.get(f'{DJANGO_API_URL}frontworkforces/')
    if response.status_code == 200:
        workforces = response.json()
        specific_day_workforces = [wf for wf in workforces if parser.parse(wf['date']).date() == specific_date and wf['object_id'] == user_object_id]
        if specific_day_workforces:
            message = f"\U0001F477 *Численность за {specific_date.strftime('%d.%m.%Y')}:*\n"
            for wf in specific_day_workforces:
                block_section_name = requests.get(f'{DJANGO_API_URL}blocksections/{wf["block_section_id"]}/').json()['name']
                work_type_name = requests.get(f'{DJANGO_API_URL}worktypes/{wf["work_type_id"]}/').json()['name']
                organization_name = requests.get(f'{DJANGO_API_URL}organizations/{wf["organization_id"]}/').json()['organization']
                message += (
                    f"{wf['workforce_count']} чел. — {work_type_name} — {block_section_name} — Этаж {wf['floor']} — {organization_name}\n\n"
                )
            await query.message.reply_text(text=message, parse_mode=ParseMode.MARKDOWN)
        else:
            await query.message.reply_text(f"Численность за {specific_date.strftime('%d.%m.%Y')} не передавалась.")
    else:
        await query.message.reply_text('Ошибка при получении данных. Попробуйте снова.')

    full_name = user_data.get('full_name', 'Пользователь')
    organization_id = user_data.get('organization_id', None)
    await send_main_menu(query.message.chat.id, context, full_name, organization_id)


async def handle_delete_workforce(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    try:
        workforce_id = context.user_data['workforce_id_to_delete']

        # Удаляем данные из базы данных
        delete_response = requests.delete(f'{DJANGO_API_URL}frontworkforces/{workforce_id}/')
        if delete_response.status_code == 200:
            await update.message.reply_text('Запись успешно удалена.')

            # Отправка данных в Google Sheets
            asyncio.create_task(delete_workforce_in_google_sheets(workforce_id))
        else:
            await update.message.reply_text('Ошибка при удалении записи из базы данных.')

        user_id = update.message.from_user.id
        user_response = requests.get(f'{DJANGO_API_URL}users/chat/{user_id}')
        if user_response.status_code == 200:
            user_data = user_response.json()
            full_name = user_data.get('full_name', 'Пользователь')
            organization_id = user_data.get('organization_id', None)
            await send_main_menu(update.message.chat.id, context, full_name, organization_id)

    except Exception as e:
        await update.message.reply_text(f'Ошибка: {str(e)}')

async def delete_workforce_in_google_sheets(workforce_id):
    async with aiohttp.ClientSession() as session:
        data = {
            "action": "delete_workforce",
            "H": workforce_id,  # Используется для поиска записи в Google Sheets
        }

        async with session.post(WEBHOOK_URL, json=data) as response:
            response.raise_for_status()
            return await response.json()


async def handle_transfer_volume(query: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    user_chat_id = query.from_user.id
    await query.message.delete()

    # Получаем данные организации пользователя
    user_response = requests.get(f'{DJANGO_API_URL}users/chat/{user_chat_id}/')
    if user_response.status_code == 200:
        user_data = user_response.json()
        organization_id = user_data['organization_id']
        context.user_data['organization_id'] = organization_id

        # Получаем данные организации, чтобы узнать object_ids
        organization_response = requests.get(f'{DJANGO_API_URL}organizations/{organization_id}/')
        if organization_response.status_code == 200:
            organization_data = organization_response.json()
            organization_object_ids = organization_data.get('object_ids', [])
        else:
            await query.message.reply_text('Ошибка при получении данных организации. Попробуйте снова.')
            return
    else:
        await query.message.reply_text('Ошибка при получении данных пользователя. Попробуйте снова.')
        return

    # Получаем объекты
    response = requests.get(f'{DJANGO_API_URL}objects/')
    if response.status_code == 200:
        objects = response.json()
        filtered_objects = [obj for obj in objects if obj['id'] in organization_object_ids]
        if filtered_objects:
            keyboard = [
                [InlineKeyboardButton(obj['name'], callback_data=f'volume_obj_{obj["id"]}')] for obj in filtered_objects
            ]
            keyboard.append([InlineKeyboardButton("Назад", callback_data='volume_menu')])
            reply_markup = InlineKeyboardMarkup(keyboard)
            await query.message.reply_text('Выберите объект:', reply_markup=reply_markup)
        else:
            await query.message.reply_text('Нет доступных объектов для данной организации.')
    else:
        await query.message.reply_text('Ошибка при получении списка объектов. Попробуйте снова.')


async def send_volume_to_google_sheets(object_name, block_section_name, floor, work_type_name, organization_name, volume_count, id_volume):
    async with aiohttp.ClientSession() as session:
        data = {
            "action": "volume",
            "A": id_volume,
            "B": object_name,
            "C": block_section_name,
            "D": floor,
            "E": work_type_name,
            "F": datetime.now().strftime("%d.%m.%Y"),
            "G": organization_name,
            "H": volume_count
        }

        async with session.post(WEBHOOK_URL, json=data) as response:
            response.raise_for_status()
            return await response.json()

async def handle_volume_count(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    try:
        volume_count = int(update.message.text)
        object_id = context.user_data['volume_object_id']
        block_section_id = context.user_data['volume_block_section_id']
        floor = context.user_data['volume_floor']
        if floor == 'roof':
            floor = 'Кровля'
        work_type_id = context.user_data['volume_work_type_id']
        organization_id = context.user_data['organization_id']
        user_id = update.message.from_user.id

        # Получаем id пользователя в базе
        user_response = requests.get(f'{DJANGO_API_URL}users/chat/{user_id}/')
        if user_response.status_code == 200:
            user_name = user_response.json()['id']
        else:
            await update.message.reply_text('Ошибка при получении id пользователя')
            return

        # Получаем имя блока/секции
        block_section_response = requests.get(f'{DJANGO_API_URL}blocksections/{block_section_id}/')
        if block_section_response.status_code == 200:
            block_section_name = block_section_response.json()['name']
            context.user_data['volume_block_section_name'] = block_section_name
        else:
            await update.message.reply_text('Ошибка при получении данных блока/секции. Попробуйте снова.')
            return

        # Получаем имя объекта
        object_response = requests.get(f'{DJANGO_API_URL}objects/{object_id}/')
        if object_response.status_code == 200:
            object_name = object_response.json()['name']
        else:
            await update.message.reply_text('Ошибка при получении данных объекта. Попробуйте снова.')
            return

        # Получаем имя вида работ
        work_type_response = requests.get(f'{DJANGO_API_URL}worktypes/{work_type_id}/')
        if work_type_response.status_code == 200:
            work_type_name = work_type_response.json()['name']
            context.user_data['volume_work_type_name'] = work_type_name
        else:
            await update.message.reply_text('Ошибка при получении данных вида работ. Попробуйте снова.')
            return

        # Получаем имя организации
        organization_response = requests.get(f'{DJANGO_API_URL}organizations/{organization_id}/')
        if organization_response.status_code == 200:
            organization_name = organization_response.json()['organization']
        else:
            await update.message.reply_text('Ошибка при получении данных организации. Попробуйте снова.')
            return

        volume_data = {
            'object_id': object_id,
            'block_section_id': block_section_id,
            'floor': floor,
            'work_type_id': work_type_id,
            'organization_id': organization_id,
            'volume': volume_count,
            'date': datetime.now().isoformat(),
            'user_id': user_name
        }

        response = requests.post(f'{DJANGO_API_URL}volumes/', json=volume_data)

        if response.status_code == 201:
            id_volume = response.json()['id']
            await update.message.reply_text(
                f"Объем успешно передан на вид работ: \n"
                f"{context.user_data['volume_work_type_name']}\n"
                f"{context.user_data['volume_block_section_name']} Этаж {floor}",
                reply_markup=InlineKeyboardMarkup(
                    [[InlineKeyboardButton(f"Повторить для {context.user_data['volume_work_type_name']}",
                                           callback_data='repeat_volume')],
                    [InlineKeyboardButton("Назад", callback_data='main_menu')]]
                )
            )
            await send_volume_to_google_sheets(
                object_name,
                block_section_name,
                floor,
                work_type_name,
                organization_name,
                volume_count,
                id_volume,
            )
        else:
            await update.message.reply_text(f'Ошибка при передаче объема: {response.text}')
    except ValueError:
        await update.message.reply_text('Пожалуйста, введите корректное число.')

    context.user_data['expecting_volume_count'] = False

async def handle_new_volume_count(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    try:
        new_volume_count = int(update.message.text)
        volume_id = context.user_data['volume_id_to_refactor']

        # Получаем данные о текущей записи
        response = requests.get(f'{DJANGO_API_URL}volumes/{volume_id}/')
        if response.status_code == 200:
            volume_data = response.json()

            # Обновляем данные в базе данных
            volume_data['volume'] = new_volume_count
            update_response = requests.patch(f'{DJANGO_API_URL}volumes/{volume_id}/',
                                             json=volume_data)

            print(update_response.json())
            if update_response.status_code == 200:
                if 'refactor_message_id' in context.user_data:
                    try:
                        await context.bot.delete_message(
                            chat_id=update.message.chat.id,
                            message_id=context.user_data['refactor_message_id']
                        )
                    except:
                        pass
                await update.message.reply_text('Объем успешно обновлен.')

                # Получение данных для вызова send_main_menu
                user_data = requests.get(f'{DJANGO_API_URL}users/chat/{update.message.chat.id}/').json()
                full_name = user_data.get('full_name', 'Пользователь')
                organization_id = user_data.get('organization_id', None)
                await send_main_menu(update.message.chat.id, context, full_name, organization_id)


                # Отправка данных в Google Sheets
                await update_volume_in_google_sheets(
                    volume_data['id'],
                    volume_data['object_id'],
                    volume_data['block_section_id'],
                    volume_data['floor'],
                    volume_data['work_type_id'],
                    volume_data['organization_id'],
                    new_volume_count
                )

            else:
                await update.message.reply_text('Ошибка при обновлении объема в базе данных.')
        else:
            await update.message.reply_text('Ошибка при получении данных объема.')
    except ValueError:
        await update.message.reply_text('Пожалуйста, введите корректное число.')

    context.user_data['expecting_new_volume_count'] = False



async def handle_delete_volume(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    try:
        volume_id = context.user_data['volume_id_to_delete']

        # Удаляем данные из базы данных
        delete_response = requests.delete(f'{DJANGO_API_URL}volumes/{volume_id}/')
        if delete_response.status_code == 200:
            await update.message.reply_text('Запись успешно удалена.')

            # Отправка данных в Google Sheets
            await delete_volume_in_google_sheets(volume_id)
        else:
            await update.message.reply_text('Ошибка при удалении записи из базы данных.')

        user_id = update.message.from_user.id
        user_response = requests.get(f'{DJANGO_API_URL}users/chat/{user_id}')
        if user_response.status_code == 200:
            user_data = user_response.json()
            full_name = user_data.get('full_name', 'Пользователь')
            organization_id = user_data.get('organization_id', None)
            await send_main_menu(update.message.chat.id, context, full_name, organization_id)

    except Exception as e:
        await update.message.reply_text(f'Ошибка: {str(e)}')

async def update_volume_in_google_sheets(volume_id, object_id, block_section_id, floor, work_type_id, organization_id, new_volume_count):
    async with aiohttp.ClientSession() as session:
        # Получение дополнительных данных
        object_data = await get_data_from_api(session, f"/objects/{object_id}")
        block_section_data = await get_data_from_api(session, f"/blocksections/{block_section_id}")
        work_type_data = await get_data_from_api(session, f"/worktypes/{work_type_id}")
        organization_data = await get_data_from_api(session, f"/organizations/{organization_id}")

        data = {
            "action": "update_volume",
            "A": volume_id,
            "B": object_data['name'],
            "C": block_section_data['name'],
            "D": floor,
            "E": work_type_data['name'],
            "F": datetime.now().strftime("%d.%m.%Y"),
            "G": organization_data['organization'],
            "H": new_volume_count
        }

        async with session.post(WEBHOOK_URL, json=data) as response:
            response.raise_for_status()
            return await response.json()


async def delete_volume_in_google_sheets(volume_id):
    async with aiohttp.ClientSession() as session:

        data = {
            "action": "delete_volume",
            "A": volume_id,
        }
        print(data)

        async with session.post(WEBHOOK_URL, json=data) as response:
            response.raise_for_status()
            return await response.json()

async def view_today_volume(query: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    user_data = requests.get(f'{DJANGO_API_URL}users/chat/{query.message.chat.id}/').json()
    user_object_id = user_data.get('object_id')
    today = datetime.now().date()

    response = requests.get(f'{DJANGO_API_URL}volumes/')
    if response.status_code == 200:
        volumes = response.json()
        today_volumes = [vol for vol in volumes if parser.parse(vol['date']).date() == today and vol['object_id'] == user_object_id]
        if today_volumes:
            message = "📐 *Объемы за сегодня:*\n"
            for vol in today_volumes:
                block_section_name = requests.get(f'{DJANGO_API_URL}blocksections/{vol["block_section_id"]}/').json()['name']
                work_type_name = requests.get(f'{DJANGO_API_URL}worktypes/{vol["work_type_id"]}/').json()['name']
                organization_name = requests.get(f'{DJANGO_API_URL}organizations/{vol["organization_id"]}/').json()['organization']
                floor_text = f"— Этаж {vol['floor']}" if vol['floor'] and vol['floor'] != "None" else ""
                message += (
                    f"{vol['volume']} м³ — {work_type_name} — {block_section_name} {floor_text} — {organization_name}\n\n"
                )
            await query.message.reply_text(
                text=message,
                parse_mode=ParseMode.MARKDOWN
            )
        else:
            await query.message.reply_text("Сегодня объемы не передавались.")
    else:
        await query.message.reply_text('Ошибка при получении данных. Попробуйте снова.')

    full_name = user_data.get('full_name', 'Пользователь')
    organization_id = user_data.get('organization_id', None)
    await send_main_menu(query.message.chat.id, context, full_name, organization_id)

async def view_specific_day_volume(query: Update, context: ContextTypes.DEFAULT_TYPE, day: int, month: int) -> None:
    user_data = requests.get(f'{DJANGO_API_URL}users/chat/{query.message.chat.id}/').json()
    user_object_id = user_data.get('object_id')
    year = datetime.now().year
    specific_date = datetime(year, month, day).date()

    response = requests.get(f'{DJANGO_API_URL}volumes/')
    if response.status_code == 200:
        volumes = response.json()
        specific_day_volumes = [vol for vol in volumes if parser.parse(vol['date']).date() == specific_date and vol['object_id'] == user_object_id]
        if specific_day_volumes:
            message = f"📐 *Объемы за {specific_date.strftime('%d.%m.%Y')}:*\n"
            for vol in specific_day_volumes:
                block_section_name = requests.get(f'{DJANGO_API_URL}blocksections/{vol["block_section_id"]}/').json()['name']
                work_type_name = requests.get(f'{DJANGO_API_URL}worktypes/{vol["work_type_id"]}/').json()['name']
                organization_name = requests.get(f'{DJANGO_API_URL}organizations/{vol["organization_id"]}/').json()['organization']
                floor_text = f"— Этаж {vol['floor']}" if vol['floor'] and vol['floor'] != "None" else ""
                message += (
                    f"{vol['volume']} м³ — {work_type_name} — {block_section_name} {floor_text} — {organization_name}\n\n"
                )
            await query.message.reply_text(text=message, parse_mode=ParseMode.MARKDOWN)
        else:
            await query.message.reply_text(f"Объемы за {specific_date.strftime('%d.%m.%Y')} не передавались.")
    else:
        await query.message.reply_text('Ошибка при получении данных. Попробуйте снова.')

    full_name = user_data.get('full_name', 'Пользователь')
    organization_id = user_data.get('organization_id', None)
    await send_main_menu(query.message.chat.id, context, full_name, organization_id)

async def send_prefab_types(chat_id, context: ContextTypes.DEFAULT_TYPE):
    response = requests.get(f'{DJANGO_API_URL}prefab_types/')
    if response.status_code == 200:
        prefab_types = response.json()
        keyboard = [
            [InlineKeyboardButton(prefab_type['name'], callback_data=f"prefab_type_{prefab_type['id']}")]
            for prefab_type in prefab_types
        ]
        keyboard.append([InlineKeyboardButton("Назад", callback_data='main_menu')])
        reply_markup = InlineKeyboardMarkup(keyboard)
        await context.bot.send_message(
            chat_id=chat_id,
            text="Выберите тип префаба:",
            reply_markup=reply_markup
        )

async def send_prefab_subtypes(chat_id, context: ContextTypes.DEFAULT_TYPE, prefab_type_id: int):
    response = requests.get(f'{DJANGO_API_URL}prefab_subtypes/?prefab_type_id={prefab_type_id}')
    if response.status_code == 200:
        prefab_subtypes = response.json()
        keyboard = [
            [InlineKeyboardButton(prefab_subtype['name'], callback_data=f"prefab_subtype_{prefab_subtype['id']}")]
            for prefab_subtype in prefab_subtypes
        ]
        keyboard.append([InlineKeyboardButton("Назад", callback_data='fact_production')])
        reply_markup = InlineKeyboardMarkup(keyboard)
        await context.bot.send_message(
            chat_id=chat_id,
            text="Выберите подтип префаба:",
            reply_markup=reply_markup
        )


async def send_prefabs(chat_id, context: ContextTypes.DEFAULT_TYPE, selected_prefab_subtype_id: int):
    # Получаем данные пользователя
    user_response = requests.get(f'{DJANGO_API_URL}users/chat/{chat_id}/')
    if user_response.status_code != 200:
        await context.bot.send_message(
            chat_id=chat_id,
            text="Ошибка при получении данных пользователя. Попробуйте снова."
        )
        return

    user_data = user_response.json()
    organization_id = user_data.get('organization_id')

    if not organization_id:
        await context.bot.send_message(
            chat_id=chat_id,
            text="Ваша организация не определена. Пожалуйста, свяжитесь с администратором."
        )
        return

    # Получаем все префабы
    response = requests.get(f'{DJANGO_API_URL}prefabs/')
    if response.status_code == 200:
        prefabs = response.json()
        # Фильтруем префабы по организации и выбранному prefab_subtype_id
        prefabs = [prefab for prefab in prefabs if prefab['organization_id'] == organization_id and prefab['prefab_subtype_id'] == selected_prefab_subtype_id]

        if prefabs:
            keyboard = []
            for prefab in prefabs:
                prefab_type_id = prefab['prefab_type_id']
                prefab_subtype_id = prefab['prefab_subtype_id']
                quantity = prefab.get('quantity', 0)

                # Получаем текущие записи в prefabs_in_work для данного prefab_id
                prefabs_in_work_response = requests.get(f'{DJANGO_API_URL}prefabs_in_work/')
                if prefabs_in_work_response.status_code == 200:
                    prefabs_in_work = prefabs_in_work_response.json()
                    total_quantity_in_work = sum([p['quantity'] for p in prefabs_in_work if p['prefab_id'] == prefab['id']])
                else:
                    total_quantity_in_work = 0

                # Вычисляем доступное количество для добавления
                remaining_quantity = quantity - total_quantity_in_work

                if remaining_quantity <= 0:
                    continue  # Пропускаем префабы с нулевым доступным количеством

                # Получаем имя prefab_type
                type_response = requests.get(f'{DJANGO_API_URL}prefab_types/{prefab_type_id}')
                if type_response.status_code == 200:
                    prefab_type_name = type_response.json().get('name', 'Неизвестный тип')
                else:
                    prefab_type_name = 'Неизвестный тип'

                # Получаем имя prefab_subtype
                subtype_response = requests.get(f'{DJANGO_API_URL}prefab_subtypes/{prefab_subtype_id}')
                if subtype_response.status_code == 200:
                    prefab_subtype_name = subtype_response.json().get('name', 'Неизвестный подтип')
                else:
                    prefab_subtype_name = 'Неизвестный подтип'

                # Получаем имя объекта
                object_id = prefab.get('object_id')
                if object_id:
                    object_response = requests.get(f'{DJANGO_API_URL}objects/{object_id}')
                    if object_response.status_code == 200:
                        object_name = object_response.json().get('name', 'Неизвестный объект')
                    else:
                        object_name = 'Неизвестный объект'
                else:
                    object_name = 'Объект не указан'

                button_text = f"{remaining_quantity} — {prefab_type_name} — {prefab_subtype_name} — {object_name}"
                keyboard.append([InlineKeyboardButton(button_text, callback_data=f"prefab_{prefab['id']}")])

            if keyboard:
                keyboard.append([InlineKeyboardButton("Назад", callback_data='main_menu')])
                reply_markup = InlineKeyboardMarkup(keyboard)
                await context.bot.send_message(
                    chat_id=chat_id,
                    text="Выберите префаб для начала производства:",
                    reply_markup=reply_markup
                )
            else:
                await context.bot.send_message(
                    chat_id=chat_id,
                    text="Нет доступных префабов выбранного типа или они уже в работе."
                )
                await send_main_menu(chat_id, context, user_data.get('full_name', 'Пользователь'),
                                     user_data.get('organization_id', None))

        else:
            await context.bot.send_message(
                chat_id=chat_id,
                text="Нет доступных префабов выбранного типа или они уже в работе."
            )
            await send_main_menu(chat_id, context, user_data.get('full_name', 'Пользователь'), user_data.get('organization_id', None))

    else:
        await context.bot.send_message(
            chat_id=chat_id,
            text="Ошибка при получении префабов. Попробуйте снова."
        )




async def handle_prefab_quantity(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if context.user_data.get('expecting_prefab_quantity'):
        try:
            quantity = int(update.message.text)
            prefab_id = context.user_data['selected_prefab_id']

            # Получаем данные о выбранном префабе
            prefab_response = requests.get(f'{DJANGO_API_URL}prefabs/{prefab_id}')
            if prefab_response.status_code != 200:
                await update.message.reply_text('Ошибка при получении данных префаба. Попробуйте снова.')
                context.user_data['expecting_prefab_quantity'] = True
                return

            prefab_data = prefab_response.json()
            available_quantity = prefab_data.get('quantity', 0)

            # Получаем текущие записи в prefabs_in_work для данного prefab_id
            prefabs_in_work_response = requests.get(f'{DJANGO_API_URL}prefabs_in_work/')
            if prefabs_in_work_response.status_code == 200:
                prefabs_in_work = prefabs_in_work_response.json()
                total_quantity_in_work = sum([p['quantity'] for p in prefabs_in_work if p['prefab_id'] == prefab_id])
            else:
                total_quantity_in_work = 0

            # Вычисляем доступное количество для добавления
            remaining_quantity = available_quantity - total_quantity_in_work

            # Проверяем доступное количество
            if remaining_quantity <= 0:
                await update.message.reply_text(
                    "Все префабы данного вида уже в производстве."
                )
                context.user_data['expecting_prefab_quantity'] = False
                return

            if quantity > remaining_quantity:
                await update.message.reply_text(
                    f"Введенное количество превышает доступное ({remaining_quantity}). Попробуйте снова."
                )
                context.user_data['expecting_prefab_quantity'] = True
                return

            # Получаем текущую дату и время
            production_date = datetime.utcnow().isoformat()

            prefabs_in_work_data = {
                'prefab_id': prefab_id,
                'quantity': quantity,
                'status': 'production',
                'production_date': production_date  # Добавляем текущую дату в production_date
            }
            response = requests.post(f'{DJANGO_API_URL}prefabs_in_work/', json=prefabs_in_work_data)
            if response.status_code == 201:
                # Получаем имя prefab_type и prefab_subtype для сообщения
                prefab_type_id = prefab_data['prefab_type_id']
                prefab_subtype_id = prefab_data['prefab_subtype_id']

                type_response = requests.get(f'{DJANGO_API_URL}prefab_types/{prefab_type_id}')
                if type_response.status_code == 200:
                    prefab_type_name = type_response.json().get('name', 'Неизвестный тип')
                else:
                    prefab_type_name = 'Неизвестный тип'

                subtype_response = requests.get(f'{DJANGO_API_URL}prefab_subtypes/{prefab_subtype_id}')
                if subtype_response.status_code == 200:
                    prefab_subtype_name = subtype_response.json().get('name', 'Неизвестный подтип')
                else:
                    prefab_subtype_name = 'Неизвестный подтип'

                await update.message.reply_text(
                    f"\U00002705 {prefab_type_name} — {prefab_subtype_name} успешно отправлены в производство.\n"
                    f"Чтобы присвоить им статус СГП, перейдите в соответствующую вкладку."
                )
                context.user_data['expecting_prefab_quantity'] = False

                # Вызываем send_main_menu
                user_data = requests.get(f'{DJANGO_API_URL}users/chat/{update.message.chat.id}/').json()
                full_name = user_data.get('full_name', 'Пользователь')
                organization_id = user_data.get('organization_id', None)
                await send_main_menu(update.message.chat.id, context, full_name, organization_id)
            else:
                await update.message.reply_text('Ошибка при создании записи. Попробуйте снова.')
                context.user_data['expecting_prefab_quantity'] = True

        except ValueError:
            await update.message.reply_text('Пожалуйста, введите корректное число.')
            context.user_data['expecting_prefab_quantity'] = True



async def handle_sgp_confirmation(update: Update, context: ContextTypes.DEFAULT_TYPE):
    query = update.callback_query
    await query.answer()

    data = query.data
    if data.startswith('sgp_prefab_'):
        prefab_in_work_id = int(data.split('_')[2])

        # Получаем данные о префабе в работе
        prefab_in_work_response = requests.get(f'{DJANGO_API_URL}prefabs_in_work/{prefab_in_work_id}')
        if prefab_in_work_response.status_code != 200:
            await query.message.reply_text('Ошибка при получении данных префаба. Попробуйте снова.')
            return

        prefab_in_work_data = prefab_in_work_response.json()
        prefab_id = prefab_in_work_data['prefab_id']
        quantity = prefab_in_work_data.get('quantity', 0)

        # Сохраняем данные в контексте
        context.user_data['prefab_in_work_id_to_update'] = prefab_in_work_id
        context.user_data['quantity_to_update'] = quantity

        # Получаем данные о префабе
        prefab_response = requests.get(f'{DJANGO_API_URL}prefabs/{prefab_id}')
        if prefab_response.status_code == 200:
            prefab_data = prefab_response.json()
            prefab_type_id = prefab_data['prefab_type_id']
            prefab_subtype_id = prefab_data['prefab_subtype_id']
            object_id = prefab_data.get('object_id')

            # Получаем имя prefab_type
            type_response = requests.get(f'{DJANGO_API_URL}prefab_types/{prefab_type_id}')
            if type_response.status_code == 200:
                prefab_type_name = type_response.json().get('name', 'Неизвестный тип')
            else:
                prefab_type_name = 'Неизвестный тип'

            # Получаем имя prefab_subtype
            subtype_response = requests.get(f'{DJANGO_API_URL}prefab_subtypes/{prefab_subtype_id}')
            if subtype_response.status_code == 200:
                prefab_subtype_name = subtype_response.json().get('name', 'Неизвестный подтип')
            else:
                prefab_subtype_name = 'Неизвестный подтип'

            # Получаем имя объекта
            object_response = requests.get(f'{DJANGO_API_URL}objects/{object_id}')
            if object_response.status_code == 200:
                object_name = object_response.json().get('name', 'Неизвестный объект')
            else:
                object_name = 'Неизвестный объект'

            # Спрашиваем пользователя о подтверждении изменения статуса
            keyboard = [
                [
                    InlineKeyboardButton("\U00002705 Да", callback_data='sgp_confirm_yes'),
                    InlineKeyboardButton("\U0000274C Нет", callback_data='sgp_confirm_no')
                ]
            ]
            reply_markup = InlineKeyboardMarkup(keyboard)
            await query.message.reply_text(
                f"Изменить статус у *{quantity}шт. — {prefab_type_name} — {prefab_subtype_name} — {object_name}* на СГП?",
                reply_markup=reply_markup,
                parse_mode=ParseMode.MARKDOWN,

            )
        else:
            await query.message.reply_text('Ошибка при получении данных префаба. Попробуйте снова.')

    elif data == 'sgp_confirm_yes':
        prefab_in_work_id = context.user_data.get('prefab_in_work_id_to_update')
        if prefab_in_work_id:
            update_data = {'status': 'sgp'}
            print(f"Updating PrefabsInWork ID: {prefab_in_work_id} with data: {update_data}")  # Логирование данных
            update_response = requests.put(
                f'{DJANGO_API_URL}prefabs_in_work/{prefab_in_work_id}/',
                json=update_data
            )
            if update_response.status_code == 200:
                # Получаем данные для сообщения
                prefab_in_work_response = requests.get(f'{DJANGO_API_URL}prefabs_in_work/{prefab_in_work_id}')
                if prefab_in_work_response.status_code == 200:
                    prefab_in_work_data = prefab_in_work_response.json()
                    prefab_id = prefab_in_work_data['prefab_id']
                    quantity = prefab_in_work_data.get('quantity', 0)

                    prefab_response = requests.get(f'{DJANGO_API_URL}prefabs/{prefab_id}')
                    if prefab_response.status_code == 200:
                        prefab_data = prefab_response.json()
                        prefab_type_id = prefab_data['prefab_type_id']
                        prefab_subtype_id = prefab_data['prefab_subtype_id']
                        object_id = prefab_data.get('object_id')

                        type_response = requests.get(f'{DJANGO_API_URL}prefab_types/{prefab_type_id}')
                        if type_response.status_code == 200:
                            prefab_type_name = type_response.json().get('name', 'Неизвестный тип')
                        else:
                            prefab_type_name = 'Неизвестный тип'

                        subtype_response = requests.get(f'{DJANGO_API_URL}prefab_subtypes/{prefab_subtype_id}')
                        if subtype_response.status_code == 200:
                            prefab_subtype_name = subtype_response.json().get('name', 'Неизвестный подтип')
                        else:
                            prefab_subtype_name = 'Неизвестный подтип'

                        object_response = requests.get(f'{DJANGO_API_URL}objects/{object_id}')
                        if object_response.status_code == 200:
                            object_name = object_response.json().get('name', 'Неизвестный объект')
                        else:
                            object_name = 'Неизвестный объект'

                        await query.message.reply_text(
                            f"\U00002705 {prefab_type_name} — {prefab_subtype_name} — {object_name} успешно изменены на статус СГП.\n"
                            f"Чтобы отправить их в производство, перейдите в соответствующую вкладку."
                        )

                        # Получаем данные пользователя для вызова send_main_menu
                        user_data = requests.get(f'{DJANGO_API_URL}users/chat/{query.message.chat.id}/').json()
                        full_name = user_data.get('full_name', 'Пользователь')
                        organization_id = user_data.get('organization_id', None)

                        # Вызываем send_main_menu
                        await send_main_menu(query.message.chat.id, context, full_name, organization_id)
                    else:
                        await query.message.reply_text('Ошибка при получении данных префаба. Попробуйте снова.')
                else:
                    await query.message.reply_text('Ошибка при получении данных префаба. Попробуйте снова.')
            else:
                await query.message.reply_text('Ошибка при изменении статуса. Попробуйте снова.')
        else:
            await query.message.reply_text('Ошибка: не найдены данные для обновления.')

    elif data == 'sgp_confirm_no':
        await query.message.reply_text('Изменение статуса отменено.')

        # Получаем данные пользователя для вызова send_main_menu
        user_data = requests.get(f'{DJANGO_API_URL}users/chat/{query.message.chat.id}/').json()
        full_name = user_data.get('full_name', 'Пользователь')
        organization_id = user_data.get('organization_id', None)

        # Вызываем send_main_menu
        await send_main_menu(query.message.chat.id, context, full_name, organization_id)


async def send_prefabs_for_shipment(chat_id, context: ContextTypes.DEFAULT_TYPE):
    user_response = requests.get(f'{DJANGO_API_URL}users/chat/{chat_id}/')
    user_data = user_response.json()
    if user_response.status_code != 200:
        await context.bot.send_message(
            chat_id=chat_id,
            text="Ошибка при получении данных пользователя. Попробуйте снова."
        )
        await send_main_menu(chat_id, context, user_data.get('full_name', 'Пользователь'), user_data.get('organization_id', None))
        return

    organization_id = user_data.get('organization_id')

    if not organization_id:
        await context.bot.send_message(
            chat_id=chat_id,
            text="Ваша организация не определена. Пожалуйста, свяжитесь с администратором."
        )
        await send_main_menu(chat_id, context, user_data.get('full_name', 'Пользователь'), user_data.get('organization_id', None))
        return

    # Получаем все префабы в работе
    response = requests.get(f'{DJANGO_API_URL}prefabs_in_work/')
    if response.status_code == 200:
        prefabs_in_work = response.json()
        prefabs_in_work = [p for p in prefabs_in_work if p['status'] == 'sgp']

        if prefabs_in_work:
            # Группируем префабы по prefab_id и суммируем их количество
            grouped_prefabs = {}
            for prefab in prefabs_in_work:
                prefab_id = prefab['prefab_id']
                quantity = prefab.get('quantity', 0)
                if prefab_id in grouped_prefabs:
                    grouped_prefabs[prefab_id]['quantity'] += quantity
                else:
                    grouped_prefabs[prefab_id] = {'quantity': quantity, 'id': prefab['id']}

            keyboard = []
            for prefab_id, data in grouped_prefabs.items():
                # Получаем данные о префабе
                prefab_response = requests.get(f'{DJANGO_API_URL}prefabs/{prefab_id}')
                if prefab_response.status_code == 200:
                    prefab_data = prefab_response.json()
                    if prefab_data['organization_id'] != organization_id:
                        continue

                    prefab_type_id = prefab_data['prefab_type_id']
                    prefab_subtype_id = prefab_data['prefab_subtype_id']
                    object_id = prefab_data.get('object_id')

                    # Получаем имя prefab_type
                    type_response = requests.get(f'{DJANGO_API_URL}prefab_types/{prefab_type_id}')
                    if type_response.status_code == 200:
                        prefab_type_name = type_response.json().get('name', 'Неизвестный тип')
                    else:
                        prefab_type_name = 'Неизвестный тип'

                    # Получаем имя prefab_subtype
                    subtype_response = requests.get(f'{DJANGO_API_URL}prefab_subtypes/{prefab_subtype_id}')
                    if subtype_response.status_code == 200:
                        prefab_subtype_name = subtype_response.json().get('name', 'Неизвестный подтип')
                    else:
                        prefab_subtype_name = 'Неизвестный подтип'

                    # Получаем имя объекта
                    object_response = requests.get(f'{DJANGO_API_URL}objects/{object_id}')
                    if object_response.status_code == 200:
                        object_name = object_response.json().get('name', 'Неизвестный объект')
                    else:
                        object_name = 'Неизвестный объект'

                    button_text = f"{data['quantity']}шт. — {prefab_type_name} — {prefab_subtype_name} — {object_name}"
                    keyboard.append([InlineKeyboardButton(button_text, callback_data=f"shipment_prefab_{data['id']}")])

            if keyboard:
                keyboard.append([InlineKeyboardButton("Назад", callback_data='main_menu')])
                reply_markup = InlineKeyboardMarkup(keyboard)
                await context.bot.send_message(
                    chat_id=chat_id,
                    text="Выберите префаб для изменения статуса на Отгрузку:",
                    reply_markup=reply_markup
                )
            else:
                await context.bot.send_message(
                    chat_id=chat_id,
                    text="Нет префабов в статусе СГП."
                )
                await send_main_menu(chat_id, context, user_data.get('full_name', 'Пользователь'), user_data.get('organization_id', None))
        else:
            await context.bot.send_message(
                chat_id=chat_id,
                text="Нет префабов в статусе СГП."
            )
            await send_main_menu(chat_id, context, user_data.get('full_name', 'Пользователь'), user_data.get('organization_id', None))
    else:
        await context.bot.send_message(
            chat_id=chat_id,
            text="Ошибка при получении префабов. Попробуйте снова."
        )
        await send_main_menu(chat_id, context, user_data.get('full_name', 'Пользователь'), user_data.get('organization_id', None))


async def handle_shipment_quantity_input(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if context.user_data.get('expecting_shipment_quantity'):
        try:
            quantity = int(update.message.text)
            selected_prefab_in_work_id = context.user_data['selected_prefab_in_work_id']

            # Получаем данные о выбранной записи prefabs_in_work
            prefab_in_work_response = requests.get(f'{DJANGO_API_URL}prefabs_in_work/{selected_prefab_in_work_id}')
            if prefab_in_work_response.status_code != 200:
                await update.message.reply_text('Ошибка при получении данных префаба. Попробуйте снова.')
                context.user_data['expecting_shipment_quantity'] = True
                return

            prefab_in_work_data = prefab_in_work_response.json()
            selected_prefab_id = prefab_in_work_data['prefab_id']

            # Получаем все записи префабов с одинаковым prefab_id и статусом 'sgp'
            prefabs_in_work_response = requests.get(f'{DJANGO_API_URL}prefabs_in_work/')
            if prefabs_in_work_response.status_code != 200:
                await update.message.reply_text('Ошибка при получении данных префабов. Попробуйте снова.')
                context.user_data['expecting_shipment_quantity'] = True
                return

            all_prefabs_in_work = prefabs_in_work_response.json()
            prefabs_to_update = [p for p in all_prefabs_in_work if p['prefab_id'] == selected_prefab_id and p['status'] == 'sgp']

            # Вычисляем общее количество доступных префабов
            total_available_quantity = sum(p['quantity'] for p in prefabs_to_update)

            # Проверяем доступное количество
            if quantity > total_available_quantity:
                await update.message.reply_text(
                    f"Введенное количество ({quantity}) превышает доступное ({total_available_quantity}). Попробуйте снова."
                )
                context.user_data['expecting_shipment_quantity'] = True
                return

            # Обновляем префабы, разделяем их если необходимо
            update_quantity_remaining = quantity

            for prefab in prefabs_to_update:
                if update_quantity_remaining <= 0:
                    break

                current_quantity = prefab['quantity']
                if current_quantity <= update_quantity_remaining:
                    # Обновляем статус на shipment для текущего префаба
                    update_response = requests.put(
                        f'{DJANGO_API_URL}prefabs_in_work/{prefab["id"]}/',
                        json={'status': 'shipment', 'shipment_date': datetime.utcnow().isoformat()}
                    )
                    if update_response.status_code == 200:
                        update_quantity_remaining -= current_quantity
                    else:
                        await update.message.reply_text('Ошибка при обновлении статуса префаба. Попробуйте снова.')
                        context.user_data['expecting_shipment_quantity'] = True
                        return
                else:
                    # Разделяем префаб на два
                    new_quantity = current_quantity - update_quantity_remaining
                    update_response = requests.patch(
                        f'{DJANGO_API_URL}prefabs_in_work/{prefab["id"]}/',
                        json={'quantity': new_quantity}
                    )
                    if update_response.status_code == 200:
                        # Создаем новую запись с количеством и статусом shipment
                        new_prefabs_in_work_data = {
                            'prefab_id': prefab['prefab_id'],
                            'quantity': update_quantity_remaining,
                            'status': 'shipment',
                            'production_date': prefab['production_date'],
                            'shipment_date': datetime.utcnow().isoformat()
                        }
                        new_prefab_response = requests.post(f'{DJANGO_API_URL}prefabs_in_work/', json=new_prefabs_in_work_data)
                        if new_prefab_response.status_code == 201:
                            update_quantity_remaining = 0
                        else:
                            await update.message.reply_text('Ошибка при создании записи. Попробуйте снова.')
                            context.user_data['expecting_shipment_quantity'] = True
                            return
                    else:
                        await update.message.reply_text('Ошибка при обновлении количества префаба. Попробуйте снова.')
                        context.user_data['expecting_shipment_quantity'] = True
                        return

            await update.message.reply_text(
                "\U00002705 Префабы успешно переданы на отгрузку."
            )
            context.user_data['expecting_shipment_quantity'] = False

            # Вызываем send_main_menu
            user_data = requests.get(f'{DJANGO_API_URL}users/chat/{update.message.chat.id}/').json()
            full_name = user_data.get('full_name', 'Пользователь')
            organization_id = user_data.get('organization_id', None)
            await send_main_menu(update.message.chat.id, context, full_name, organization_id)

        except ValueError:
            await update.message.reply_text('Пожалуйста, введите корректное число.')
            context.user_data['expecting_shipment_quantity'] = True

async def handle_shipment_confirmation(update: Update, context: ContextTypes.DEFAULT_TYPE):
    query = update.callback_query
    await query.answer()

    data = query.data
    if data.startswith('shipment_prefab_'):
        prefabs_in_work_id = int(data.split('_')[2])
        context.user_data['selected_prefab_in_work_id'] = prefabs_in_work_id
        context.user_data['expecting_shipment_quantity'] = True
        await query.message.reply_text(
            "Введите количество префабов для отправки на отгрузку:"
        )

    elif data == 'shipment_confirm_yes':
        prefab_in_work_id = context.user_data.get('prefab_in_work_id_to_update')
        if prefab_in_work_id:
            update_data = {'status': 'shipment'}
            update_response = requests.put(
                f'{DJANGO_API_URL}prefabs_in_work/{prefab_in_work_id}/',
                json=update_data
            )
            if update_response.status_code == 200:
                # Получаем данные для сообщения
                prefab_in_work_response = requests.get(f'{DJANGO_API_URL}prefabs_in_work/{prefab_in_work_id}')
                if prefab_in_work_response.status_code == 200:
                    prefab_in_work_data = prefab_in_work_response.json()
                    prefab_id = prefab_in_work_data['prefab_id']
                    quantity = prefab_in_work_data.get('quantity', 0)

                    prefab_response = requests.get(f'{DJANGO_API_URL}prefabs/{prefab_id}')
                    if prefab_response.status_code == 200:
                        prefab_data = prefab_response.json()
                        prefab_type_id = prefab_data['prefab_type_id']
                        prefab_subtype_id = prefab_data['prefab_subtype_id']
                        object_id = prefab_data.get('object_id')

                        type_response = requests.get(f'{DJANGO_API_URL}prefab_types/{prefab_type_id}')
                        if type_response.status_code == 200:
                            prefab_type_name = type_response.json().get('name', 'Неизвестный тип')
                        else:
                            prefab_type_name = 'Неизвестный тип'

                        subtype_response = requests.get(f'{DJANGO_API_URL}prefab_subtypes/{prefab_subtype_id}')
                        if subtype_response.status_code == 200:
                            prefab_subtype_name = subtype_response.json().get('name', 'Неизвестный подтип')
                        else:
                            prefab_subtype_name = 'Неизвестный подтип'

                        object_response = requests.get(f'{DJANGO_API_URL}objects/{object_id}')
                        if object_response.status_code == 200:
                            object_name = object_response.json().get('name', 'Неизвестный объект')
                        else:
                            object_name = 'Неизвестный объект'

                        await query.message.reply_text(
                            f"Префаб \U00002705 {prefab_type_name} — {prefab_subtype_name} — {object_name} успешно изменен на статус Отгрузка.\n\n"
                            f"Если инженер по качеству оставит замечания из-за бракованного префаба, Вам придет уведомление, а префабы под замену окажутся во вкладке *Факт на производство*.",
                            parse_mode=ParseMode.MARKDOWN
                        )

                        # Получаем данные пользователя для вызова send_main_menu
                        user_data = requests.get(f'{DJANGO_API_URL}users/chat/{query.message.chat.id}/').json()
                        full_name = user_data.get('full_name', 'Пользователь')
                        organization_id = user_data.get('organization_id', None)

                        # Вызываем send_main_menu
                        await send_main_menu(query.message.chat.id, context, full_name, organization_id)
                    else:
                        await query.message.reply_text('Ошибка при получении данных префаба. Попробуйте снова.')
                else:
                    await query.message.reply_text('Ошибка при получении данных префаба. Попробуйте снова.')
            else:
                await query.message.reply_text('Ошибка при изменении статуса. Попробуйте снова.')
        else:
            await query.message.reply_text('Ошибка: не найдены данные для обновления.')

    elif data == 'shipment_confirm_no':
        await query.message.reply_text('Изменение статуса отменено.')

        # Получаем данные пользователя для вызова send_main_menu
        user_data = requests.get(f'{DJANGO_API_URL}users/chat/{query.message.chat.id}/').json()
        full_name = user_data.get('full_name', 'Пользователь')
        organization_id = user_data.get('organization_id', None)

        # Вызываем send_main_menu
        await send_main_menu(query.message.chat.id, context, full_name, organization_id)

async def send_prefabs_list(chat_id, context: ContextTypes.DEFAULT_TYPE, status: str):
    user_response = requests.get(f'{DJANGO_API_URL}users/chat/{chat_id}/')
    user_data = user_response.json()
    if user_response.status_code != 200:
        await context.bot.send_message(
            chat_id=chat_id,
            text="Ошибка при получении данных пользователя. Попробуйте снова."
        )
        await send_main_menu(chat_id, context, user_data.get('full_name', 'Пользователь'), user_data.get('organization_id', None))
        return

    organization_id = user_data.get('organization_id')
    user_object_id = user_data.get('object_id')

    if not organization_id or not user_object_id:
        await context.bot.send_message(
            chat_id=chat_id,
            text="Ваша организация или объект не определены. Пожалуйста, свяжитесь с администратором."
        )
        await send_main_menu(chat_id, context, user_data.get('full_name', 'Пользователь'), user_data.get('organization_id', None))
        return

    response = requests.get(f'{DJANGO_API_URL}prefabs_in_work/')
    if response.status_code == 200:
        prefabs_in_work = response.json()
        prefabs_in_work = [p for p in prefabs_in_work if p['status'] == status]

        if prefabs_in_work:
            aggregated_prefabs = {}
            for prefab in prefabs_in_work:
                prefab_id = prefab['prefab_id']
                quantity = prefab.get('quantity', 0)

                prefab_response = requests.get(f'{DJANGO_API_URL}prefabs/{prefab_id}')
                if prefab_response.status_code == 200:
                    prefab_data = prefab_response.json()

                    if prefab_data['object_id'] != user_object_id:
                        continue  # Пропускаем префабы, не относящиеся к объекту пользователя

                    prefab_type_id = prefab_data['prefab_type_id']
                    prefab_subtype_id = prefab_data['prefab_subtype_id']
                    prefab_org_id = prefab_data.get('organization_id')

                    type_response = requests.get(f'{DJANGO_API_URL}prefab_types/{prefab_type_id}')
                    if type_response.status_code == 200:
                        prefab_type_name = type_response.json().get('name', 'Неизвестный тип')
                    else:
                        prefab_type_name = 'Неизвестный тип'

                    subtype_response = requests.get(f'{DJANGO_API_URL}prefab_subtypes/{prefab_subtype_id}')
                    if subtype_response.status_code == 200:
                        prefab_subtype_name = subtype_response.json().get('name', 'Неизвестный подтип')
                    else:
                        prefab_subtype_name = 'Неизвестный подтип'

                    org_response = requests.get(f'{DJANGO_API_URL}organizations/{prefab_org_id}')
                    if org_response.status_code == 200:
                        org_name = org_response.json().get('organization', 'Неизвестная организация')
                    else:
                        org_name = 'Неизвестная организация'

                    key = f"{org_name} — {prefab_type_name} — {prefab_subtype_name}"
                    if key not in aggregated_prefabs:
                        aggregated_prefabs[key] = 0
                    aggregated_prefabs[key] += quantity

            if aggregated_prefabs:
                messages = [f"{key} — {quantity}шт." for key, quantity in aggregated_prefabs.items()]
                await context.bot.send_message(
                    chat_id=chat_id,
                    text="\n".join(messages)
                )
            else:
                await context.bot.send_message(
                    chat_id=chat_id,
                    text="Нет префабов с указанным статусом."
                )
        else:
            await context.bot.send_message(
                chat_id=chat_id,
                text="Нет префабов с указанным статусом."
            )
    else:
        await context.bot.send_message(
            chat_id=chat_id,
            text="Ошибка при получении префабов. Попробуйте снова."
        )
    await send_main_menu(chat_id, context, user_data.get('full_name', 'Пользователь'), user_data.get('organization_id', None))



async def send_warehouses_list(chat_id, context: ContextTypes.DEFAULT_TYPE):
    response = requests.get(f'{DJANGO_API_URL}warehouses/')
    if response.status_code == 200:
        warehouses = response.json()
        if warehouses:
            keyboard = [[InlineKeyboardButton(warehouse['name'], callback_data=f'select_warehouse_{warehouse["id"]}')]
                        for warehouse in warehouses]
            keyboard.append([InlineKeyboardButton("Назад", callback_data='prefabsoptionlist')])
            reply_markup = InlineKeyboardMarkup(keyboard)
            await context.bot.send_message(
                chat_id=chat_id,
                text="Выберите склад:",
                reply_markup=reply_markup
            )
        else:
            await context.bot.send_message(
                chat_id=chat_id,
                text="Склады не найдены."
            )

    else:
        await context.bot.send_message(
            chat_id=chat_id,
            text="Ошибка при получении списка складов. Попробуйте снова."
        )

async def handle_prefab_sgp_quantity(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if context.user_data.get('expecting_sgp_quantity'):
        try:
            quantity = int(update.message.text)
            selected_prefab_id = context.user_data['selected_prefab_in_work_id']

            # Получаем все записи префабов с одинаковым prefab_id
            prefabs_in_work_response = requests.get(f'{DJANGO_API_URL}prefabs_in_work/')
            if prefabs_in_work_response.status_code != 200:
                await update.message.reply_text('Ошибка при получении данных префабов. Попробуйте снова.')
                context.user_data['expecting_sgp_quantity'] = True
                return

            selected_prefab_id = requests.get(f'{DJANGO_API_URL}prefabs_in_work/{selected_prefab_id}').json()['prefab_id']
            all_prefabs_in_work = prefabs_in_work_response.json()
            prefabs_to_update = [p for p in all_prefabs_in_work if p['prefab_id'] == selected_prefab_id and p['status'] == 'production']


            # Вычисляем общее количество доступных префабов
            total_available_quantity = sum(p['quantity'] for p in prefabs_to_update)

            # Проверяем доступное количество
            if quantity > total_available_quantity:
                await update.message.reply_text(
                    f"Введенное количество ({quantity}) превышает доступное ({total_available_quantity}). Попробуйте снова."
                )
                context.user_data['expecting_sgp_quantity'] = True
                return

            # Обновляем префабы, разделяем их если необходимо
            update_quantity_remaining = quantity

            for prefab in prefabs_to_update:
                if update_quantity_remaining <= 0:
                    break

                current_quantity = prefab['quantity']
                if current_quantity <= update_quantity_remaining:
                    # Обновляем статус на sgp для текущего префаба
                    update_response = requests.put(
                        f'{DJANGO_API_URL}prefabs_in_work/{prefab["id"]}/',
                        json={'status': 'sgp', 'sgp_date': datetime.utcnow().isoformat()}
                    )
                    if update_response.status_code == 200:
                        update_quantity_remaining -= current_quantity
                    else:
                        await update.message.reply_text('Ошибка при обновлении статуса префаба. Попробуйте снова.')
                        context.user_data['expecting_sgp_quantity'] = True
                        return
                else:
                    # Разделяем префаб на два
                    new_quantity = current_quantity - update_quantity_remaining
                    update_response = requests.patch(
                        f'{DJANGO_API_URL}prefabs_in_work/{prefab["id"]}/',
                        json={'quantity': new_quantity}
                    )
                    if update_response.status_code == 200:
                        # Создаем новую запись с количеством и статусом sgp
                        new_prefabs_in_work_data = {
                            'prefab_id': prefab['prefab_id'],
                            'quantity': update_quantity_remaining,
                            'status': 'sgp',
                            'production_date': prefab['production_date'],
                            'sgp_date': datetime.utcnow().isoformat()
                        }
                        new_prefab_response = requests.post(f'{DJANGO_API_URL}prefabs_in_work/', json=new_prefabs_in_work_data)
                        if new_prefab_response.status_code == 201:
                            update_quantity_remaining = 0
                        else:
                            await update.message.reply_text('Ошибка при создании записи. Попробуйте снова.')
                            context.user_data['expecting_sgp_quantity'] = True
                            return
                    else:
                        await update.message.reply_text('Ошибка при обновлении количества префаба. Попробуйте снова.')
                        context.user_data['expecting_sgp_quantity'] = True
                        return

            # Получаем данные о префабе
            prefab_data = requests.get(f'{DJANGO_API_URL}prefabs/{selected_prefab_id}').json()
            prefab_type_id = prefab_data['prefab_type_id']
            prefab_subtype_id = prefab_data['prefab_subtype_id']

            # Получаем имя prefab_type и prefab_subtype для сообщения
            type_response = requests.get(f'{DJANGO_API_URL}prefab_types/{prefab_type_id}')
            if type_response.status_code == 200:
                prefab_type_name = type_response.json().get('name', 'Неизвестный тип')
            else:
                prefab_type_name = 'Неизвестный тип'

            subtype_response = requests.get(f'{DJANGO_API_URL}prefab_subtypes/{prefab_subtype_id}')
            if subtype_response.status_code == 200:
                prefab_subtype_name = subtype_response.json().get('name', 'Неизвестный подтип')
            else:
                prefab_subtype_name = 'Неизвестный подтип'

            await update.message.reply_text(
                f"\U00002705 {prefab_type_name} — {prefab_subtype_name} — {quantity}шт. успешно отправлены на СГП.\n"
                f"Чтобы отправить их в Отгрузку, перейдите в соответствующую вкладку."
            )
            context.user_data['expecting_sgp_quantity'] = False

            # Вызываем send_main_menu
            user_data = requests.get(f'{DJANGO_API_URL}users/chat/{update.message.chat.id}/').json()
            full_name = user_data.get('full_name', 'Пользователь')
            organization_id = user_data.get('organization_id', None)
            await send_main_menu(update.message.chat.id, context, full_name, organization_id)

        except ValueError:
            await update.message.reply_text('Пожалуйста, введите корректное число.')
            context.user_data['expecting_sgp_quantity'] = True


async def send_prefabs_in_production(chat_id, context: ContextTypes.DEFAULT_TYPE):
    user_response = requests.get(f'{DJANGO_API_URL}users/chat/{chat_id}/')
    user_data = user_response.json()
    if user_response.status_code != 200:
        await context.bot.send_message(
            chat_id=chat_id,
            text="Ошибка при получении данных пользователя. Попробуйте снова."
        )
        await send_main_menu(chat_id, context, user_data.get('full_name', 'Пользователь'), user_data.get('organization_id', None))
        return

    organization_id = user_data.get('organization_id')

    if not organization_id:
        await context.bot.send_message(
            chat_id=chat_id,
            text="Ваша организация не определена. Пожалуйста, свяжитесь с администратором."
        )
        await send_main_menu(chat_id, context, user_data.get('full_name', 'Пользователь'), user_data.get('organization_id', None))
        return

    # Получаем все префабы в работе
    response = requests.get(f'{DJANGO_API_URL}prefabs_in_work/')
    if response.status_code == 200:
        prefabs_in_work = response.json()
        prefabs_in_production = [p for p in prefabs_in_work if p['status'] == 'production']

        if prefabs_in_production:
            # Группируем префабы по prefab_id и суммируем их количество
            grouped_prefabs = {}
            for prefab in prefabs_in_production:
                prefab_id = prefab['prefab_id']
                quantity = prefab.get('quantity', 0)
                if prefab_id in grouped_prefabs:
                    grouped_prefabs[prefab_id]['quantity'] += quantity
                else:
                    grouped_prefabs[prefab_id] = {'quantity': quantity, 'id': prefab['id']}

            keyboard = []
            for prefab_id, data in grouped_prefabs.items():
                # Получаем данные о префабе
                prefab_response = requests.get(f'{DJANGO_API_URL}prefabs/{prefab_id}')
                if prefab_response.status_code == 200:
                    prefab_data = prefab_response.json()
                    if prefab_data['organization_id'] != organization_id:
                        continue

                    prefab_type_id = prefab_data['prefab_type_id']
                    prefab_subtype_id = prefab_data['prefab_subtype_id']
                    object_id = prefab_data.get('object_id')

                    # Получаем имя prefab_type
                    type_response = requests.get(f'{DJANGO_API_URL}prefab_types/{prefab_type_id}')
                    if type_response.status_code == 200:
                        prefab_type_name = type_response.json().get('name', 'Неизвестный тип')
                    else:
                        prefab_type_name = 'Неизвестный тип'

                    # Получаем имя prefab_subtype
                    subtype_response = requests.get(f'{DJANGO_API_URL}prefab_subtypes/{prefab_subtype_id}')
                    if subtype_response.status_code == 200:
                        prefab_subtype_name = subtype_response.json().get('name', 'Неизвестный подтип')
                    else:
                        prefab_subtype_name = 'Неизвестный подтип'

                    # Получаем имя объекта
                    object_response = requests.get(f'{DJANGO_API_URL}objects/{object_id}')
                    if object_response.status_code == 200:
                        object_name = object_response.json().get('name', 'Неизвестный объект')
                    else:
                        object_name = 'Неизвестный объект'

                    button_text = f"{data['quantity']}шт. — {prefab_type_name} — {prefab_subtype_name} — {object_name}"
                    keyboard.append([InlineKeyboardButton(button_text, callback_data=f"sgp_prefab_{data['id']}")])

            if keyboard:
                keyboard.append([InlineKeyboardButton("Назад", callback_data='main_menu')])
                reply_markup = InlineKeyboardMarkup(keyboard)
                await context.bot.send_message(
                    chat_id=chat_id,
                    text="Выберите префаб для изменения статуса на СГП:",
                    reply_markup=reply_markup
                )
            else:
                await context.bot.send_message(
                    chat_id=chat_id,
                    text="Нет префабов в производстве."
                )
                await send_main_menu(chat_id, context, user_data.get('full_name', 'Пользователь'), user_data.get('organization_id', None))
        else:
            await context.bot.send_message(
                chat_id=chat_id,
                text="Нет префабов в производстве."
            )
            await send_main_menu(chat_id, context, user_data.get('full_name', 'Пользователь'), user_data.get('organization_id', None))
    else:
        await context.bot.send_message(
            chat_id=chat_id,
            text="Ошибка при получении префабов. Попробуйте снова."
        )
        await send_main_menu(chat_id, context, user_data.get('full_name', 'Пользователь'), user_data.get('organization_id', None))

async def handle_prefab_shipment_quantity(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if context.user_data.get('expecting_shipment_quantity'):
        try:
            quantity = int(update.message.text)
            selected_prefab_in_work_id = context.user_data['selected_prefab_in_work_id']

            # Получаем данные о выбранной записи prefabs_in_work
            prefabs_in_work_response = requests.get(f'{DJANGO_API_URL}prefabs_in_work/{selected_prefab_in_work_id}')
            if prefabs_in_work_response.status_code != 200:
                await update.message.reply_text('Ошибка при получении данных префаба. Попробуйте снова.')
                context.user_data['expecting_shipment_quantity'] = True
                return

            prefab_in_work_data = prefabs_in_work_response.json()
            selected_prefab_id = prefab_in_work_data['prefab_id']

            # Получаем все записи префабов с одинаковым prefab_id и статусом 'sgp'
            all_prefabs_in_work_response = requests.get(f'{DJANGO_API_URL}prefabs_in_work/')
            if all_prefabs_in_work_response.status_code != 200:
                await update.message.reply_text('Ошибка при получении данных префабов. Попробуйте снова.')
                context.user_data['expecting_shipment_quantity'] = True
                return

            all_prefabs_in_work = all_prefabs_in_work_response.json()
            prefabs_to_update = [p for p in all_prefabs_in_work if p['prefab_id'] == selected_prefab_id and p['status'] == 'sgp']

            # Вычисляем общее количество доступных префабов
            total_available_quantity = sum(p['quantity'] for p in prefabs_to_update)

            # Проверяем доступное количество
            if quantity > total_available_quantity:
                await update.message.reply_text(
                    f"Введенное количество ({quantity}) превышает доступное ({total_available_quantity}). Попробуйте снова."
                )
                context.user_data['expecting_shipment_quantity'] = True
                return

            # Обновляем префабы, разделяем их если необходимо
            update_quantity_remaining = quantity

            for prefab in prefabs_to_update:
                if update_quantity_remaining <= 0:
                    break

                current_quantity = prefab['quantity']
                if current_quantity <= update_quantity_remaining:
                    # Обновляем статус на shipment для текущего префаба
                    update_response = requests.put(
                        f'{DJANGO_API_URL}prefabs_in_work/{prefab["id"]}/',
                        json={'status': 'shipment', 'shipping_date': datetime.utcnow().isoformat()}
                    )
                    if update_response.status_code == 200:
                        update_quantity_remaining -= current_quantity
                    else:
                        await update.message.reply_text('Ошибка при обновлении статуса префаба. Попробуйте снова.')
                        context.user_data['expecting_shipment_quantity'] = True
                        return
                else:
                    # Разделяем префаб на два
                    new_quantity = current_quantity - update_quantity_remaining
                    update_response = requests.patch(
                        f'{DJANGO_API_URL}prefabs_in_work/{prefab["id"]}/',
                        json={'quantity': new_quantity}
                    )
                    if update_response.status_code == 200:
                        # Создаем новую запись с количеством и статусом shipment
                        new_prefabs_in_work_data = {
                            'prefab_id': prefab['prefab_id'],
                            'quantity': update_quantity_remaining,
                            'status': 'shipment',
                            'production_date': prefab['production_date'],
                            'sgp_date': prefab['sgp_date'],
                            'shipping_date': datetime.utcnow().isoformat()
                        }
                        new_prefab_response = requests.post(f'{DJANGO_API_URL}prefabs_in_work/', json=new_prefabs_in_work_data)
                        if new_prefab_response.status_code == 201:
                            update_quantity_remaining = 0
                        else:
                            await update.message.reply_text('Ошибка при создании записи. Попробуйте снова.')
                            context.user_data['expecting_shipment_quantity'] = True
                            return
                    else:
                        await update.message.reply_text('Ошибка при обновлении количества префаба. Попробуйте снова.')
                        context.user_data['expecting_shipment_quantity'] = True
                        return

                # Получаем данные о префабе
            prefab_data = requests.get(f'{DJANGO_API_URL}prefabs/{selected_prefab_id}').json()
            prefab_type_id = prefab_data['prefab_type_id']
            prefab_subtype_id = prefab_data['prefab_subtype_id']

            # Получаем имя prefab_type и prefab_subtype для сообщения
            type_response = requests.get(f'{DJANGO_API_URL}prefab_types/{prefab_type_id}')
            if type_response.status_code == 200:
                prefab_type_name = type_response.json().get('name', 'Неизвестный тип')
            else:
                prefab_type_name = 'Неизвестный тип'

            subtype_response = requests.get(f'{DJANGO_API_URL}prefab_subtypes/{prefab_subtype_id}')
            if subtype_response.status_code == 200:
                prefab_subtype_name = subtype_response.json().get('name', 'Неизвестный подтип')
            else:
                prefab_subtype_name = 'Неизвестный подтип'

            await update.message.reply_text(
                f"\U00002705 {prefab_type_name} — {prefab_subtype_name} — {quantity}шт. успешно отправлены в Отгрузке.\n"
                f"Ожидайте приемки на площадке."
            )
            context.user_data['expecting_shipment_quantity'] = False

            # Вызываем send_main_menu
            user_data = requests.get(f'{DJANGO_API_URL}users/chat/{update.message.chat.id}/').json()
            full_name = user_data.get('full_name', 'Пользователь')
            organization_id = user_data.get('organization_id', None)
            await send_main_menu(update.message.chat.id, context, full_name, organization_id)

        except ValueError:
            await update.message.reply_text('Пожалуйста, введите корректное число.')
            context.user_data['expecting_shipment_quantity'] = True


async def send_prefabs_list_for_shipment(chat_id, context: ContextTypes.DEFAULT_TYPE, warehouse_id: int):
    user_response = requests.get(f'{DJANGO_API_URL}users/chat/{chat_id}/')
    user_data = user_response.json()
    if user_response.status_code != 200:
        await context.bot.send_message(
            chat_id=chat_id,
            text="Ошибка при получении данных пользователя. Попробуйте снова."
        )
        await send_main_menu(chat_id, context, user_data.get('full_name', 'Пользователь'), user_data.get('organization_id', None))
        return

    organization_id = user_data.get('organization_id')
    user_object_id = user_data.get('object_id')

    if not organization_id or not user_object_id:
        await context.bot.send_message(
            chat_id=chat_id,
            text="Ваша организация или объект не определены. Пожалуйста, свяжитесь с администратором."
        )
        await send_main_menu(chat_id, context, user_data.get('full_name', 'Пользователь'), user_data.get('organization_id', None))
        return

    # Получаем все префабы со статусом "shipment"
    response = requests.get(f'{DJANGO_API_URL}prefabs_in_work/')
    if response.status_code == 200:
        prefabs_in_work = response.json()
        prefabs_in_work = [p for p in prefabs_in_work if p['status'] == 'shipment']

        if prefabs_in_work:
            # Группируем префабы по prefab_id и суммируем их количество
            grouped_prefabs = {}
            for prefab in prefabs_in_work:
                prefab_id = prefab['prefab_id']
                quantity = prefab.get('quantity', 0)
                if prefab_id in grouped_prefabs:
                    grouped_prefabs[prefab_id]['quantity'] += quantity
                else:
                    grouped_prefabs[prefab_id] = {'quantity': quantity, 'id': prefab['id']}

            keyboard = []
            for prefab_id, data in grouped_prefabs.items():
                # Получаем данные о префабе
                prefab_response = requests.get(f'{DJANGO_API_URL}prefabs/{prefab_id}')
                if prefab_response.status_code == 200:
                    prefab_data = prefab_response.json()

                    if prefab_data['object_id'] != user_object_id:
                        continue  # Пропускаем префабы, не относящиеся к объекту пользователя

                    prefab_type_id = prefab_data['prefab_type_id']
                    prefab_subtype_id = prefab_data['prefab_subtype_id']
                    prefab_org_id = prefab_data.get('organization_id')

                    # Получаем имя prefab_type
                    type_response = requests.get(f'{DJANGO_API_URL}prefab_types/{prefab_type_id}')
                    if type_response.status_code == 200:
                        prefab_type_name = type_response.json().get('name', 'Неизвестный тип')
                    else:
                        prefab_type_name = 'Неизвестный тип'

                    # Получаем имя prefab_subtype
                    subtype_response = requests.get(f'{DJANGO_API_URL}prefab_subtypes/{prefab_subtype_id}')
                    if subtype_response.status_code == 200:
                        prefab_subtype_name = subtype_response.json().get('name', 'Неизвестный подтип')
                    else:
                        prefab_subtype_name = 'Неизвестный подтип'

                    # Получаем имя организации
                    org_response = requests.get(f'{DJANGO_API_URL}organizations/{prefab_org_id}')
                    if org_response.status_code == 200:
                        org_name = org_response.json().get('organization', 'Неизвестная организация')
                    else:
                        org_name = 'Неизвестная организация'

                    button_text = f"{org_name} — {prefab_subtype_name} — {data['quantity']}шт."
                    keyboard.append([InlineKeyboardButton(button_text, callback_data=f"prefabinstock_{data['id']}")])

            if keyboard:
                keyboard.append([InlineKeyboardButton("Назад", callback_data='placespace')])
                reply_markup = InlineKeyboardMarkup(keyboard)
                await context.bot.send_message(
                    chat_id=chat_id,
                    text="Выберите префаб для отгрузки на склад:",
                    reply_markup=reply_markup
                )
            else:
                await context.bot.send_message(
                    chat_id=chat_id,
                    text="Нет префабов со статусом 'Отгружено'."
                )
                await send_main_menu(chat_id, context, user_data.get('full_name', 'Пользователь'), user_data.get('organization_id', None))
        else:
            await context.bot.send_message(
                chat_id=chat_id,
                text="Нет префабов со статусом 'Отгружено'."
            )
            await send_main_menu(chat_id, context, user_data.get('full_name', 'Пользователь'), user_data.get('organization_id', None))
    else:
        await context.bot.send_message(
            chat_id=chat_id,
            text="Ошибка при получении префабов. Попробуйте снова."
        )
        await send_main_menu(chat_id, context, user_data.get('full_name', 'Пользователь'), user_data.get('organization_id', None))

async def handle_prefab_stock_quantity(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if context.user_data.get('expecting_stock_quantity'):
        try:
            quantity = int(update.message.text)
            selected_prefab_id = context.user_data['selected_prefab_in_work_id']

            # Получаем данные о выбранной записи prefabs_in_work
            prefabs_in_work_response = requests.get(f'{DJANGO_API_URL}prefabs_in_work/{selected_prefab_id}')
            if prefabs_in_work_response.status_code != 200:
                await update.message.reply_text('Ошибка при получении данных префаба. Попробуйте снова.')
                context.user_data['expecting_stock_quantity'] = True
                return

            selected_prefab_id = requests.get(f'{DJANGO_API_URL}prefabs_in_work/{selected_prefab_id}').json()['prefab_id']
            all_prefabs_in_work = requests.get(f'{DJANGO_API_URL}prefabs_in_work/').json()
            prefabs_to_update = [p for p in all_prefabs_in_work if p['prefab_id'] == selected_prefab_id and p['status'] == 'shipment']

            # Вычисляем общее количество доступных префабов
            total_available_quantity = sum(p['quantity'] for p in prefabs_to_update)

            # Проверяем доступное количество
            if quantity > total_available_quantity:
                await update.message.reply_text(
                    f"Введенное количество ({quantity}) превышает доступное ({total_available_quantity}). Попробуйте снова."
                )
                context.user_data['expecting_stock_quantity'] = True
                return

            context.user_data['stock_quantity'] = quantity
            context.user_data['prefabs_to_update'] = prefabs_to_update
            context.user_data['total_available_quantity'] = total_available_quantity

            keyboard = [
                [InlineKeyboardButton("Принять", callback_data='acceptstockquantity')],
                [InlineKeyboardButton("Принять с замечаниями", callback_data='acceptstockquantity_with_comments')],
                [InlineKeyboardButton("Замечание", callback_data='remark_stockquantity')]
            ]
            reply_markup = InlineKeyboardMarkup(keyboard)

            await update.message.reply_text(
                f"Вы ввели {quantity} шт.\n\n"
                f"*Принять* - привязать префаб к складу,\n\n"
                f"*Принять с замечаниями* - привязать префаб к складу и отправить замечание заводу\n\n"
                f"*Замечание* - отправить замечание заводу вместе с дефектными префабами.",
                reply_markup=reply_markup,
                parse_mode=ParseMode.MARKDOWN
            )
            context.user_data['expecting_stock_quantity'] = False

        except ValueError:
            await update.message.reply_text('Пожалуйста, введите корректное число.')
            context.user_data['expecting_stock_quantity'] = True



async def send_to_google_sheets_warehouse(prefab_data):
    async with aiohttp.ClientSession() as session:
        # Получение данных о префабе
        prefab_id = prefab_data['prefab_id']
        prefab_info = await get_data_from_api(session, f"/prefabs/{prefab_id}")
        object_info = await get_data_from_api(session, f"/objects/{prefab_info['object_id']}")
        prefab_type_info = await get_data_from_api(session, f"/prefab_types/{prefab_info['prefab_type_id']}")
        prefab_subtype_info = await get_data_from_api(session, f"/prefab_subtypes/{prefab_info['prefab_subtype_id']}")
        warehouse_info = await get_data_from_api(session, f"/warehouses/{prefab_data['warehouse_id']}")



        # Получение информации об организации
        organization_id = prefab_info['organization_id']
        organization_info = await get_data_from_api(session, f"/organizations/{organization_id}")
        organization_name = organization_info['organization']


        # Используем текущую дату
        current_date = datetime.now().strftime("%d.%m.%Y")

        data = {
            "action": "warehouse",
            "A": prefab_data['id'],
            "B": prefab_subtype_info['name'],
            "C": prefab_type_info['name'],
            "D": current_date,
            "E": warehouse_info['name'],
            "F": object_info['name'],
            "G": None,
            "H": prefab_data['quantity'],
            "I": organization_name

        }

        async with session.post(WEBHOOK_URL, json=data) as response:
            response.raise_for_status()
            return await response.json()


async def handle_accept_stock_quantity(update: Update, context: ContextTypes.DEFAULT_TYPE, with_comments=False, remark=False):
    query = update.callback_query
    await query.answer()

    quantity = context.user_data['stock_quantity']
    prefabs_to_update = context.user_data['prefabs_to_update']
    warehouse_id = context.user_data['selected_warehouse_id']
    warehouse_name = requests.get(f'{DJANGO_API_URL}warehouses/{warehouse_id}/').json().get('name')
    new_status = 'stock' if not remark else 'remark'

    update_quantity_remaining = quantity
    updated_prefabs = []
    summarized_prefabs = {}

    for prefab in prefabs_to_update:
        if update_quantity_remaining <= 0:
            break

        current_quantity = prefab['quantity']
        prefab_in_work_id = prefab['id']
        prefab_id = prefab['prefab_id']

        # Получаем данные о префабе из таблицы prefabs
        prefab_details = requests.get(f'{DJANGO_API_URL}prefabs/{prefab_id}/').json()
        prefab_type_id = prefab_details.get('prefab_type_id')
        prefab_subtype_id = prefab_details.get('prefab_subtype_id')

        # Получаем названия типа и подтипа префаба
        prefab_type_name = requests.get(f'{DJANGO_API_URL}prefab_types/{prefab_type_id}/').json().get('name')
        prefab_subtype_name = requests.get(f'{DJANGO_API_URL}prefab_subtypes/{prefab_subtype_id}/').json().get('name')

        if current_quantity <= update_quantity_remaining:
            update_data = {'status': new_status, 'stock_date': datetime.now().isoformat()}
            if not remark:
                update_data['warehouse_id'] = warehouse_id

            update_response = requests.put(
                f'{DJANGO_API_URL}prefabs_in_work/{prefab_in_work_id}/',
                json=update_data
            )
            if update_response.status_code == 200:
                updated_prefabs.append(prefab_in_work_id)
                update_quantity_remaining -= current_quantity

                if not remark:
                    # Суммируем количество одинаковых префабов
                    key = f"{prefab_type_name} - {prefab_subtype_name}"
                    if key in summarized_prefabs:
                        summarized_prefabs[key] += current_quantity
                    else:
                        summarized_prefabs[key] = current_quantity

                    # Убираем отправку в Google Sheets, если статус remark
                    updated_prefab = update_response.json()
                    asyncio.create_task(send_to_google_sheets_warehouse(updated_prefab))
            else:
                await query.message.reply_text('Ошибка при обновлении статуса префаба. Попробуйте снова.')
                return
        else:
            new_quantity = current_quantity - update_quantity_remaining
            update_response = requests.patch(
                f'{DJANGO_API_URL}prefabs_in_work/{prefab_in_work_id}/',
                json={'quantity': new_quantity}
            )
            if update_response.status_code == 200:
                new_prefabs_in_work_data = {
                    'prefab_id': prefab_id,
                    'quantity': update_quantity_remaining,
                    'status': new_status,
                    'production_date': prefab['production_date'],
                    'sgp_date': prefab['sgp_date'],
                    'shipping_date': prefab['shipping_date'],
                    'stock_date': datetime.now().isoformat()
                }
                if not remark:
                    new_prefabs_in_work_data['warehouse_id'] = warehouse_id

                new_prefab_response = requests.post(f'{DJANGO_API_URL}prefabs_in_work/', json=new_prefabs_in_work_data)
                if new_prefab_response.status_code == 201:
                    new_prefab_in_work_id = new_prefab_response.json().get('id')
                    updated_prefabs.append(new_prefab_in_work_id)

                    if not remark:
                        # Суммируем количество одинаковых префабов
                        key = f"{prefab_type_name} - {prefab_subtype_name}"
                        if key in summarized_prefabs:
                            summarized_prefabs[key] += update_quantity_remaining
                        else:
                            summarized_prefabs[key] = update_quantity_remaining

                        # Убираем отправку в Google Sheets, если статус remark
                        new_prefab = new_prefab_response.json()
                        asyncio.create_task(send_to_google_sheets_warehouse(new_prefab))

                    update_quantity_remaining = 0
                else:
                    await query.message.reply_text('Ошибка при создании записи. Попробуйте снова.')
                    return
            else:
                await query.message.reply_text('Ошибка при обновлении количества префаба. Попробуйте снова.')
                return

    context.user_data['updated_prefabs'] = updated_prefabs

    if with_comments or remark:
        message_text = "Пожалуйста, введите комментарий для префабов:\n"
    else:
        message_text = "\U00002705 Префабы успешно отправлены на склад. Теперь добавьте фотографии (до 10) и нажмите /done для завершения.\n"

    # Формируем сообщение с суммированными данными
    for key, total_quantity in summarized_prefabs.items():
        message_text += f"\U00002705 {key}: {total_quantity} шт. успешно отправлены на {warehouse_name}\n"

    await query.message.reply_text(message_text)

    context.user_data['stage'] = 'attach_comments_prefab_in_work' if with_comments or remark else 'attach_photos_prefab_in_work'




async def send_warehouses_list_montage(chat_id, context: ContextTypes.DEFAULT_TYPE):
    response = requests.get(f'{DJANGO_API_URL}warehouses/')
    if response.status_code == 200:
        warehouses = response.json()
        if warehouses:
            keyboard = [[InlineKeyboardButton(warehouse['name'], callback_data=f'selectwarehouse_for_montage_{warehouse["id"]}')]
                        for warehouse in warehouses]
            keyboard.append([InlineKeyboardButton("Назад", callback_data='prefabsoptionlist')])
            reply_markup = InlineKeyboardMarkup(keyboard)
            await context.bot.send_message(
                chat_id=chat_id,
                text="Выберите склад для монтажа:",
                reply_markup=reply_markup
            )
        else:
            await context.bot.send_message(
                chat_id=chat_id,
                text="Склады не найдены."
            )
    else:
        await context.bot.send_message(
            chat_id=chat_id,
            text="Ошибка при получении списка складов. Попробуйте снова."
        )

async def send_prefab_types_montage(chat_id, context: ContextTypes.DEFAULT_TYPE):
    response = requests.get(f'{DJANGO_API_URL}prefab_types/')
    if response.status_code == 200:
        prefab_types = response.json()
        if prefab_types:
            keyboard = [[InlineKeyboardButton(prefab_type['name'], callback_data=f'select_prefab_type_for_montage_{prefab_type["id"]}')]
                        for prefab_type in prefab_types]
            keyboard.append([InlineKeyboardButton("Назад", callback_data='montage')])
            reply_markup = InlineKeyboardMarkup(keyboard)
            await context.bot.send_message(
                chat_id=chat_id,
                text="Выберите тип префаба:",
                reply_markup=reply_markup
            )
        else:
            await context.bot.send_message(
                chat_id=chat_id,
                text="Типы префабов не найдены."
            )
    else:
        await context.bot.send_message(
            chat_id=chat_id,
            text="Ошибка при получении списка типов префабов. Попробуйте снова."
        )

async def send_prefab_subtypes_montage(chat_id, context: ContextTypes.DEFAULT_TYPE):
    prefab_type_id = context.user_data['selected_prefab_type_id']
    response = requests.get(f'{DJANGO_API_URL}prefab_subtypes/?prefab_type_id={prefab_type_id}')
    if response.status_code == 200:
        prefab_subtypes = response.json()
        if prefab_subtypes:
            keyboard = [[InlineKeyboardButton(subtype['name'], callback_data=f'select_prefab_subtype_for_montage_{subtype["id"]}')]
                        for subtype in prefab_subtypes]
            keyboard.append([InlineKeyboardButton("Назад", callback_data='montage')])
            reply_markup = InlineKeyboardMarkup(keyboard)
            await context.bot.send_message(
                chat_id=chat_id,
                text="Выберите подтип префаба:",
                reply_markup=reply_markup
            )
        else:
            await context.bot.send_message(
                chat_id=chat_id,
                text="Подтипы префабов не найдены."
            )
    else:
        await context.bot.send_message(
            chat_id=chat_id,
            text="Ошибка при получении списка подтипов префабов. Попробуйте снова."
        )

async def send_prefabs_list_montage(chat_id, context: ContextTypes.DEFAULT_TYPE):
    prefab_type_id = context.user_data['selected_prefab_type_id']
    prefab_subtype_id = context.user_data['selected_prefab_subtype_id']
    warehouse_id = context.user_data['selected_warehouse_id']

    response = requests.get(f'{DJANGO_API_URL}prefabs/')
    if response.status_code == 200:
        all_prefabs = response.json()
        prefabs = [p for p in all_prefabs if p['prefab_type_id'] == prefab_type_id and p['prefab_subtype_id'] == prefab_subtype_id]

        if prefabs:
            prefab_in_stock_dict = {}
            for prefab in prefabs:
                prefab_id = prefab['id']
                prefabs_in_work_response = requests.get(f'{DJANGO_API_URL}prefabs_in_work/')
                if prefabs_in_work_response.status_code == 200:
                    all_prefabs_in_work = prefabs_in_work_response.json()
                    prefabs_in_stock_filtered = [p for p in all_prefabs_in_work if p['prefab_id'] == prefab_id and p['status'] == 'stock' and p['warehouse_id'] == warehouse_id]

                    for prefab_in_stock in prefabs_in_stock_filtered:
                        prefab_in_work_id = prefab_in_stock['id']
                        quantity = prefab_in_stock['quantity']

                        if prefab_id not in prefab_in_stock_dict:
                            prefab_in_stock_dict[prefab_id] = {
                                'quantity': 0,
                                'prefabs_in_work_ids': []
                            }
                        prefab_in_stock_dict[prefab_id]['quantity'] += quantity
                        prefab_in_stock_dict[prefab_id]['prefabs_in_work_ids'].append({
                            'id': prefab_in_work_id,
                            'quantity': quantity
                        })

            prefabs_in_stock = []
            for prefab_id, data in prefab_in_stock_dict.items():
                prefab = next(p for p in prefabs if p['id'] == prefab_id)
                prefab_type_response = requests.get(f'{DJANGO_API_URL}prefab_types/{prefab_type_id}')
                prefab_subtype_response = requests.get(f'{DJANGO_API_URL}prefab_subtypes/{prefab_subtype_id}')
                organization_response = requests.get(f'{DJANGO_API_URL}organizations/{prefab["organization_id"]}')

                if prefab_type_response.status_code == 200 and prefab_subtype_response.status_code == 200 and organization_response.status_code == 200:
                    prefab_type_name = prefab_type_response.json()['name']
                    prefab_subtype_name = prefab_subtype_response.json()['name']
                    organization_name = organization_response.json()['organization']

                    button_text = f"{data['quantity']} шт. - {prefab_type_name} - {prefab_subtype_name} - {organization_name}"
                    prefabs_in_stock.append(
                        [InlineKeyboardButton(button_text, callback_data=f'prefabin_stock_{prefab_id}')]
                    )

            if prefabs_in_stock:
                prefabs_in_stock.append([InlineKeyboardButton("Назад", callback_data='montage')])
                reply_markup = InlineKeyboardMarkup(prefabs_in_stock)
                await context.bot.send_message(
                    chat_id=chat_id,
                    text="Выберите префаб со статусом 'на складе':",
                    reply_markup=reply_markup
                )
            else:
                await context.bot.send_message(
                    chat_id=chat_id,
                    text="Нет префабов со статусом 'на складе'."
                )
                user_response = requests.get(f'{DJANGO_API_URL}users/chat/{chat_id}')
                if user_response.status_code == 200:
                    user_data = user_response.json()
                    full_name = user_data['full_name']
                    organization_id = user_data['organization_id']
                    await send_main_menu(chat_id, context, full_name, organization_id)
                else:
                    await context.bot.send_message(chat_id=chat_id, text="Ошибка при получении данных пользователя.")
        else:
            await context.bot.send_message(
                chat_id=chat_id,
                text="Префабы не найдены."
            )
            user_response = requests.get(f'{DJANGO_API_URL}users/chat/{chat_id}')
            if user_response.status_code == 200:
                user_data = user_response.json()
                full_name = user_data['full_name']
                organization_id = user_data['organization_id']
                await send_main_menu(chat_id, context, full_name, organization_id)
            else:
                await context.bot.send_message(chat_id=chat_id, text="Ошибка при получении данных пользователя.")
    else:
        await context.bot.send_message(
            chat_id=chat_id,
            text="Ошибка при получении списка префабов. Попробуйте снова."
        )
        user_response = requests.get(f'{DJANGO_API_URL}users/chat/{chat_id}')
        if user_response.status_code == 200:
            user_data = user_response.json()
            full_name = user_data['full_name']
            organization_id = user_data['organization_id']
            await send_main_menu(chat_id, context, full_name, organization_id)
        else:
            await context.bot.send_message(chat_id=chat_id, text="Ошибка при получении данных пользователя.")


async def handle_montage_quantity(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    user_input = update.message.text
    if not user_input.isdigit():
        await update.message.reply_text("Пожалуйста, введите допустимое количество.")
        return

    quantity = int(user_input)
    prefab_id = context.user_data['selected_prefab_id']

    response = requests.get(f'{DJANGO_API_URL}prefabs_in_work/')
    if response.status_code != 200:
        await update.message.reply_text("Ошибка при получении префабов. Попробуйте снова.")
        return

    all_prefabs_in_work = response.json()
    prefabs_in_work = [p for p in all_prefabs_in_work if p['prefab_id'] == prefab_id and p['status'] == 'stock']

    total_quantity = sum(p['quantity'] for p in prefabs_in_work)

    if total_quantity < quantity:
        prefab_info_response = requests.get(f'{DJANGO_API_URL}prefabs/{prefab_id}')
        if prefab_info_response.status_code == 200:
            prefab_info = prefab_info_response.json()
            prefab_subtype_id = prefab_info.get('prefab_subtype_id')
            if prefab_subtype_id:
                prefab_subtype_response = requests.get(f'{DJANGO_API_URL}prefab_subtypes/{prefab_subtype_id}')
                if prefab_subtype_response.status_code == 200:
                    prefab_subtype_info = prefab_subtype_response.json()
                    prefab_name = prefab_subtype_info.get('name', 'неизвестный подтип префаба')
                else:
                    prefab_name = 'неизвестный подтип префаба'
            else:
                prefab_name = 'неизвестный подтип префаба'
        else:
            prefab_name = 'неизвестный префаб'

        await update.message.reply_text(f"Недостаточно префабов на складе. Максимальное количество: {total_quantity} ({prefab_name})")
        return

    remaining_quantity = quantity
    prefabs_for_montage = []

    for prefab in prefabs_in_work:
        if remaining_quantity <= 0:
            break

        if prefab['quantity'] <= remaining_quantity:
            prefabs_for_montage.append({
                'id': prefab['id'],
                'quantity': prefab['quantity'],
                'status': 'montage'
            })
            remaining_quantity -= prefab['quantity']
        else:
            prefabs_for_montage.append({
                'id': prefab['id'],
                'quantity': remaining_quantity,
                'status': 'montage'
            })
            remaining_quantity = 0

    context.user_data['prefabs_for_montage'] = prefabs_for_montage
    context.user_data['prefabs_in_work'] = prefabs_in_work  # Сохраняем данные в контексте для последующего использования
    await send_block_sections_list(update.message.chat.id, context)


async def send_block_sections_list(chat_id, context: ContextTypes.DEFAULT_TYPE):
    user_chat_id = chat_id
    user_response = requests.get(f'{DJANGO_API_URL}users/chat/{user_chat_id}')

    if user_response.status_code == 200:
        user_data = user_response.json()
        organization_id = user_data['organization_id']
        object_id = user_data['object_id']

        if organization_id is None:
            await context.bot.send_message(
                chat_id=chat_id,
                text="Ошибка: пользователь не принадлежит ни одной организации."
            )
            return

        response = requests.get(f'{DJANGO_API_URL}objects/{object_id}/blocksections/')
        if response.status_code == 200:
            block_sections = response.json()
            filtered_block_sections = [bs for bs in block_sections if bs['object_id'] == object_id]

            if filtered_block_sections:
                keyboard = [
                    [InlineKeyboardButton(bs['name'], callback_data=f'select_block_section_{bs["id"]}')]
                    for bs in filtered_block_sections
                ]
                reply_markup = InlineKeyboardMarkup(keyboard)
                await context.bot.send_message(
                    chat_id=chat_id,
                    text="Выберите секцию:",
                    reply_markup=reply_markup
                )
            else:
                await context.bot.send_message(
                    chat_id=chat_id,
                    text="Секции не найдены."
                )
        else:
            await context.bot.send_message(
                chat_id=chat_id,
                text="Ошибка при получении списка секций. Попробуйте снова."
            )
    else:
        await context.bot.send_message(
            chat_id=chat_id,
            text="Ошибка при получении данных пользователя. Попробуйте снова."
        )


async def send_floors_list(chat_id, context: ContextTypes.DEFAULT_TYPE):
    block_section_id = context.user_data['selected_block_section_id']
    response = requests.get(f'{DJANGO_API_URL}blocksections/{block_section_id}')
    if response.status_code == 200:
        block_section = response.json()
        if block_section and block_section['number_of_floors']:
            number_of_floors_bottom = block_section['number_of_floors_bottom']
            number_of_floors = block_section['number_of_floors']

            # Генерация кнопок этажей в две колонки, исключая 0
            keyboard = []
            for i in range(number_of_floors_bottom, number_of_floors + 1):
                if i == 0:
                    continue
                if len(keyboard) == 0 or len(keyboard[-1]) == 2:
                    keyboard.append([InlineKeyboardButton(f'Этаж {i}', callback_data=f'select_floor_{i}')])
                else:
                    keyboard[-1].append(InlineKeyboardButton(f'Этаж {i}', callback_data=f'select_floor_{i}'))

            # Добавление кнопки кровли
            keyboard.append([InlineKeyboardButton('Кровля', callback_data='select_floor_roof')])

            reply_markup = InlineKeyboardMarkup(keyboard)
            await context.bot.send_message(
                chat_id=chat_id,
                text="Выберите этаж:",
                reply_markup=reply_markup
            )
        else:
            await context.bot.send_message(
                chat_id=chat_id,
                text="Этажи не найдены."
            )
    else:
        await context.bot.send_message(
            chat_id=chat_id,
            text="Ошибка при получении списка этажей. Попробуйте снова."
        )


async def handle_select_block_section(query: Update, context: ContextTypes.DEFAULT_TYPE):
    await query.message.delete()
    block_section_id = int(query.data.split('_')[-1])
    context.user_data['selected_block_section_id'] = block_section_id
    await send_floors_list(query.message.chat.id, context)


async def send_to_google_sheets_montage(prefab_data):
    async with aiohttp.ClientSession() as session:
        # Получение данных о префабе
        prefab_id = prefab_data['prefab_id']
        prefab_info = await get_data_from_api(session, f"/prefabs/{prefab_id}")
        object_info = await get_data_from_api(session, f"/objects/{prefab_info['object_id']}")
        prefab_type_info = await get_data_from_api(session, f"/prefab_types/{prefab_info['prefab_type_id']}")
        prefab_subtype_info = await get_data_from_api(session, f"/prefab_subtypes/{prefab_info['prefab_subtype_id']}")
        warehouse_info = await get_data_from_api(session, f"/warehouses/{prefab_data['warehouse_id']}")
        block_section_info = await get_data_from_api(session, f"/blocksections/{prefab_data['block_section_id']}")

        # Получение информации об организации
        organization_id = prefab_info['organization_id']
        organization_info = await get_data_from_api(session, f"/organizations/{organization_id}")
        organization_name = organization_info['organization']

        # Используем текущую дату
        current_date = datetime.now().strftime("%d.%m.%Y")

        data = {

            "action": "montage",
            "A": prefab_data['id'],
            "B": object_info['name'],
            "C": current_date,
            "D": block_section_info['name'],
            "E": prefab_data['floor'],
            "F": prefab_type_info['name'],
            "G": prefab_subtype_info['name'],
            "H": prefab_data['quantity'],
            "I": None,
            "J": warehouse_info['name'],
            "K": organization_name

        }

        async with session.post(WEBHOOK_URL, json=data) as response:
            response.raise_for_status()
            return await response.json()

async def handle_select_floor(query: Update, context: ContextTypes.DEFAULT_TYPE):
    await query.message.delete()
    floor = query.data.split('_')[-1]
    context.user_data['selected_floor'] = floor

    prefabs_for_montage = context.user_data['prefabs_for_montage']
    block_section_id = context.user_data['selected_block_section_id']

    for prefab in prefabs_for_montage:
        prefab_in_work = next(p for p in context.user_data['prefabs_in_work'] if p['id'] == prefab['id'])
        if prefab['quantity'] < prefab_in_work['quantity']:
            remaining_stock_quantity = prefab_in_work['quantity'] - prefab['quantity']
            # Создаем новый префаб для оставшегося количества со статусом 'stock'
            new_prefab_data = {
                'prefab_id': prefab_in_work['prefab_id'],
                'quantity': remaining_stock_quantity,
                'status': 'stock',
                'warehouse_id': prefab_in_work['warehouse_id'],
                'production_date': prefab_in_work.get('production_date'),
                'sgp_date': prefab_in_work.get('sgp_date'),
                'shipping_date': prefab_in_work.get('shipping_date'),
                'stock_date': prefab_in_work.get('stock_date')

            }
            response = requests.post(f'{DJANGO_API_URL}prefabs_in_work/', json=new_prefab_data)
            if response.status_code != 201:
                await query.message.reply_text("Ошибка при создании нового префаба со статусом 'stock'. Попробуйте снова.")
                return

            # Обновляем количество и статус текущего префаба
            response = requests.patch(f'{DJANGO_API_URL}prefabs_in_work/{prefab["id"]}', json={'quantity': prefab['quantity'], 'status': 'montage', 'block_section_id': block_section_id, 'floor': floor, 'montage_date': datetime.utcnow().isoformat()})
            if response.status_code != 200:
                await query.message.reply_text("Ошибка при обновлении статуса и количества. Попробуйте снова.")
                return

            # Отправка данных в Google Sheets
            updated_prefab = response.json()
            asyncio.create_task(send_to_google_sheets_montage(updated_prefab))
        else:
            response = requests.patch(f'{DJANGO_API_URL}prefabs_in_work/{prefab["id"]}', json={'status': 'montage', 'block_section_id': block_section_id, 'floor': floor, 'montage_date': datetime.utcnow().isoformat()})
            if response.status_code != 200:
                await query.message.reply_text("Ошибка при обновлении статуса. Попробуйте снова.")
                return

            # Отправка данных в Google Sheets
            updated_prefab = response.json()
            asyncio.create_task(send_to_google_sheets_montage(updated_prefab))

    await query.message.reply_text("Количество на складе обновлено. Теперь добавьте фотографии (до 10) и нажмите /done для завершения.")
    context.user_data['stage'] = 'attach_photos_prefab_in_montage'




async def handle_prefab_photo_upload(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    prefab_ids = [p['id'] for p in context.user_data['prefabs_for_montage']]
    photos = context.user_data.get('photos', [])
    file_id = update.message.photo[-1].file_id
    photos.append(file_id)
    context.user_data['photos'] = photos

    if 'last_photo_message_id' in context.user_data:
        try:
            await context.bot.delete_message(
                chat_id=update.message.chat.id,
                message_id=context.user_data['last_photo_message_id']
            )
        except telegram.error.BadRequest as e:
            if str(e) != "Message to delete not found":
                raise

    if len(photos) < 10:
        reply_keyboard = [
            [KeyboardButton("/done")]
        ]
        reply_markup = ReplyKeyboardMarkup(reply_keyboard, resize_keyboard=True, one_time_keyboard=False)

        message = await update.message.reply_text(
            f'Фотография {len(photos)} из 10 успешно загружена. Прикрепите ещё или нажмите /done для завершения.',
            reply_markup=reply_markup
        )
        context.user_data['last_photo_message_id'] = message.message_id
    else:
        await finalize_photo_montage(update, context)


async def finalize_photo_montage(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    prefab_ids = [p['id'] for p in context.user_data['prefabs_for_montage']]
    new_photos = context.user_data.get('photos', [])

    for prefab_id in prefab_ids:
        response = requests.get(f'{DJANGO_API_URL}prefabs_in_work/{prefab_id}')
        if response.status_code != 200:
            await update.message.reply_text('Ошибка при получении существующих фотографий. Попробуйте снова.')
            return

        prefab = response.json()
        existing_photos = prefab.get('photos', [])
        all_photos = existing_photos + new_photos

        update_data = {'photos': all_photos}
        response = requests.patch(f'{DJANGO_API_URL}prefabs_in_work/{prefab_id}', json=update_data)
        if response.status_code != 200:
            await update.message.reply_text('Ошибка при загрузке фотографий. Попробуйте снова.')
            return

    reply_markup_kb_main = ReplyKeyboardMarkup(reply_keyboard_main, resize_keyboard=True, one_time_keyboard=False)

    await update.message.reply_text("\U00002705 Фотографии успешно загружены. Монтаж завершён.",
                                    reply_markup=reply_markup_kb_main)

    context.user_data['stage'] = None
    context.user_data['photos'] = []

    user_id = context.user_data.get('user_id')
    if not user_id:
        user_id = update.message.from_user.id

    response = requests.get(f'{DJANGO_API_URL}users/chat/{user_id}')
    if response.status_code == 200:
        user_data = response.json()
        full_name = user_data['full_name']
        organization_id = user_data['organization_id']
        await send_main_menu(update.message.chat.id, context, full_name, organization_id)
    else:
        await update.message.reply_text('Ошибка при получении данных пользователя.')


async def handle_support_question(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    user_id = update.message.from_user.id
    user_response = requests.get(f'{DJANGO_API_URL}users/chat/{user_id}/')

    if user_response.status_code == 200:
        user_data = user_response.json()
        question = update.message.text

        ticket_data = {
            "sender_id": user_data['id'],
            "question": question,
        }
        logger.info(f"Sending support ticket data: {ticket_data}")
        response = requests.post(f'{DJANGO_API_URL}support_tickets/', json=ticket_data)
        if response.status_code == 201:
            context.user_data['ticket_id'] = response.json()['id']
            context.user_data['stage'] = 'attach_photos_support_ticket'
            reply_keyboard = [
                [KeyboardButton("/done")]
            ]
            reply_markup = ReplyKeyboardMarkup(reply_keyboard, resize_keyboard=True, one_time_keyboard=False)
            await update.message.reply_text(
                "Ваш вопрос отправлен в тех. поддержку. Прикрепите фотографии (если нужно) или нажмите /done для завершения.",
                reply_markup=reply_markup
            )
        else:
            logger.error(f"Error creating support ticket: {response.status_code} {response.text}")
            await update.message.reply_text("Ошибка при отправке вопроса. Попробуйте снова.")
    else:
        logger.error(f"Error fetching user data: {user_response.status_code} {user_response.text}")
        await update.message.reply_text("Ошибка при получении данных пользователя. Попробуйте снова.")


async def handle_ticket_selection(update: Update, context: ContextTypes.DEFAULT_TYPE, ticket_id: int) -> None:
    query = update.callback_query
    response = requests.get(f'{DJANGO_API_URL}support_tickets/{ticket_id}/')
    if response.status_code != 200:
        await query.message.reply_text('Ошибка при получении данных запроса.')
        return

    ticket_data = response.json()
    sender_id = ticket_data["sender_id"]

    user_response = requests.get(f'{DJANGO_API_URL}users/{sender_id}/')
    if user_response.status_code == 200:
        user_data = user_response.json()
        full_name = user_data['full_name']
        organization_id = user_data["organization_id"]
        org_response = requests.get(f'{DJANGO_API_URL}organizations/{organization_id}/')
        if org_response.status_code == 200:
            organization_name = org_response.json().get('organization', 'Неизвестная организация')
        else:
            organization_name = 'Неизвестная организация'
    else:
        full_name = "Неизвестный пользователь"
        organization_name = "Неизвестная организация"

    created_at = ticket_data['created_at'].split('T')[0]

    ticket_info = (f"Запрос от *{full_name} — {organization_name}*\n"
                   f"Дата: {created_at}\n"
                   f"Вопрос: {ticket_data['question']}\n"

                   )
    photo_ids = ticket_data.get('photo_ids', [])
    media_group = []

    for idx, photo_id in enumerate(photo_ids):
        if photo_id:
            if idx == 0:
                media_group.append(InputMediaPhoto(media=photo_id, caption=ticket_info, parse_mode=ParseMode.MARKDOWN))
            else:
                media_group.append(InputMediaPhoto(media=photo_id))

    keyboard = [
        [InlineKeyboardButton("Ответить", callback_data=f"answer_{ticket_id}")],
        [InlineKeyboardButton("Назад", callback_data="main_menu")]
    ]
    reply_markup = InlineKeyboardMarkup(keyboard)

    if media_group:
        await context.bot.send_media_group(chat_id=query.message.chat.id, media=media_group, parse_mode=ParseMode.MARKDOWN,)
        await context.bot.send_message(
            chat_id=query.message.chat.id,
            text="Выберите действие:",
            reply_markup=reply_markup
        )
    else:
        await query.message.reply_text(ticket_info, parse_mode=ParseMode.MARKDOWN)
        await context.bot.send_message(
            chat_id=query.message.chat.id,
            text="Выберите действие:",
            reply_markup=reply_markup
        )

async def handle_support_answer(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    ticket_id = context.user_data.get('ticket_id')
    answer = update.message.text

    response = requests.get(f'{DJANGO_API_URL}support_tickets/{ticket_id}').json()
    sender_question = response['question']
    # Получаем ID пользователя
    user_response = requests.get(f'{DJANGO_API_URL}users/chat/{update.message.chat_id}/')
    if user_response.status_code == 200:
        user_data = user_response.json()
        respondent_id = user_data['id']
    else:
        await update.message.reply_text('Ошибка при получении данных пользователя.')
        return

    # Обновляем запрос техподдержки
    update_data = {
        'answer': answer,
        'respondent_id': respondent_id,
        'status': 'closed'
    }
    response = requests.patch(f'{DJANGO_API_URL}support_tickets/{ticket_id}', json=update_data)
    if response.status_code == 200:
        await update.message.reply_text('Ответ отправлен. Запрос закрыт.')

        context.user_data['stage'] = None
        # Отправка уведомления отправителю
        ticket_data = response.json()
        sender_id = ticket_data["sender_id"]
        sender_response = requests.get(f'{DJANGO_API_URL}users/{sender_id}/')
        if sender_response.status_code == 200:
            sender_data = sender_response.json()
            sender_chat_id = sender_data['chat_id']

            await context.bot.send_message(chat_id=sender_chat_id, text=f'*Ваш запрос получил ответ.*\n\n'
                                                                        f'*Ваш запрос:* {sender_question} \n\n'
                                                                        f'*Технический специалист направил Вам ответ:* {answer}',
                                           parse_mode=ParseMode.MARKDOWN)
        else:
            await update.message.reply_text('Ошибка при отправке уведомления отправителю.')
    else:
        await update.message.reply_text('Ошибка при обновлении запроса техподдержки.')

    # Возврат в главное меню
    await send_main_menu(update.message.chat_id, context, user_data['full_name'], user_data['organization_id'])

async def get_remarks_for_factory(organization_id: int):
    # Получаем все префабы
    response = requests.get(f'{DJANGO_API_URL}prefabs/')
    if response.status_code != 200:
        return []

    prefabs = response.json()
    prefab_ids = [prefab['id'] for prefab in prefabs if prefab['organization_id'] == organization_id]

    # Получаем все prefabs_in_work и фильтруем по статусу "remark" и prefab_id
    response = requests.get(f'{DJANGO_API_URL}prefabs_in_work/')
    if response.status_code != 200:
        return []

    all_prefabs_in_work = response.json()
    remarks = [prefab_in_work for prefab_in_work in all_prefabs_in_work if prefab_in_work['status'] == 'remark' and prefab_in_work['prefab_id'] in prefab_ids]

    return remarks

async def send_remarks(chat_id, context, organization_id):
    remarks = await get_remarks_for_factory(organization_id)

    if not remarks:
        await context.bot.send_message(
            chat_id=chat_id,
            text="Нет замечаний для вашего завода."
        )
        # Вызов функции send_main_menu после отправки сообщения
        user_data = requests.get(f'{DJANGO_API_URL}users/chat/{chat_id}/').json()
        full_name = user_data.get('full_name', 'Пользователь')
        await send_main_menu(chat_id, context, full_name, organization_id)
        return

    grouped_remarks = {}
    for remark in remarks:
        prefab_id = remark['prefab_id']
        if prefab_id not in grouped_remarks:
            grouped_remarks[prefab_id] = {
                'quantity': 0,
                'remarks': []
            }
        grouped_remarks[prefab_id]['quantity'] += remark['quantity']
        grouped_remarks[prefab_id]['remarks'].append(remark)

    keyboard = []
    for prefab_id, data in grouped_remarks.items():
        prefab_response = requests.get(f'{DJANGO_API_URL}prefabs/{prefab_id}')
        if prefab_response.status_code != 200:
            continue

        prefab = prefab_response.json()
        prefab_type_response = requests.get(f'{DJANGO_API_URL}prefab_types/{prefab["prefab_type_id"]}')
        prefab_subtype_response = requests.get(f'{DJANGO_API_URL}prefab_subtypes/{prefab["prefab_subtype_id"]}')
        object_response = requests.get(f'{DJANGO_API_URL}objects/{prefab["object_id"]}')

        if prefab_type_response.status_code != 200 or prefab_subtype_response.status_code != 200 or object_response.status_code != 200:
            continue

        prefab_type = prefab_type_response.json()
        prefab_subtype = prefab_subtype_response.json()
        object_ = object_response.json()

        button_text = f'{data["quantity"]} шт. - {prefab_type["name"]} - {prefab_subtype["name"]} - {object_["name"]}'
        keyboard.append([InlineKeyboardButton(button_text, callback_data=f'remark_{prefab_id}')])

    keyboard.append([InlineKeyboardButton("Назад", callback_data='main_menu')])
    reply_markup = InlineKeyboardMarkup(keyboard)
    await context.bot.send_message(
        chat_id=chat_id,
        text="Замечания:",
        reply_markup=reply_markup
    )


async def remarks_prefab_newstatus(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if context.user_data.get('expecting_new_status_quantity'):
        try:
            quantity = int(update.message.text)
            selected_prefab_id = context.user_data['selected_prefab_id']
            selected_stage = context.user_data['selected_stage']

            all_prefabs_in_work_response = requests.get(f'{DJANGO_API_URL}prefabs_in_work/')
            if all_prefabs_in_work_response.status_code != 200:
                await update.message.reply_text('Ошибка при получении данных префабов. Попробуйте снова.')
                context.user_data['expecting_new_status_quantity'] = True
                return
            all_prefabs_in_work = all_prefabs_in_work_response.json()
            prefabs_to_update = [p for p in all_prefabs_in_work if p['prefab_id'] == selected_prefab_id and p['status'] == 'remark']
            print(prefabs_to_update)
            total_available_quantity = sum(p['quantity'] for p in prefabs_to_update)
            if quantity > total_available_quantity:
                await update.message.reply_text(
                    f"Введенное количество ({quantity}) превышает доступное ({total_available_quantity}). Попробуйте снова."
                )
                context.user_data['expecting_new_status_quantity'] = True
                return

            update_quantity_remaining = quantity
            updated_prefabs = []

            for prefab in prefabs_to_update:
                if update_quantity_remaining <= 0:
                    break

                current_quantity = prefab['quantity']
                if current_quantity <= update_quantity_remaining:
                    update_data = {'status': selected_stage}
                    update_response = requests.put(
                        f'{DJANGO_API_URL}prefabs_in_work/{prefab["id"]}/',
                        json=update_data
                    )
                    if update_response.status_code == 200:
                        updated_prefabs.append(prefab["id"])
                        update_quantity_remaining -= current_quantity
                    else:
                        await update.message.reply_text('Ошибка при обновлении статуса префаба. Попробуйте снова.')
                        return
                else:
                    new_quantity = current_quantity - update_quantity_remaining
                    update_response = requests.patch(
                        f'{DJANGO_API_URL}prefabs_in_work/{prefab["id"]}/',
                        json={'quantity': new_quantity}
                    )
                    if update_response.status_code == 200:
                        new_prefabs_in_work_data = {
                            'prefab_id': prefab['prefab_id'],
                            'quantity': update_quantity_remaining,
                            'status': selected_stage,
                            'production_date': prefab['production_date'],
                            'sgp_date': prefab['sgp_date'],
                            'shipping_date': prefab['shipping_date']
                        }
                        new_prefab_response = requests.post(f'{DJANGO_API_URL}prefabs_in_work/', json=new_prefabs_in_work_data)
                        if new_prefab_response.status_code == 201:
                            new_prefab_in_work_id = new_prefab_response.json().get('id')
                            updated_prefabs.append(new_prefab_in_work_id)
                            update_quantity_remaining = 0
                        else:
                            await update.message.reply_text('Ошибка при создании записи. Попробуйте снова.')
                            return
                    else:
                        await update.message.reply_text('Ошибка при обновлении количества префаба. Попробуйте снова.')
                        return

            context.user_data['updated_prefabs'] = updated_prefabs

            await update.message.reply_text(
                "\U00002705 Префабы успешно обновлены."
            )
            context.user_data['expecting_new_status_quantity'] = False

            # Вызываем send_main_menu
            user_data = requests.get(f'{DJANGO_API_URL}users/chat/{update.message.chat.id}/').json()
            full_name = user_data.get('full_name', 'Пользователь')
            organization_id = user_data.get('organization_id', None)
            await send_main_menu(update.message.chat.id, context, full_name, organization_id)

        except ValueError:
            await update.message.reply_text('Пожалуйста, введите корректное число.')
            context.user_data['expecting_new_status_quantity'] = True

async def choose_stage_for_prefab(query, context, remark_id):
    keyboard = [
        [InlineKeyboardButton("Производство", callback_data=f'prefab_stage_production_{remark_id}')],
        [InlineKeyboardButton("СГП", callback_data=f'prefab_stage_sgp_{remark_id}')],
        [InlineKeyboardButton("Отгрузка", callback_data=f'prefab_stage_shipment_{remark_id}')]
    ]
    reply_markup = InlineKeyboardMarkup(keyboard)

    await query.message.reply_text("Выберите на какой этап вернуть префаб:", reply_markup=reply_markup)


async def handle_remark_quantity(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    user_input = update.message.text

    if not user_input.isdigit():
        await update.message.reply_text("Пожалуйста, введите корректное число.")
        return

    quantity = int(user_input)
    prefab_in_work_id = context.user_data.get('prefab_in_work_id')
    stage = context.user_data.get('selected_stage')
    chat_id = update.message.chat_id

    # Получаем информацию о префабе
    response = requests.get(f'{DJANGO_API_URL}prefabs_in_work/{prefab_in_work_id}')
    if response.status_code != 200:
        await update.message.reply_text("Ошибка при получении данных префаба.")
        return

    prefab_in_work = response.json()

    if quantity > prefab_in_work['quantity']:
        await update.message.reply_text("Введенное количество превышает количество префабов. Попробуйте еще раз.")
        return
    elif quantity == prefab_in_work['quantity']:
        # Обновляем статус префаба
        payload = {"status": stage}
        response = requests.patch(f'{DJANGO_API_URL}prefabs_in_work/{prefab_in_work_id}', json=payload)
        if response.status_code == 200:
            await update.message.reply_text("Префаб переведен на новый статус")
        else:
            await update.message.reply_text("Ошибка при обновлении статуса префаба.")
    else:
        # Обновляем количество у старого префаба
        new_quantity = prefab_in_work['quantity'] - quantity
        response = requests.patch(f'{DJANGO_API_URL}prefabs_in_work/{prefab_in_work_id}',
                                  json={"quantity": new_quantity})
        if response.status_code != 200:
            await update.message.reply_text("Ошибка при обновлении количества префаба.")
            return

        # Создаем новый префаб с новым количеством и статусом
        new_prefab = prefab_in_work.copy()
        new_prefab['quantity'] = quantity
        new_prefab['status'] = stage
        del new_prefab['id']  # Удаляем id чтобы создать новый префаб
        response = requests.post(f'{DJANGO_API_URL}prefabs_in_work/', json=new_prefab)
        if response.status_code == 201:
            await update.message.reply_text(f"Префаб в указанном количестве {quantity} отправлен на выбранный статус.")
        else:
            await update.message.reply_text("Ошибка при создании нового префаба.")

    # Получаем данные пользователя для передачи в send_main_menu
    user_response = requests.get(f'{DJANGO_API_URL}users/chat/{chat_id}/')
    if user_response.status_code == 200:
        user_data = user_response.json()
        full_name = user_data.get('full_name', 'Пользователь')
        organization_id = user_data.get('organization_id')
        await send_main_menu(chat_id, context, full_name, organization_id)
    else:
        await update.message.reply_text("Ошибка при получении данных пользователя.")

    # Сбрасываем состояние
    context.user_data.pop('expecting_remark_quantity', None)
    context.user_data.pop('prefab_in_work_id', None)
    context.user_data.pop('selected_stage', None)


async def send_prefabs_in_stage(chat_id, context, stage):
    if 'user_id' not in context.user_data:
        user_id = chat_id  # Используем chat_id в качестве user_id
        context.user_data['user_id'] = user_id
    else:
        user_id = context.user_data['user_id']

    # Получаем информацию о пользователе по user_id
    user_response = requests.get(f'{DJANGO_API_URL}users/chat/{user_id}/')
    if user_response.status_code != 200:
        await context.bot.send_message(
            chat_id=chat_id,
            text="Ошибка при получении данных пользователя. Попробуйте снова."
        )
        return
    user_data = user_response.json()
    user_organization_id = user_data['organization_id']

    # Получаем все префабы в работе
    response = requests.get(f'{DJANGO_API_URL}prefabs_in_work/')
    if response.status_code != 200:
        await context.bot.send_message(
            chat_id=chat_id,
            text="Ошибка при получении данных префабов. Попробуйте снова."
        )
        return

    prefabs_in_work = response.json()
    prefabs_in_stage = [p for p in prefabs_in_work if p['status'] == stage]

    grouped_prefabs = {}
    for prefab in prefabs_in_stage:
        prefab_id = prefab['prefab_id']

        # Получаем информацию о префабе, чтобы узнать организацию
        prefab_response = requests.get(f'{DJANGO_API_URL}prefabs/{prefab_id}/')
        if prefab_response.status_code != 200:
            continue

        prefab_data = prefab_response.json()
        if prefab_data['organization_id'] != user_organization_id:
            continue

        if prefab_id not in grouped_prefabs:
            grouped_prefabs[prefab_id] = {
                'quantity': 0,
                'prefabs': []
            }
        grouped_prefabs[prefab_id]['quantity'] += prefab['quantity']
        grouped_prefabs[prefab_id]['prefabs'].append(prefab)

    keyboard = []
    for prefab_id, data in grouped_prefabs.items():
        prefab_response = requests.get(f'{DJANGO_API_URL}prefabs/{prefab_id}')
        if prefab_response.status_code != 200:
            continue

        prefab = prefab_response.json()
        prefab_type_response = requests.get(f'{DJANGO_API_URL}prefab_types/{prefab["prefab_type_id"]}')
        prefab_subtype_response = requests.get(f'{DJANGO_API_URL}prefab_subtypes/{prefab["prefab_subtype_id"]}')
        object_response = requests.get(f'{DJANGO_API_URL}objects/{prefab["object_id"]}')

        if prefab_type_response.status_code != 200 or prefab_subtype_response.status_code != 200 or object_response.status_code != 200:
            continue

        prefab_type = prefab_type_response.json()
        prefab_subtype = prefab_subtype_response.json()
        object_ = object_response.json()

        button_text = f'{data["quantity"]} шт. — {prefab_type["name"]} — {prefab_subtype["name"]} — {object_["name"]}'
        keyboard.append([InlineKeyboardButton(button_text, callback_data=f'refactorprefab_{prefab_id}')])

    keyboard.append([InlineKeyboardButton("Назад", callback_data='edit_prefab')])
    reply_markup = InlineKeyboardMarkup(keyboard)
    await context.bot.send_message(
        chat_id=chat_id,
        text=f"Префабы в выбранном статусе:",
        reply_markup=reply_markup
    )



async def handle_refactor_prefab_quantity(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if context.user_data.get('refactor_prefab_count'):
        try:
            quantity = int(update.message.text)
            context.user_data['quantity'] = quantity
            selected_prefab_id = context.user_data['selected_prefab_id']

            all_prefabs_in_work = requests.get(f'{DJANGO_API_URL}prefabs_in_work/').json()
            prefabs_to_update = [p for p in all_prefabs_in_work if p['prefab_id'] == selected_prefab_id and p['status'] == context.user_data['selected_stage']]

            total_available_quantity = sum(p['quantity'] for p in prefabs_to_update)

            if quantity > total_available_quantity:
                await update.message.reply_text(
                    f"Введенное количество ({quantity}) превышает доступное ({total_available_quantity}). Попробуйте снова."
                )
                context.user_data['refactor_prefab_count'] = True
                return

            update_quantity_remaining = quantity
            updated_prefabs = []

            for prefab in prefabs_to_update:
                if update_quantity_remaining <= 0:
                    break

                current_quantity = prefab['quantity']
                if current_quantity <= update_quantity_remaining:
                    updated_prefabs.append(prefab["id"])
                    update_quantity_remaining -= current_quantity
                else:
                    new_quantity = current_quantity - update_quantity_remaining
                    update_response = requests.patch(
                        f'{DJANGO_API_URL}prefabs_in_work/{prefab["id"]}/',
                        json={'quantity': new_quantity}
                    )
                    if update_response.status_code == 200:
                        new_prefabs_in_work_data = {
                            'prefab_id': prefab['prefab_id'],
                            'quantity': update_quantity_remaining,
                            'status': context.user_data['selected_stage'],
                            'production_date': prefab['production_date'],
                            'sgp_date': prefab['sgp_date'],
                            'shipping_date': prefab['shipping_date']
                        }
                        new_prefab_response = requests.post(f'{DJANGO_API_URL}prefabs_in_work/', json=new_prefabs_in_work_data)
                        if new_prefab_response.status_code == 201:
                            new_prefab_in_work_id = new_prefab_response.json().get('id')
                            updated_prefabs.append(new_prefab_in_work_id)
                            update_quantity_remaining = 0
                        else:
                            await update.message.reply_text('Ошибка при создании записи. Попробуйте снова.')
                            return
                    else:
                        await update.message.reply_text('Ошибка при обновлении количества префаба. Попробуйте снова.')
                        return

            context.user_data['updated_prefabs'] = updated_prefabs
            context.user_data['expecting_new_status_prefab'] = True

            keyboard = [
                [InlineKeyboardButton("Тендер", callback_data='new_status_tender')],
                [InlineKeyboardButton("Производство", callback_data='new_status_production')],
                [InlineKeyboardButton("СГП", callback_data='new_status_sgp')],
                [InlineKeyboardButton("Отгрузка", callback_data='new_status_shipment')]
            ]

            if 'refactor_message_id' in context.user_data:
                try:
                    await context.bot.delete_message(
                        chat_id=update.message.chat.id,
                        message_id=context.user_data['refactor_message_id']
                    )
                except:
                    pass

            reply_markup = InlineKeyboardMarkup(keyboard)
            await update.message.reply_text("Выберите новый статус для префаба:", reply_markup=reply_markup)
        except ValueError:
            await update.message.reply_text('Пожалуйста, введите корректное число.')
            context.user_data['refactor_prefab_count'] = True
        finally:
            context.user_data['refactor_prefab_count'] = False

async def handle_new_status(update: Update, context: ContextTypes.DEFAULT_TYPE):
    query = update.callback_query
    await query.answer()

    selected_status = query.data.split('_')[2]
    updated_prefabs = context.user_data['updated_prefabs']

    # Если выбран статус "Тендер", просто вычитаем количество и возвращаемся в главное меню
    if selected_status == 'tender':
        print(context.user_data)
        update_quantity_remaining = context.user_data['quantity']
        for prefab_id in updated_prefabs:
            prefab_data = requests.get(f'{DJANGO_API_URL}prefabs_in_work/{prefab_id}/').json()
            current_quantity = prefab_data['quantity']

            if current_quantity <= update_quantity_remaining:
                # Удаляем запись
                delete_response = requests.delete(f'{DJANGO_API_URL}prefabs_in_work/{prefab_id}/')
                if delete_response.status_code != 200:
                    await query.message.reply_text(f'Ошибка при удалении префаба ID {prefab_id}. Попробуйте снова.')
                    return
                update_quantity_remaining -= current_quantity
            else:
                # Обновляем запись с новым количеством
                new_quantity = current_quantity - update_quantity_remaining
                update_response = requests.patch(
                    f'{DJANGO_API_URL}prefabs_in_work/{prefab_id}/',
                    json={'quantity': new_quantity}
                )
                if update_response.status_code != 200:
                    await query.message.reply_text(f'Ошибка при обновлении количества префаба ID {prefab_id}. Попробуйте снова.')
                    return
                update_quantity_remaining = 0

            if update_quantity_remaining <= 0:
                break

        await query.message.reply_text(f"\U00002705 Введенное количество успешно вычтено как тендер.")
        context.user_data['expecting_new_status_prefab'] = False

        # Вызываем send_main_menu
        user_data = requests.get(f'{DJANGO_API_URL}users/chat/{query.message.chat.id}/').json()
        full_name = user_data.get('full_name', 'Пользователь')
        organization_id = user_data.get('organization_id', None)
        await send_main_menu(query.message.chat.id, context, full_name, organization_id)
        return

    # Инициализируем переменные для хранения информации о типе и подвиде префабов
    prefab_type_name = 'Неизвестный тип'
    prefab_subtype_name = 'Неизвестный подтип'
    total_quantity = 0

    # Обновляем статус префабов
    for prefab_id in updated_prefabs:
        update_data = {'status': selected_status}
        response = requests.put(f'{DJANGO_API_URL}prefabs_in_work/{prefab_id}/', json=update_data)
        if response.status_code != 200:
            await query.message.reply_text(f'Ошибка при обновлении статуса для префаба ID {prefab_id}. Попробуйте снова.')
            return

        # Получаем данные о префабе для сообщения
        prefab_response = requests.get(f'{DJANGO_API_URL}prefabs_in_work/{prefab_id}/')
        if prefab_response.status_code == 200:
            prefab_data = prefab_response.json()
            prefab_id = prefab_data['prefab_id']
            total_quantity += prefab_data['quantity']  # Увеличиваем общее количество

            prefab_info_response = requests.get(f'{DJANGO_API_URL}prefabs/{prefab_id}')
            if prefab_info_response.status_code == 200:
                prefab_info = prefab_info_response.json()
                prefab_type_id = prefab_info['prefab_type_id']
                prefab_subtype_id = prefab_info['prefab_subtype_id']

                type_response = requests.get(f'{DJANGO_API_URL}prefab_types/{prefab_type_id}')
                if type_response.status_code == 200:
                    prefab_type_name = type_response.json().get('name', 'Неизвестный тип')

                subtype_response = requests.get(f'{DJANGO_API_URL}prefab_subtypes/{prefab_subtype_id}')
                if subtype_response.status_code == 200:
                    prefab_subtype_name = subtype_response.json().get('name', 'Неизвестный подтип')

    # Формируем и отправляем сообщение
    await query.message.reply_text(
        f"\U00002705 {prefab_type_name} — {prefab_subtype_name} - {total_quantity} штук переведены на статус '{selected_status}'."
    )
    context.user_data['expecting_new_status_prefab'] = False

    # Вызываем send_main_menu
    user_data = requests.get(f'{DJANGO_API_URL}users/chat/{query.message.chat.id}/').json()
    full_name = user_data.get('full_name', 'Пользователь')
    organization_id = user_data.get('organization_id', None)
    await send_main_menu(query.message.chat.id, context, full_name, organization_id)

async def send_objects_list(chat_id, context: ContextTypes.DEFAULT_TYPE):
    user_response = requests.get(f'{DJANGO_API_URL}users/chat/{chat_id}/')
    if user_response.status_code == 200:
        user_data = user_response.json()
        organization_id = user_data.get('organization_id')

        if organization_id:
            prefabs_response = requests.get(f'{DJANGO_API_URL}prefabs/')
            if prefabs_response.status_code == 200:
                prefabs = prefabs_response.json()
                object_ids = list(set([prefab['object_id'] for prefab in prefabs if prefab['organization_id'] == organization_id]))

                if object_ids:
                    keyboard = []
                    for object_id in object_ids:
                        object_response = requests.get(f'{DJANGO_API_URL}objects/{object_id}/')
                        if object_response.status_code == 200:
                            object_data = object_response.json()
                            object_name = object_data.get('name', 'Неизвестный объект')
                            keyboard.append([InlineKeyboardButton(object_name, callback_data=f'selectobjectprefabs_{object_id}')])

                    reply_markup = InlineKeyboardMarkup(keyboard)
                    await context.bot.send_message(
                        chat_id=chat_id,
                        text="Выберите объект для получения сводки:",
                        reply_markup=reply_markup
                    )
                else:
                    await context.bot.send_message(
                        chat_id=chat_id,
                        text="Нет объектов для отображения."
                    )
            else:
                await context.bot.send_message(
                    chat_id=chat_id,
                    text="Ошибка при получении списка префабов. Попробуйте снова."
                )
        else:
            await context.bot.send_message(
                chat_id=chat_id,
                text="Ошибка при получении данных пользователя."
            )
    else:
        await context.bot.send_message(
            chat_id=chat_id,
            text="Ошибка при получении данных пользователя."
        )

STATUS_TRANSLATION = {
    'production': 'в производстве',
    'sgp': 'СГП',
    'shipment': 'отгружено',
    'stock': 'на складе',
    'montage': 'в монтаже'
}

STATUS_ORDER = ['production', 'sgp', 'shipment', 'stock', 'montage']

async def send_prefab_summary(chat_id, context: ContextTypes.DEFAULT_TYPE, object_id: int):
    # Получаем сводку по префабам с помощью нового эндпоинта
    summary_response = requests.get(f'{DJANGO_API_URL}prefab_summary/{chat_id}/{object_id}')

    if summary_response.status_code != 200:
        await context.bot.send_message(
            chat_id=chat_id,
            text="Ошибка при получении сводки по префабам."
        )
        return

    summary_data = summary_response.json().get("summary", [])

    if not summary_data:
        await context.bot.send_message(
            chat_id=chat_id,
            text="Нет данных по префабам для выбранного объекта."
        )
        return

    # Формируем текст сообщения для отправки пользователю
    summary_text = f"📊 Сводка префабов по объекту {summary_data[0]['object_name']}:\n\n"
    for prefab in summary_data:
        summary_text += (
            f"📋 {prefab['prefab_subtype_name']} (количество по тендеру: {prefab['prefab_quantity']}):\n"
            f"  {STATUS_TRANSLATION.get('production', 'production')}: {prefab['production_quantity']}\n"
            f"  {STATUS_TRANSLATION.get('sgp', 'sgp')}: {prefab['sgp_quantity']}\n"
            f"  {STATUS_TRANSLATION.get('shipment', 'shipment')}: {prefab['shipment_quantity']}\n"
            "\n"
        )

    # Отправляем сформированный текст пользователю
    await context.bot.send_message(
        chat_id=chat_id,
        text=summary_text
    )


API_URL = "http://127.0.0.1:8000"
DATABASE_URL = "postgresql://postgres:qwerty22375638@176.123.163.235:5432/tgfrontbrusnika"


async def report_today_pdf(chat_id, context):
    async with aiohttp.ClientSession() as session:
        # Получение object_id по chat_id
        async with session.get(f"{API_URL}/users/chat/{chat_id}") as response:
            if response.status != 200:
                await context.bot.send_message(chat_id, 'Ошибка при получении данных пользователя. Попробуйте позже.')
                return

            user_data = await response.json()
            object_id = user_data.get('object_id')

            if not object_id:
                await context.bot.send_message(chat_id,
                                               'Не удалось получить object_id. Пожалуйста, проверьте данные пользователя.')
                return

            selected_date = datetime.today().strftime('%Y-%m-%d')  # Текущая дата

            # Подключение к базе данных
            conn = psycopg2.connect(DATABASE_URL)
            cur = conn.cursor()

            # Создание пустого Excel файла
            random_number = random.randint(10000000, 99999999)
            excel_file = f'report_{random_number}.xlsx'

            with ExcelWriter(excel_file) as writer:
                ### Статусы
                query_statuses = f"SELECT * FROM get_statuses_report('{selected_date}', {object_id})"
                df_statuses = pd.read_sql(query_statuses, conn)
                df_statuses.to_excel(writer, index=False, sheet_name='Статусы')

                ### Площадка
                query_warehouse = f"SELECT * FROM get_warehouse_report('{selected_date}', {object_id})"
                df_warehouse = pd.read_sql(query_warehouse, conn)
                df_warehouse.to_excel(writer, index=False, sheet_name='Площадка')

                ### Монтаж
                query_montage = f"SELECT * FROM get_montage_report('{selected_date}', {object_id})"
                df_montage = pd.read_sql(query_montage, conn)
                df_montage.to_excel(writer, index=False, sheet_name='Монтаж')

            # Закрытие соединения с базой данных
            conn.close()

            print(f"Данные успешно сохранены в файл {excel_file}")

            ### Преобразование Excel в Word с горизонтальной ориентацией
            doc = Document()
            selected_date = datetime.strptime(selected_date, '%Y-%m-%d').strftime('%d.%m.%Y')
            doc.add_heading(f'Отчет за {selected_date}', 0)

            # Установка альбомной ориентации для всех разделов документа
            section = doc.sections[-1]
            section.page_width = Inches(11.69)  # A4 landscape width
            section.page_height = Inches(8.27)  # A4 landscape height
            section.orientation = 1  # Альбомная ориентация

            # Чтение каждого листа из Excel-файла и добавление его в Word-документ
            for sheet_name in ['Статусы', 'Площадка', 'Монтаж']:
                df = pd.read_excel(excel_file, sheet_name=sheet_name)
                doc.add_heading(sheet_name, level=1)

                # Преобразование DataFrame в таблицу Word
                table = doc.add_table(rows=df.shape[0] + 1, cols=df.shape[1])

                # Добавление заголовков
                for j, col in enumerate(df.columns):
                    cell = table.cell(0, j)
                    cell.text = col
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.bold = True
                            run.font.size = Pt(10)  # Установка меньшего размера шрифта

                # Заполнение таблицы данными
                for i in range(df.shape[0]):
                    for j in range(df.shape[1]):
                        cell = table.cell(i + 1, j)
                        cell.text = str(df.iat[i, j]) if df.iat[i, j] != "" else ""
                        cell.paragraphs[0].alignment = 1  # Выравнивание по центру
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.size = Pt(10)  # Установка меньшего размера шрифта

                # Добавление границ к таблице
                table.style = 'Table Grid'

                doc.add_paragraph()  # Разделение между таблицами

            # Сохранение Word-документа
            word_file = f'report_{random_number}.docx'
            doc.save(word_file)
            print(f"Данные успешно сохранены в файл {word_file}")

            # Конвертация в PDF с помощью LibreOffice
            pdf_file = f'Отчет_от_{selected_date}_{random_number}.pdf'

            if platform.system() == "Windows":
                libreoffice_path = "C:\\Program Files\\LibreOffice\\program\\soffice.exe"
            else:
                libreoffice_path = "libreoffice"  # для Linux предполагается, что LibreOffice доступен в PATH


            subprocess.run([libreoffice_path, '--headless', '--convert-to', 'pdf', word_file])

            # Переименовать результат в 'temp_<random_number>.pdf'
            os.rename(word_file.replace('.docx', '.pdf'), pdf_file)

            # Отправка PDF пользователю
            with open(pdf_file, 'rb') as pdf_file_obj:
                await context.bot.send_document(chat_id, pdf_file_obj)

            # Удаление временных файлов
            os.remove(excel_file)
            os.remove(word_file)
            os.remove(pdf_file)

            print(f"Файл {pdf_file} успешно отправлен и временные файлы удалены")

async def report_specific_day_pdf(chat_id, context, selected_date):
    async with aiohttp.ClientSession() as session:
        # Получение object_id по chat_id
        async with session.get(f"{API_URL}/users/chat/{chat_id}") as response:
            if response.status != 200:
                await context.bot.send_message(chat_id, 'Ошибка при получении данных пользователя. Попробуйте позже.')
                return

            user_data = await response.json()
            object_id = user_data.get('object_id')

            if not object_id:
                await context.bot.send_message(chat_id,
                                               'Не удалось получить object_id. Пожалуйста, проверьте данные пользователя.')
                return

            # Подключение к базе данных
            conn = psycopg2.connect(DATABASE_URL)
            cur = conn.cursor()

            # Создание пустого Excel файла
            random_number = random.randint(10000000, 99999999)
            excel_file = f'report_{random_number}.xlsx'

            with ExcelWriter(excel_file) as writer:
                ### Статусы
                query_statuses = f"SELECT * FROM get_statuses_report('{selected_date}', {object_id})"
                df_statuses = pd.read_sql(query_statuses, conn)
                df_statuses.to_excel(writer, index=False, sheet_name='Статусы')

                ### Площадка
                query_warehouse = f"SELECT * FROM get_warehouse_report('{selected_date}', {object_id})"
                df_warehouse = pd.read_sql(query_warehouse, conn)
                df_warehouse.to_excel(writer, index=False, sheet_name='Площадка')

                ### Монтаж
                query_montage = f"SELECT * FROM get_montage_report('{selected_date}', {object_id})"
                df_montage = pd.read_sql(query_montage, conn)
                df_montage.to_excel(writer, index=False, sheet_name='Монтаж')

            # Закрытие соединения с базой данных
            conn.close()

            print(f"Данные успешно сохранены в файл {excel_file}")

            # Преобразование Excel в PDF аналогично `report_today_pdf`
            await convert_excel_to_pdf_and_send(excel_file, random_number, chat_id, context, selected_date)


async def convert_excel_to_pdf_and_send(excel_file, random_number, chat_id, context, selected_date):
    # Преобразование Excel в Word с горизонтальной ориентацией
    doc = Document()
    selected_date = datetime.strptime(selected_date, '%Y-%m-%d').strftime('%d.%m.%Y')
    doc.add_heading(f'Отчет {selected_date}', 0)

    # Установка альбомной ориентации для всех разделов документа
    section = doc.sections[-1]
    section.page_width = Inches(11.69)  # A4 landscape width
    section.page_height = Inches(8.27)  # A4 landscape height
    section.orientation = 1  # Альбомная ориентация

    # Чтение каждого листа из Excel-файла и добавление его в Word-документ
    excel_workbook = load_workbook(excel_file)
    for sheet_name in excel_workbook.sheetnames:
        df = pd.read_excel(excel_file, sheet_name=sheet_name)
        doc.add_heading(sheet_name, level=1)

        # Преобразование DataFrame в таблицу Word
        table = doc.add_table(rows=df.shape[0] + 1, cols=df.shape[1])

        # Добавление заголовков
        for j, col in enumerate(df.columns):
            cell = table.cell(0, j)
            cell.text = col
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.size = Pt(10)  # Установка меньшего размера шрифта

        # Заполнение таблицы данными
        for i in range(df.shape[0]):
            for j in range(df.shape[1]):
                cell = table.cell(i + 1, j)
                cell.text = str(df.iat[i, j]) if df.iat[i, j] != "" else ""
                cell.paragraphs[0].alignment = 1  # Выравнивание по центру
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(10)  # Установка меньшего размера шрифта

        # Добавление границ к таблице
        table.style = 'Table Grid'

        doc.add_paragraph()  # Разделение между таблицами

    # Сохранение Word-документа
    word_file = f'report_{random_number}.docx'
    doc.save(word_file)
    print(f"Данные успешно сохранены в файл {word_file}")

    # Конвертация в PDF с помощью LibreOffice
    pdf_file = f'Отчет_от_{selected_date}_{random_number}.pdf'

    if platform.system() == "Windows":
        libreoffice_path = "C:\\Program Files\\LibreOffice\\program\\soffice.exe"
    else:
        libreoffice_path = "libreoffice"  # для Linux предполагается, что LibreOffice доступен в PATH

    # Конвертация документа Word в PDF
    subprocess.run([libreoffice_path, '--headless', '--convert-to', 'pdf', word_file])

    # Переименовать результат в 'temp_<random_number>.pdf'
    os.rename(word_file.replace('.docx', '.pdf'), pdf_file)

    # Отправка PDF пользователю
    with open(pdf_file, 'rb') as pdf_file_obj:
        await context.bot.send_document(chat_id, pdf_file_obj)

    # Удаление временных файлов
    os.remove(excel_file)
    os.remove(word_file)
    os.remove(pdf_file)

    print(f"Файл {pdf_file} успешно отправлен и временные файлы удалены")



async def button(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    query = update.callback_query
    await query.answer()

    data = query.data
    user_id = query.from_user.id

    if data.startswith('org_'):
        org_id = int(data.split('_')[1])

        # Получаем текущие данные пользователя
        response = requests.get(f'{DJANGO_API_URL}users/chat/{user_id}/')
        if response.status_code == 200:
            user_data = response.json()
            user_data['organization_id'] = org_id  # Обновляем поле organization_id

            logger.info(f"Отправка данных в API: {json.dumps(user_data, indent=2)}")
            response = requests.put(f'{DJANGO_API_URL}users/{user_data["id"]}/', json=user_data)
            logger.info(f"Ответ от API: {response.status_code}, {response.text}")

            if response.status_code == 200:
                reply_keyboard = [
                    [KeyboardButton("/info")],
                    [KeyboardButton("/start")],
                    [KeyboardButton("/choice")]
                ]
                reply_markup_kb = ReplyKeyboardMarkup(reply_keyboard, resize_keyboard=True, one_time_keyboard=False)
                await query.message.delete()
                await query.message.reply_text(
                    'Организация успешно выбрана!',
                    reply_markup=reply_markup_kb
                )
                user_data = response.json()
                await send_main_menu(query.message.chat.id, context, user_data['full_name'], user_data['organization_id'])
            else:
                await query.message.reply_text('Ошибка при сохранении данных. Попробуйте снова.')
        else:
            await query.message.reply_text('Ошибка при получении данных пользователя. Попробуйте снова.')


    elif data.startswith('work_'):
        work_type_id = int(data.split('_')[1])
        await choose_block_section(query, context, work_type_id)

    elif data.startswith('block_'):
        block_section_id = int(data.split('_')[1])
        await choose_floor(query, context, block_section_id)

    elif data.startswith('floor_'):
        floor = data.split('_')[1]
        await confirm_transfer_data(query, context, floor)

    elif data == 'confirm_yes':
        await query.edit_message_reply_markup(reply_markup=None)
        await handle_transfer_confirmation(query, context, confirmed=True)

    elif data == 'confirm_no':
        # await query.message.delete()  # Удаление предыдущего сообщения
        await handle_transfer_confirmation(query, context, confirmed=False)

    elif data == 'front_menu':
        await query.message.delete()
        keyboard = [
            [InlineKeyboardButton("\U0001F91D Передать фронт", callback_data='transfer')],
            [InlineKeyboardButton("\U0001F9F1 Принять фронт", callback_data='accept_fronts')],
            [InlineKeyboardButton("Завершить фронт", callback_data='endfront')],
            [InlineKeyboardButton("Назад", callback_data='main_menu')]
        ]
        reply_markup = InlineKeyboardMarkup(keyboard)
        await query.message.reply_text('Выберите действие для фронта:', reply_markup=reply_markup)

    elif data == 'frontbutton':
        await query.message.delete()
        keyboard = [
            [InlineKeyboardButton("\U0001F4C4 Просмотр переданных фронтов", callback_data='view_fronts')],
            [InlineKeyboardButton("\U0001F6E0 Просмотр фронтов в работе", callback_data='fronts_in_process')],
            [InlineKeyboardButton("\U0001F4CB Выдать фронт", callback_data='issue_front')],
            [InlineKeyboardButton("Назад", callback_data='main_menu')]
        ]
        reply_markup = InlineKeyboardMarkup(keyboard)
        await query.message.reply_text('Выберите действие для фронта:', reply_markup=reply_markup)

    elif data == 'workforce_menu':
        await query.message.delete()
        keyboard = [
            [InlineKeyboardButton("\U0001F477 Передать численность", callback_data='workforce_transfer')],
            [InlineKeyboardButton("\U0000270F Редактировать численность", callback_data='workforce_refactor')],
            [InlineKeyboardButton("\U0000274C Удалить запись", callback_data='workforce_delete')],
            [InlineKeyboardButton("Назад", callback_data='main_menu')]
        ]
        reply_markup = InlineKeyboardMarkup(keyboard)
        await query.message.reply_text('Выберите действие для численности:', reply_markup=reply_markup)

    elif data == 'transfer':
        await query.message.delete()
        # Удаление сообщения с основным меню
        if 'main_menu_message_id' in context.user_data:
            try:
                await context.bot.delete_message(
                    chat_id=query.message.chat.id,
                    message_id=context.user_data['main_menu_message_id']
                )

            except telegram.error.BadRequest:
                logger.warning("Message to delete not found")

        response = requests.get(f'{DJANGO_API_URL}users/chat/{user_id}/')
        if response.status_code == 200:
            user_data = response.json()
            organization_id = user_data.get('organization_id')
            if not user_data['organization_id']:
                await query.message.reply_text('Пожалуйста, выберите организацию командой /choice.')
                return

        else:
            await query.message.reply_text('Ошибка при получении данных пользователя. Попробуйте снова.')
            return

        # Получаем данные организации
        response_org = requests.get(f'{DJANGO_API_URL}organizations/{organization_id}/')
        if response_org.status_code == 200:
            organization_data = response_org.json()
            organization_object_ids = organization_data.get('object_ids', [])
        else:
            await query.message.reply_text('Ошибка при получении данных организации. Попробуйте снова.')
            return

        response = requests.get(f'{DJANGO_API_URL}objects/')
        if response.status_code == 200:
            objects = response.json()

            filtered_objects = [obj for obj in objects if obj['id'] in organization_object_ids]

            if filtered_objects:
                keyboard = [
                    [InlineKeyboardButton(obj['name'], callback_data=f'obj_{obj["id"]}')] for obj in filtered_objects
                ]
                keyboard.append([InlineKeyboardButton("Назад", callback_data='main_menu')])
                reply_markup = InlineKeyboardMarkup(keyboard)
                await query.message.reply_text('Выберите объект:', reply_markup=reply_markup)
                context.user_data['stage'] = 'choose_object'
            else:
                await query.message.reply_text('Нет доступных объектов для данной организации.')

        else:
            await query.message.reply_text('Ошибка при получении списка объектов. Попробуйте снова.')

    elif data.startswith('obj_'):
        user_id = query.from_user.id  # Получаем user_id из данных обновления
        object_id = int(data.split('_')[1])

        # Получаем все фронты, где статус "in_process"
        response = requests.get(f'{DJANGO_API_URL}fronttransfers/?status=in_process')
        if response.status_code == 200:
            fronts = response.json()
            print(fronts)
            # Фильтруем фронты, где пользователь является отправителем
            user_fronts = [front for front in fronts if front['sender_chat_id'] == str(user_id)]
            user_has_fronts_in_object = any(front['object_id'] == object_id for front in user_fronts)

            if user_has_fronts_in_object:
                await choose_existing_front(query, context, user_fronts, object_id)
            else:
                await choose_work_type(query, context, object_id)
        else:
            await query.message.reply_text('Ошибка при получении списка фронтов. Попробуйте снова.')

    elif data == 'endfront':
        await query.message.delete()
        # Удаление сообщения с основным меню
        if 'main_menu_message_id' in context.user_data:
            try:
                await context.bot.delete_message(
                    chat_id=query.message.chat.id,
                    message_id=context.user_data['main_menu_message_id']
                )

            except telegram.error.BadRequest:
                logger.warning("Message to delete not found")

        response = requests.get(f'{DJANGO_API_URL}users/chat/{user_id}/')
        if response.status_code == 200:
            user_data = response.json()
            organization_id = user_data.get('organization_id')
            if not user_data['organization_id']:
                await query.message.reply_text('Пожалуйста, выберите организацию командой /choice.')
                return

        else:
            await query.message.reply_text('Ошибка при получении данных пользователя. Попробуйте снова.')
            return

        # Получаем данные организации
        response_org = requests.get(f'{DJANGO_API_URL}organizations/{organization_id}/')
        if response_org.status_code == 200:
            organization_data = response_org.json()
            organization_object_ids = organization_data.get('object_ids', [])
        else:
            await query.message.reply_text('Ошибка при получении данных организации. Попробуйте снова.')
            return

        response = requests.get(f'{DJANGO_API_URL}objects/')
        if response.status_code == 200:
            objects = response.json()

            filtered_objects = [obj for obj in objects if obj['id'] in organization_object_ids]

            if filtered_objects:
                keyboard = [
                    [InlineKeyboardButton(obj['name'], callback_data=f'endfrontobj_{obj["id"]}')] for obj in filtered_objects
                ]
                keyboard.append([InlineKeyboardButton("Назад", callback_data='main_menu')])
                reply_markup = InlineKeyboardMarkup(keyboard)
                await query.message.reply_text('Выберите объект:', reply_markup=reply_markup)
                context.user_data['stage'] = 'choose_object'
            else:
                await query.message.reply_text('Нет доступных объектов для данной организации.')

        else:
            await query.message.reply_text('Ошибка при получении списка объектов. Попробуйте снова.')

    elif data.startswith('endfrontobj_'):
        await query.message.delete()
        user_id = query.from_user.id
        object_id = int(query.data.split('_')[1])

        # Получаем все фронты со статусом "in_process"
        response = requests.get(f'{DJANGO_API_URL}fronttransfers/?status=in_process')
        if response.status_code == 200:
            fronts = response.json()
            filtered_fronts = [front for front in fronts if
                               front['sender_chat_id'] == str(user_id) and front['object_id'] == object_id
                               and front['is_finish'] == False]

            if not filtered_fronts:
                await query.message.reply_text("Нет фронтов для завершения на выбранном объекте.")
                return

            keyboard = []
            for front in filtered_fronts:
                object_name = front.get('object_name', 'неизвестно')
                work_type_name = front.get('work_type_name', 'неизвестно')
                block_section_name = front.get('block_section_name', 'неизвестно')

                if not object_name or object_name == 'неизвестно':
                    object_response = requests.get(f'{DJANGO_API_URL}objects/{front["object_id"]}/')
                    if object_response.status_code == 200:
                        object_name = object_response.json().get('name', 'неизвестно')

                if not work_type_name or work_type_name == 'неизвестно':
                    work_type_response = requests.get(f'{DJANGO_API_URL}worktypes/{front["work_type_id"]}/')
                    if work_type_response.status_code == 200:
                        work_type_name = work_type_response.json().get('name', 'неизвестно')

                if not block_section_name or block_section_name == 'неизвестно':
                    block_section_response = requests.get(f'{DJANGO_API_URL}blocksections/{front["block_section_id"]}/')
                    if block_section_response.status_code == 200:
                        block_section_name = block_section_response.json().get('name', 'неизвестно')

                button_text = f"{object_name} - {work_type_name} - {block_section_name} - Этаж {front['floor']}"
                callback_data = f"endfront_{front['id']}"
                keyboard.append([InlineKeyboardButton(button_text, callback_data=callback_data)])

            keyboard.append([InlineKeyboardButton("Назад", callback_data='endfront')])

            reply_markup = InlineKeyboardMarkup(keyboard)
            await query.message.reply_text("Выберите фронт для завершения работ:", reply_markup=reply_markup)
        else:
            await query.message.reply_text("Ошибка при получении списка фронтов. Попробуйте снова.")

    elif data.startswith('endfront_'):
        await query.message.delete()
        front_id = int(query.data.split('_')[1])

        # Получаем данные о выбранном фронте
        front_response = requests.get(f'{DJANGO_API_URL}fronttransfers/{front_id}/')
        if front_response.status_code == 200:
            front_data = front_response.json()

            # Обновляем статус и поле is_finish
            front_data['is_finish'] = True

            object_name = front_data.get('object_name', 'неизвестно')
            work_type_name = front_data.get('work_type_name', 'неизвестно')
            block_section_name = front_data.get('block_section_name', 'неизвестно')

            if not object_name or object_name == 'неизвестно':
                object_response = requests.get(f'{DJANGO_API_URL}objects/{front_data["object_id"]}/')
                if object_response.status_code == 200:
                    object_name = object_response.json().get('name', 'неизвестно')

            if not work_type_name or work_type_name == 'неизвестно':
                work_type_response = requests.get(f'{DJANGO_API_URL}worktypes/{front_data["work_type_id"]}/')
                if work_type_response.status_code == 200:
                    work_type_name = work_type_response.json().get('name', 'неизвестно')

            if not block_section_name or block_section_name == 'неизвестно':
                block_section_response = requests.get(f'{DJANGO_API_URL}blocksections/{front_data["block_section_id"]}/')
                if block_section_response.status_code == 200:
                    block_section_name = block_section_response.json().get('name', 'неизвестно')

            # Отправляем обновленные данные на сервер
            update_response = requests.put(f'{DJANGO_API_URL}fronttransfers/{front_id}/', json=front_data)

            asyncio.create_task(send_to_google_sheets(front_id, action='update'))
            if update_response.status_code == 200:
                await query.message.reply_text(
                    f"Фронт успешно завершен: {object_name} - {work_type_name} - {block_section_name} - Этаж {front_data['floor']}")
            else:
                await query.message.reply_text(f"Ошибка при завершении фронта. Попробуйте снова.")
        else:
            await query.message.reply_text("Ошибка при получении данных фронта. Попробуйте снова.")


    elif data.startswith('existing_front_'):
        front_id = int(data.split('_')[2])
        await handle_existing_front_selection(query, context, front_id)

    elif data == 'accept_fronts':
        await list_accept_fronts(update, context)


    elif data.startswith('accept_front_'):
        front_id = int(data.split('_')[2])
        await accept_front(query, context, front_id)

    elif data.startswith('accept_'):
        front_id = int(data.split('_')[1])
        await show_front_details(query, context, front_id)

    elif data == 'view_fronts':
        await view_fronts(update, context)

    elif data.startswith('front_'):
        front_id = int(data.split('_')[1])
        await view_front_details(query, context, front_id)

    elif data.startswith('rework_'):
        front_id = int(data.split('_')[1])
        await handle_rework(query, context, front_id)

    elif data.startswith('approve_'):
        front_id = int(data.split('_')[1])
        await approve_front(query, context, front_id)


    elif data.startswith('delete_error_'):
        front_id = int(data.split('_')[2])
        await handle_delete_error(query, context, front_id)

    elif data.startswith('decline_front_'):
        front_id = int(data.split('_')[2])
        await decline_front(query, context, front_id)

    elif data.startswith('transfer_'):
        if 'org_' in data:
            org_id = int(data.split('_')[2])
            await choose_transfer_user(query, context, org_id)

        elif 'user_' in data:
            user_id = int(data.split('_')[2])
            await choose_transfer_work_type(query, context, user_id)

        elif 'work_' in data:
            work_type_id = int(data.split('_')[2])
            await confirm_transfer(query, context, work_type_id)

        else:
            front_id = int(data.split('_')[1])
            await handle_transfer(query, context, front_id)

    elif data == 'back':
        await send_main_menu(query.message.chat.id, context, context.user_data['full_name'],
                             context.user_data['organization'])
    elif data == 'main_menu':
        await query.message.delete()
        await query.answer("Назад...")
        response = requests.get(f'{DJANGO_API_URL}users/chat/{user_id}/')
        if response.status_code == 200:
            user_data = response.json()
            full_name = user_data.get('full_name', 'Пользователь')
            organization_id = user_data.get('organization_id', None)
            await send_main_menu(query.message.chat.id, context, full_name, organization_id)

        else:
            await query.message.reply_text('Ошибка при получении данных пользователя. Попробуйте снова.')


    elif data.startswith('fronts_info_'):
        front_id = int(data.split('_')[2])
        await show_front_info(update, context, front_id)

    elif data == 'fronts_in_process':
        await view_fronts_in_process(update, context)

    elif data.startswith('object_'):
        await change_object_id(update, context)

    elif data == 'issue_front':

        await query.message.delete()
        # Удаление сообщения с основным меню
        # if 'main_menu_message_id' in context.user_data:
        #     await context.bot.delete_message(
        #         chat_id=query.message.chat.id,
        #         message_id=context.user_data['main_menu_message_id']
        #     )

        # Получаем данные пользователя
        response_user = requests.get(f'{DJANGO_API_URL}users/chat/{user_id}/')
        if response_user.status_code == 200:
            user_data = response_user.json()
            user_object_id = user_data.get('object_id')

            # Получаем объекты
            response_objects = requests.get(f'{DJANGO_API_URL}objects/')
            if response_objects.status_code == 200:
                objects = response_objects.json()

                if objects:
                    # Сортируем объекты по имени в алфавитном порядке
                    objects.sort(key=lambda obj: obj['name'])
                    # Создаем клавиатуру с кнопками в две колонки
                    keyboard = []
                    for i in range(0, len(objects), 2):
                        row = [
                            InlineKeyboardButton(objects[i]['name'], callback_data=f'issue_obj_{objects[i]["id"]}')
                        ]
                        if i + 1 < len(objects):
                            row.append(InlineKeyboardButton(objects[i + 1]['name'],
                                                            callback_data=f'issue_obj_{objects[i + 1]["id"]}'))
                        keyboard.append(row)


                    keyboard.append([InlineKeyboardButton("Назад", callback_data='frontbutton')])
                    reply_markup = InlineKeyboardMarkup(keyboard)
                    await query.message.reply_text('Выберите объект:', reply_markup=reply_markup)
                    context.user_data['stage'] = 'issue_choose_object'
                else:
                    await query.message.reply_text('Нет доступных объектов для данного пользователя.')

            else:
                await query.message.reply_text('Ошибка при получении списка объектов. Попробуйте снова.')

        else:
            await query.message.reply_text('Ошибка при получении данных пользователя. Попробуйте снова.')


    elif data.startswith('issue_obj_'):
        await query.message.delete()  # Удаление предыдущего сообщения
        object_id = int(data.split('_')[2])
        context.user_data['issue_object_id'] = object_id
        response = requests.get(f'{DJANGO_API_URL}organizations/')
        if response.status_code == 200:
            organizations = response.json()

            # Исключаем организацию с id = 3
            filtered_organizations = [org for org in organizations if org['organization'] != "БОС"]
            # Сортируем организации по алфавиту
            filtered_organizations.sort(key=lambda org: org['organization'])
            # Создание кнопок в две колонки
            keyboard = []
            for i in range(0, len(filtered_organizations), 2):
                row = [
                    InlineKeyboardButton(filtered_organizations[i]['organization'],
                                         callback_data=f'issue_org_{filtered_organizations[i]["id"]}')
                ]
                if i + 1 < len(filtered_organizations):
                    row.append(InlineKeyboardButton(filtered_organizations[i + 1]['organization'],
                                                    callback_data=f'issue_org_{filtered_organizations[i + 1]["id"]}'))
                keyboard.append(row)


            if filtered_organizations:
                keyboard.append([InlineKeyboardButton("Назад", callback_data='issue_front')])
                reply_markup = InlineKeyboardMarkup(keyboard)
                await query.message.reply_text('Выберите организацию:', reply_markup=reply_markup)
                context.user_data['stage'] = 'issue_choose_organization'
            else:
                await query.message.reply_text('Нет доступных организаций для выбранного объекта.')


        else:
            await query.message.reply_text('Ошибка при получении списка организаций. Попробуйте снова.')



    elif data.startswith('issue_org_'):
        await query.message.delete()  # Удаление предыдущего сообщения
        org_id = int(data.split('_')[2])
        context.user_data['issue_org_id'] = org_id
        response = requests.get(f'{DJANGO_API_URL}users/')

        if response.status_code == 200:
            users = response.json()
            filtered_users = [user for user in users if user['organization_id'] == org_id]

            if filtered_users:
                keyboard = [
                    [InlineKeyboardButton(user['full_name'], callback_data=f'issue_user_{user["chat_id"]}')] for user in
                    filtered_users
                ]
                keyboard.append([InlineKeyboardButton("Назад", callback_data='issue_front')])
                reply_markup = InlineKeyboardMarkup(keyboard)
                await query.message.reply_text('Выберите пользователя:', reply_markup=reply_markup)
                context.user_data['stage'] = 'issue_choose_user'

            else:
                await query.message.reply_text('В этой организации нет доступных пользователей.')
        else:
            await query.message.reply_text('Ошибка при получении списка пользователей. Попробуйте снова.')



    elif data.startswith('issue_user_'):
        await query.message.delete()  # Удаление предыдущего сообщения
        user_chat_id = int(data.split('_')[2])
        context.user_data['issue_user_chat_id'] = user_chat_id

        # Получение пересечения work_types_ids
        common_work_types_ids = await get_common_work_types(context.user_data['issue_object_id'],
                                                            context.user_data['issue_org_id'])

        if common_work_types_ids:
            ids_query = "&".join([f"ids={id}" for id in common_work_types_ids])
            async with aiohttp.ClientSession() as session:
                async with session.get(f'{DJANGO_API_URL}worktypes/?{ids_query}') as response:
                    if response.status == 200:
                        work_types = await response.json()
                        keyboard = [
                            [InlineKeyboardButton(work['name'], callback_data=f'issue_work_{work["id"]}')] for work in
                            work_types
                        ]
                        keyboard.append([InlineKeyboardButton("Назад", callback_data='issue_front')])
                        reply_markup = InlineKeyboardMarkup(keyboard)
                        await query.message.reply_text('Выберите вид работ:', reply_markup=reply_markup)
                        context.user_data['stage'] = 'issue_choose_work_type'
                    else:
                        await query.message.reply_text('Ошибка при получении списка видов работ. Попробуйте снова.')
        else:
            await query.message.reply_text('Нет доступных видов работ для выбранного объекта и организации.')


    elif data.startswith('issue_work_'):
        await query.message.delete()  # Удаление предыдущего сообщения
        work_type_id = int(data.split('_')[2])
        context.user_data['issue_work_type_id'] = work_type_id
        response = requests.get(f'{DJANGO_API_URL}objects/{context.user_data["issue_object_id"]}/blocksections/')

        if response.status_code == 200:
            block_sections = response.json()
            keyboard = [
                [InlineKeyboardButton(block['name'], callback_data=f'issue_block_{block["id"]}')] for block in
                block_sections
            ]
            keyboard.append([InlineKeyboardButton("Назад", callback_data='issue_front')])
            reply_markup = InlineKeyboardMarkup(keyboard)
            await query.message.reply_text('Выберите блок или секцию:', reply_markup=reply_markup)
            context.user_data['stage'] = 'issue_choose_block_section'

        else:
            await query.message.reply_text('Ошибка при получении списка блоков или секций. Попробуйте снова.')

    elif data.startswith('issue_block_'):
        await query.message.delete()  # Удаление предыдущего сообщения
        block_section_id = int(data.split('_')[2])
        context.user_data['issue_block_section_id'] = block_section_id
        response = requests.get(f'{DJANGO_API_URL}blocksections/{block_section_id}/')

        if response.status_code == 200:
            block_section = response.json()
            number_of_floors_bottom = block_section['number_of_floors_bottom']
            number_of_floors = block_section['number_of_floors']

            # Генерация кнопок этажей в две колонки, исключая 0
            keyboard = []
            for i in range(number_of_floors_bottom, number_of_floors + 1):
                if i == 0:
                    continue
                if len(keyboard) == 0 or len(keyboard[-1]) == 2:
                    keyboard.append([InlineKeyboardButton(f'{i} этаж', callback_data=f'issue_floor_{i}')])
                else:
                    keyboard[-1].append(InlineKeyboardButton(f'{i} этаж', callback_data=f'issue_floor_{i}'))

            # Добавление кнопки кровли
            keyboard.append([InlineKeyboardButton('Кровля', callback_data='issue_floor_roof')])

            # Добавление кнопки "Назад"
            keyboard.append([InlineKeyboardButton("Назад", callback_data='issue_front')])
            reply_markup = InlineKeyboardMarkup(keyboard)
            await query.message.reply_text('Выберите этаж:', reply_markup=reply_markup)
            context.user_data['stage'] = 'issue_choose_floor'

        else:
            await query.message.reply_text('Ошибка при получении информации о блоке или секции. Попробуйте снова.')



    elif data.startswith('issue_floor_'):
        await query.message.delete()  # Удаление предыдущего сообщения
        floor = data.split('_')[2]
        floor_number = floor if floor != 'roof' else 'Кровля'
        context.user_data['issue_floor'] = floor_number

        # Получаем все данные до подтверждения
        user_chat_id = context.user_data['issue_user_chat_id']
        user_response = requests.get(f'{DJANGO_API_URL}users/chat/{user_chat_id}/')
        object_response = requests.get(f'{DJANGO_API_URL}objects/{context.user_data["issue_object_id"]}/')
        work_type_response = requests.get(f'{DJANGO_API_URL}worktypes/{context.user_data["issue_work_type_id"]}/')
        block_section_response = requests.get(
            f'{DJANGO_API_URL}blocksections/{context.user_data["issue_block_section_id"]}/')

        if user_response.status_code == 200 and object_response.status_code == 200 and work_type_response.status_code == 200 and block_section_response.status_code == 200:
            user_data = user_response.json()
            context.user_data['issue_user_id'] = user_data['id']
            context.user_data['issue_object_name'] = object_response.json()['name']
            context.user_data['issue_work_type_name'] = work_type_response.json()['name']
            context.user_data['issue_block_section_name'] = block_section_response.json()['name']

            transfer_info = (
                f"Объект: {context.user_data['issue_object_name']}\n"
                f"Вид работ: {context.user_data['issue_work_type_name']}\n"
                f"Блок/Секция: {context.user_data['issue_block_section_name']}\n"
                f"Этаж: {floor_number}\n"
            )

            keyboard = [
                [InlineKeyboardButton("\U00002705 Да", callback_data='issue_confirm_yes')],
                [InlineKeyboardButton("\U0000274C Нет", callback_data='issue_confirm_no')]
            ]
            reply_markup = InlineKeyboardMarkup(keyboard)
            await query.message.reply_text(f'Подтвердите данные:\n{transfer_info}', reply_markup=reply_markup)

        else:
            await query.message.reply_text('Ошибка при получении данных. Попробуйте снова.')


    elif data == 'issue_confirm_yes':
        await query.edit_message_reply_markup(reply_markup=None)

        # Используем данные из контекста, чтобы избежать повторных запросов
        user_id = context.user_data['issue_user_id']
        user_chat_id = str(context.user_data['issue_user_chat_id'])  # Преобразуем в строку
        object_name = context.user_data['issue_object_name']
        work_type_name = context.user_data['issue_work_type_name']
        block_section_name = context.user_data['issue_block_section_name']
        floor_number = context.user_data['issue_floor']

        boss_id = requests.get(f'{DJANGO_API_URL}users/chat/{query.from_user.id}/').json()['id']


        transfer_data = {
            'sender_id': user_id,
            'sender_chat_id': user_chat_id,
            'object_id': context.user_data['issue_object_id'],
            'work_type_id': context.user_data['issue_work_type_id'],
            'block_section_id': context.user_data['issue_block_section_id'],
            'floor': context.user_data['issue_floor'],
            'status': 'on_consideration',
            'created_at': datetime.now().isoformat(),
            'approval_at': datetime.now().isoformat(),
            'receiver_id': user_id,
            'boss_id': boss_id,
            'photos': [],
            'is_finish': False

        }

        # Добавим отладочную информацию для проверки данных

        logger.info(f"Отправка данных в API для создания нового фронта: {json.dumps(transfer_data, indent=2)}")
        response = requests.post(f'{DJANGO_API_URL}fronttransfers/', json=transfer_data)

        # Проверим ответ от API
        logger.info(f"Ответ от API при создании нового фронта: {response.status_code}, {response.text}")

        if response.status_code == 200:
            transfer = response.json()
            await query.message.reply_text('Фронт успешно выдан.')
            formatted_datetime = datetime.now().strftime('%d-%m-%Y %H:%M:%S')

            # Уведомление пользователя
            message_text = (

                f"Вам был выдан фронт работ:\n"
                f"*Объект:* {object_name}\n"
                f"*Вид работ:* {work_type_name}\n"
                f"*Блок/Секция:* {block_section_name}\n"
                f"*Этаж:* {floor_number}\n"
                f"*Дата выдачи (МСК):* {formatted_datetime}\n"

            )

            keyboard = [
                [InlineKeyboardButton("Список фронт работ", callback_data='accept_fronts')],
            ]

            reply_markup = InlineKeyboardMarkup(keyboard)

            try:
                # Попробуем отправить сообщение и логируем результат
                logger.info(f"Отправка сообщения пользователю с chat_id: {user_chat_id}")
                await context.bot.send_message(
                    chat_id=user_chat_id,
                    text=message_text,
                    parse_mode=ParseMode.MARKDOWN,
                    reply_markup=reply_markup

                )

                logger.info("Сообщение успешно отправлено пользователю")

            except Exception as e:

                # Логируем ошибку, если отправка сообщения не удалась
                logger.error(f"Ошибка при отправке сообщения пользователю: {e}")
                await query.message.reply_text(f"Ошибка при отправке сообщения пользователю: {e}")

            # Возвращение в главное меню

            user_id = query.from_user.id
            response = requests.get(f'{DJANGO_API_URL}users/chat/{user_id}/')

            if response.status_code == 200:
                user_data = response.json()
                full_name = user_data.get('full_name', 'Пользователь')
                organization_id = user_data.get('organization_id', None)
                await send_main_menu(user_id, context, full_name, organization_id)

            else:

                await query.message.reply_text('Ошибка при получении данных пользователя. Попробуйте снова.')


        else:

            await query.message.reply_text(f'Ошибка при создании фронта: {response.text}')


    elif data == 'issue_confirm_no':
        await query.message.delete()  # Удаление предыдущего сообщения
        await query.message.reply_text('Выдача фронта отменена.')
        user_id = query.from_user.id
        response = requests.get(f'{DJANGO_API_URL}users/chat/{user_id}/')

        if response.status_code == 200:
            user_data = response.json()
            full_name = user_data.get('full_name', 'Пользователь')
            organization_id = user_data.get('organization_id', None)
            await send_main_menu(user_id, context, full_name, organization_id)

        else:
            await query.message.reply_text('Ошибка при получении данных пользователя. Попробуйте снова.')

    elif data == 'workforce_transfer':
        await handle_transfer_workforce(query, context)



    elif data.startswith('workforce_obj_'):
        await query.message.delete()
        object_id = int(data.split('_')[2])
        context.user_data['workforce_object_id'] = object_id
        common_work_types_ids = await get_common_work_types(object_id, context.user_data['organization_id'])
        if common_work_types_ids:
            ids_query = "&".join([f"ids={id}" for id in common_work_types_ids])
            response = requests.get(f'{DJANGO_API_URL}worktypes/?{ids_query}')
            if response.status_code == 200:
                work_types = response.json()
                keyboard = [
                    [InlineKeyboardButton(work['name'], callback_data=f'workforce_work_{work["id"]}')] for work in
                    work_types
                ]
                keyboard.append([InlineKeyboardButton("Назад", callback_data='workforce_menu')])
                reply_markup = InlineKeyboardMarkup(keyboard)
                await query.message.reply_text('Выберите вид работ:', reply_markup=reply_markup)
            else:
                await query.message.reply_text('Ошибка при получении списка видов работ. Попробуйте снова.')
        else:
            await query.message.reply_text('Нет доступных видов работ для выбранного объекта и организации.')


    elif data.startswith('workforce_block_'):
        await query.message.delete()
        block_section_id = int(data.split('_')[2])
        context.user_data['workforce_block_section_id'] = block_section_id
        response = requests.get(f'{DJANGO_API_URL}blocksections/{block_section_id}/')
        if response.status_code == 200:
            block_section = response.json()
            number_of_floors_bottom = block_section['number_of_floors_bottom']
            number_of_floors = block_section['number_of_floors']

            # Генерация кнопок этажей в две колонки, исключая 0
            keyboard = []
            for i in range(number_of_floors_bottom, number_of_floors + 1):
                if i == 0:
                    continue
                if len(keyboard) == 0 or len(keyboard[-1]) == 2:
                    keyboard.append([InlineKeyboardButton(f'{i} этаж', callback_data=f'workforce_floor_{i}')])
                else:
                    keyboard[-1].append(InlineKeyboardButton(f'{i} этаж', callback_data=f'workforce_floor_{i}'))

            # Добавление кнопки кровли и кнопки "Назад"
            keyboard.append([InlineKeyboardButton('Кровля', callback_data='workforce_floor_roof')])
            keyboard.append([InlineKeyboardButton("Назад", callback_data='workforce_menu')])

            reply_markup = InlineKeyboardMarkup(keyboard)
            await query.message.reply_text('Выберите этаж:', reply_markup=reply_markup)

        else:
            await query.message.reply_text('Ошибка при получении информации о блоке или секции. Попробуйте снова.')

    elif data.startswith('workforce_floor_'):
        await query.message.delete()
        floor = data.split('_')[2]
        context.user_data['workforce_floor'] = floor

        keyboard = []
        keyboard.append([InlineKeyboardButton("\U0000274C Отмена", callback_data='main_menu')])
        reply_markup = InlineKeyboardMarkup(keyboard)

        await query.message.reply_text('Введите численность:', reply_markup=reply_markup)
        context.user_data['expecting_workforce_count'] = True


    elif data.startswith('workforce_work_'):
        await query.message.delete()
        work_type_id = int(data.split('_')[2])
        context.user_data['workforce_work_type_id'] = work_type_id
        work_type_response = requests.get(f'{DJANGO_API_URL}worktypes/{work_type_id}/')
        if work_type_response.status_code == 200:
            work_type_name = work_type_response.json()['name']
            context.user_data['workforce_work_type_name'] = work_type_name
            response = requests.get(
                f'{DJANGO_API_URL}objects/{context.user_data["workforce_object_id"]}/blocksections/')

            if response.status_code == 200:
                block_sections = response.json()
                keyboard = [
                    [InlineKeyboardButton(block['name'], callback_data=f'workforce_block_{block["id"]}')] for block in
                    block_sections
                ]
                keyboard.append([InlineKeyboardButton("Назад", callback_data='workforce_menu')])
                reply_markup = InlineKeyboardMarkup(keyboard)
                await query.message.reply_text('Выберите блок или секцию:', reply_markup=reply_markup)
            else:
                await query.message.reply_text('Ошибка при получении списка блоков или секций. Попробуйте снова.')
        else:
            await query.message.reply_text('Ошибка при получении данных вида работ. Попробуйте снова.')


    elif data.startswith('delete_workforce_'):
        await query.message.delete()
        workforce_id = int(data.split('_')[2])
        context.user_data['workforce_id_to_delete'] = workforce_id
        await query.message.reply_text('Вы действительно хотите удалить эту запись?',
                                       reply_markup=InlineKeyboardMarkup([
                                           [InlineKeyboardButton("\U00002705 Да", callback_data='confirm_delete_workforce')],
                                           [InlineKeyboardButton("\U0000274C Нет", callback_data='workforce_menu')]
                                       ]))


    elif data == 'confirm_delete_workforce':
        await query.message.delete()
        await handle_delete_workforce(query, context)

    elif data == 'repeat_workforce':
        await query.edit_message_reply_markup(reply_markup=None)
        response = requests.get(f'{DJANGO_API_URL}objects/{context.user_data["workforce_object_id"]}/blocksections/')
        if response.status_code == 200:
            block_sections = response.json()
            keyboard = [
                [InlineKeyboardButton(block['name'], callback_data=f'workforce_block_{block["id"]}')] for block in
                block_sections
            ]
            reply_markup = InlineKeyboardMarkup(keyboard)
            await query.message.reply_text('Выберите блок или секцию:', reply_markup=reply_markup)
        else:
            await query.message.reply_text('Ошибка при получении списка блоков или секций. Попробуйте снова.')

    elif data == 'workforce_refactor':
        await query.message.delete()
        today = datetime.now().date().isoformat()
        response = requests.get(f'{DJANGO_API_URL}frontworkforces/')

        if response.status_code == 200:
            workforces = response.json()
            user_data = requests.get(f'{DJANGO_API_URL}users/chat/{user_id}/').json()
            user_id = user_data.get('id')

            # Фильтруем записи только за сегодняшний день и для текущего объекта
            today_workforces = [
                wf for wf in workforces
                if parser.parse(wf['date']).date().isoformat() == today and wf['user_id'] == user_id
            ]

            if today_workforces:
                buttons = []
                for wf in today_workforces:
                    work_type_name = requests.get(f'{DJANGO_API_URL}worktypes/{wf["work_type_id"]}/').json()['name']
                    block_section_name = \
                    requests.get(f'{DJANGO_API_URL}blocksections/{wf["block_section_id"]}/').json()['name']
                    button_text = f"{wf['workforce_count']} - {work_type_name} - {block_section_name} - Этаж {wf['floor']}"
                    buttons.append([InlineKeyboardButton(button_text, callback_data=f'refactor_{wf["id"]}')])
                buttons.append([InlineKeyboardButton("Назад", callback_data='workforce_menu')])
                reply_markup = InlineKeyboardMarkup(buttons)
                await query.message.reply_text('Выберите численность для редактирования:', reply_markup=reply_markup)
            else:
                await query.message.reply_text('Сегодня нет численностей для редактирования.')

                user_id = update.effective_user.id
                user_response = requests.get(f'{DJANGO_API_URL}users/chat/{user_id}')
                if user_response.status_code == 200:
                    user_data = user_response.json()
                    await send_main_menu(query.message.chat.id, context, user_data['full_name'],
                                         user_data['organization_id'])
                else:
                    await query.message.reply_text('Ошибка при получении данных пользователя.')

        else:
            await query.message.reply_text('Ошибка при получении списка численностей. Попробуйте снова.')

    elif data == 'workforce_delete':
        await query.message.delete()
        today = datetime.now().date().isoformat()
        response = requests.get(f'{DJANGO_API_URL}frontworkforces/')

        if response.status_code == 200:
            workforces = response.json()
            user_data = requests.get(f'{DJANGO_API_URL}users/chat/{user_id}/').json()
            user_id = user_data.get('id')

            # Фильтруем записи только за сегодняшний день и для текущего объекта
            today_workforces = [
                wf for wf in workforces
                if parser.parse(wf['date']).date().isoformat() == today and wf['user_id'] == user_id
            ]
            if today_workforces:
                buttons = []
                for wf in today_workforces:
                    work_type_name = requests.get(f'{DJANGO_API_URL}worktypes/{wf["work_type_id"]}/').json()['name']
                    block_section_name = \
                        requests.get(f'{DJANGO_API_URL}blocksections/{wf["block_section_id"]}/').json()['name']
                    button_text = f"{wf['workforce_count']} - {work_type_name} - {block_section_name} - Этаж {wf['floor']}"
                    buttons.append([InlineKeyboardButton(button_text, callback_data=f'delete_workforce_{wf["id"]}')])
                buttons.append([InlineKeyboardButton("Назад", callback_data='workforce_menu')])
                reply_markup = InlineKeyboardMarkup(buttons)
                await query.message.reply_text('Выберите запись численности для удаления:', reply_markup=reply_markup)
            else:
                await query.message.reply_text('Сегодня не было передано численностей.')

                user_id = update.effective_user.id
                user_response = requests.get(f'{DJANGO_API_URL}users/chat/{user_id}')
                if user_response.status_code == 200:
                    user_data = user_response.json()
                    await send_main_menu(query.message.chat.id, context, user_data['full_name'],
                                         user_data['organization_id'])
                else:
                    await query.message.reply_text('Ошибка при получении данных пользователя.')

        else:
            await query.message.reply_text('Ошибка при получении данных о численности.')


    elif data.startswith('refactor_'):
        await query.message.delete()
        workforce_id = int(data.split('_')[1])
        context.user_data['workforce_id_to_refactor'] = workforce_id
        keyboard = []
        keyboard.append([InlineKeyboardButton("\U0000274C Отмена", callback_data='main_menu')])
        reply_markup = InlineKeyboardMarkup(keyboard)
        message = await query.message.reply_text('Введите новую численность:', reply_markup=reply_markup)
        context.user_data['expecting_new_workforce_count'] = True
        context.user_data['refactor_message_id'] = message.message_id


    elif data == 'view_workforce':
        await query.message.delete()
        keyboard = [
            [InlineKeyboardButton("\U0001F4E2 Оповестить", callback_data='notify_organizations')],
            [InlineKeyboardButton("\U0001F4C6 Сегодня", callback_data='view_today_workforce')],
            [InlineKeyboardButton("\U0001F50D Определенный день", callback_data='choose_month')],
            [InlineKeyboardButton("Назад", callback_data='main_menu')],
        ]
        reply_markup = InlineKeyboardMarkup(keyboard)
        await query.message.reply_text('Выберите действие:', reply_markup=reply_markup)

    elif data == 'notify_organizations':
        await query.message.delete()
        # Получаем данные пользователя
        response = requests.get(f'{DJANGO_API_URL}users/chat/{user_id}/')
        if response.status_code == 200:
            user_data = response.json()
            user_object_id = user_data.get('object_id')

            if user_object_id is None:
                await query.message.reply_text("У вас нет назначенного объекта.")
                return

            # Получаем список организаций, у которых в object_ids есть наш object_id
            response = requests.get(f'{DJANGO_API_URL}organizations/')
            if response.status_code == 200:
                organizations = response.json()
                relevant_organizations = [
                    org for org in organizations if org.get('object_ids') and user_object_id in org['object_ids']
                ]

                if not relevant_organizations:
                    await query.message.reply_text("Нет организаций для уведомления.")
                    return

                organization_ids = [org['id'] for org in relevant_organizations]

                # Получаем список пользователей, у которых organization_id есть в списке relevant_organizations
                response = requests.get(f'{DJANGO_API_URL}users/')
                if response.status_code == 200:
                    users = response.json()
                    relevant_users = [
                        user for user in users if user['organization_id'] in organization_ids
                    ]

                    if not relevant_users:
                        await query.message.reply_text("Нет пользователей для уведомления.")
                        return

                    chat_ids = [user['chat_id'] for user in relevant_users]

                    # Отправка уведомления
                    notification_text = "\U00002757 **Генеральный подрядчик напоминает вам внести численность для текущего фронта работ.**"
                    for chat_id in chat_ids:
                        try:
                            await context.bot.send_message(
                                chat_id=chat_id,
                                text=notification_text,
                                parse_mode=ParseMode.MARKDOWN
                            )
                        except Exception as e:
                            logger.error(f"Ошибка при отправке сообщения пользователю {chat_id}: {e}")

                    await query.message.reply_text("Уведомления успешно отправлены.")

                    # Возвращение в главное меню
                    full_name = user_data.get('full_name', 'Пользователь')
                    organization_id = user_data.get('organization_id', None)
                    await send_main_menu(user_id, context, full_name, organization_id)

                else:
                    await query.message.reply_text("Ошибка при получении списка пользователей.")
            else:
                await query.message.reply_text("Ошибка при получении списка организаций.")
        else:
            await query.message.reply_text("Ошибка при получении данных пользователя.")


    elif data == 'view_today_workforce':
        await query.message.delete()
        await view_today_workforce(query, context)


    elif data == 'choose_month':
        await query.message.delete()

        # Разбиваем месяцы на три строки по четыре месяца в строке
        months = [
            ["Январь", "Февраль", "Март", "Апрель"],
            ["Май", "Июнь", "Июль", "Август"],
            ["Сентябрь", "Октябрь", "Ноябрь", "Декабрь"]
        ]

        keyboard = [
            [InlineKeyboardButton(month, callback_data=f'month_{i * 4 + j + 1}') for j, month in enumerate(row)]
            for i, row in enumerate(months)
        ]
        keyboard.append([InlineKeyboardButton("Назад", callback_data='view_workforce')])
        reply_markup = InlineKeyboardMarkup(keyboard)
        await query.message.reply_text('Выберите месяц:', reply_markup=reply_markup)


    elif data.startswith('month_'):
        await query.message.delete()
        month = int(data.split('_')[1])
        context.user_data['selected_month'] = month
        days_in_month = (datetime(2024, month % 12 + 1, 1) - timedelta(days=1)).day

        # Разбиваем дни на строки по 7 дней в каждой строке
        keyboard = [
            [InlineKeyboardButton(str(day), callback_data=f'day_{day}') for day in
             range(i, min(i + 7, days_in_month + 1))]
            for i in range(1, days_in_month + 1, 7)
        ]
        keyboard.append([InlineKeyboardButton("Назад", callback_data='view_workforce')])
        reply_markup = InlineKeyboardMarkup(keyboard)

        await query.message.reply_text('Выберите день:', reply_markup=reply_markup)

    elif data.startswith('day_'):
        await query.message.delete()
        day = int(data.split('_')[1])
        month = context.user_data['selected_month']
        await view_specific_day_workforce(query, context, day, month)

    elif data == 'volume_menu':
        await query.message.delete()
        keyboard = [
            [InlineKeyboardButton("📐 Передать объем", callback_data='volume_transfer')],
            [InlineKeyboardButton("\U0000270F Редактировать объем", callback_data='volume_refactor')],
            [InlineKeyboardButton("\U0000274C Удалить запись", callback_data='volume_delete')],
            [InlineKeyboardButton("Назад", callback_data='main_menu')]
        ]
        reply_markup = InlineKeyboardMarkup(keyboard)
        await query.message.reply_text('Выберите действие для объема:', reply_markup=reply_markup)

    elif data == 'volume_transfer':
        await handle_transfer_volume(query, context)

    elif data.startswith('volume_obj_'):
        await query.message.delete()
        object_id = int(data.split('_')[2])
        context.user_data['volume_object_id'] = object_id
        common_work_types_ids = await get_common_work_types(object_id, context.user_data['organization_id'])
        if common_work_types_ids:
            ids_query = "&".join([f"ids={id}" for id in common_work_types_ids])
            response = requests.get(f'{DJANGO_API_URL}worktypes/?{ids_query}')
            if response.status_code == 200:
                work_types = response.json()
                keyboard = [
                    [InlineKeyboardButton(work['name'], callback_data=f'volume_work_{work["id"]}')] for work in
                    work_types
                ]
                keyboard.append([InlineKeyboardButton("Назад", callback_data='volume_transfer')])
                reply_markup = InlineKeyboardMarkup(keyboard)
                await query.message.reply_text('Выберите вид работ:', reply_markup=reply_markup)
            else:
                await query.message.reply_text('Ошибка при получении списка видов работ. Попробуйте снова.')
        else:
            await query.message.reply_text('Нет доступных видов работ для выбранного объекта и организации.')

    elif data.startswith('volume_block_'):
        await query.message.delete()
        block_section_id = int(data.split('_')[2])
        context.user_data['volume_block_section_id'] = block_section_id
        response = requests.get(f'{DJANGO_API_URL}blocksections/{block_section_id}/')
        if response.status_code == 200:
            block_section = response.json()
            number_of_floors_bottom = block_section['number_of_floors_bottom']
            number_of_floors = block_section['number_of_floors']

            # Генерация кнопок этажей в две колонки, исключая 0
            keyboard = []
            for i in range(number_of_floors_bottom, number_of_floors + 1):
                if i == 0:
                    continue
                if len(keyboard) == 0 or len(keyboard[-1]) == 2:
                    keyboard.append([InlineKeyboardButton(f'{i} этаж', callback_data=f'volume_floor_{i}')])
                else:
                    keyboard[-1].append(InlineKeyboardButton(f'{i} этаж', callback_data=f'volume_floor_{i}'))

            # Добавление кнопки кровли и кнопки "Назад"
            keyboard.append([InlineKeyboardButton('Кровля', callback_data='volume_floor_roof')])
            keyboard.append([InlineKeyboardButton("Назад", callback_data='volume_transfer')])

            reply_markup = InlineKeyboardMarkup(keyboard)
            await query.message.reply_text('Выберите этаж:', reply_markup=reply_markup)
        else:
            await query.message.reply_text('Ошибка при получении информации о блоке или секции. Попробуйте снова.')

    elif data.startswith('volume_floor_'):
        await query.message.delete()
        floor = data.split('_')[2]
        context.user_data['volume_floor'] = floor

        keyboard = []
        keyboard.append([InlineKeyboardButton("\U0000274C Отмена", callback_data='main_menu')])
        reply_markup = InlineKeyboardMarkup(keyboard)

        await query.message.reply_text('Введите объем в м³:', reply_markup=reply_markup)
        context.user_data['expecting_volume_count'] = True

    elif data.startswith('volume_work_'):
        await query.message.delete()
        work_type_id = int(data.split('_')[2])
        context.user_data['volume_work_type_id'] = work_type_id
        work_type_response = requests.get(f'{DJANGO_API_URL}worktypes/{work_type_id}/')
        if work_type_response.status_code == 200:
            work_type_name = work_type_response.json()['name']
            context.user_data['volume_work_type_name'] = work_type_name
            response = requests.get(
                f'{DJANGO_API_URL}objects/{context.user_data["volume_object_id"]}/blocksections/')

            if response.status_code == 200:
                block_sections = response.json()
                keyboard = [
                    [InlineKeyboardButton(block['name'], callback_data=f'volume_block_{block["id"]}')] for block in
                    block_sections
                ]
                keyboard.append([InlineKeyboardButton("Назад", callback_data='volume_transfer')])
                reply_markup = InlineKeyboardMarkup(keyboard)
                await query.message.reply_text('Выберите блок или секцию:', reply_markup=reply_markup)
            else:
                await query.message.reply_text('Ошибка при получении списка блоков или секций. Попробуйте снова.')
        else:
            await query.message.reply_text('Ошибка при получении данных вида работ. Попробуйте снова.')

    elif data.startswith('delete_volume_'):
        await query.message.delete()
        volume_id = int(data.split('_')[2])
        context.user_data['volume_id_to_delete'] = volume_id
        await query.message.reply_text('Вы действительно хотите удалить эту запись?',
                                       reply_markup=InlineKeyboardMarkup([
                                           [InlineKeyboardButton("\U00002705 Да",
                                                                 callback_data='confirm_delete_volume')],
                                           [InlineKeyboardButton("\U0000274C Нет", callback_data='volume_menu')]
                                       ]))

    elif data == 'confirm_delete_volume':
        await query.message.delete()
        await handle_delete_volume(query, context)

    elif data == 'repeat_volume':
        await query.edit_message_reply_markup(reply_markup=None)
        response = requests.get(f'{DJANGO_API_URL}objects/{context.user_data["volume_object_id"]}/blocksections/')
        if response.status_code == 200:
            block_sections = response.json()
            keyboard = [
                [InlineKeyboardButton(block['name'], callback_data=f'volume_block_{block["id"]}')] for block in
                block_sections
            ]
            reply_markup = InlineKeyboardMarkup(keyboard)
            await query.message.reply_text('Выберите блок или секцию:', reply_markup=reply_markup)
        else:
            await query.message.reply_text('Ошибка при получении списка блоков или секций. Попробуйте снова.')

    elif data == 'volume_refactor':
        await query.message.delete()
        today = datetime.now().date().isoformat()
        response = requests.get(f'{DJANGO_API_URL}volumes/')

        if response.status_code == 200:
            volumes = response.json()
            user_data = requests.get(f'{DJANGO_API_URL}users/chat/{user_id}/').json()
            user_id = user_data.get('id')

            # Фильтруем записи только за сегодняшний день и для текущего объекта
            today_volumes = [
                vol for vol in volumes
                if parser.parse(vol['date']).date().isoformat() == today and vol['user_id'] == user_id
            ]

            if today_volumes:
                buttons = []
                for vol in today_volumes:
                    work_type_name = requests.get(f'{DJANGO_API_URL}worktypes/{vol["work_type_id"]}/').json()['name']
                    block_section_name = \
                        requests.get(f'{DJANGO_API_URL}blocksections/{vol["block_section_id"]}/').json()['name']
                    button_text = f"{vol['volume']} - {work_type_name} - {block_section_name} - Этаж {vol['floor']}"
                    buttons.append([InlineKeyboardButton(button_text, callback_data=f'refactorvolume_{vol["id"]}')])

                buttons.append([InlineKeyboardButton("Назад", callback_data='volume_menu')])

                reply_markup = InlineKeyboardMarkup(buttons)
                await query.message.reply_text('Выберите объем для редактирования:', reply_markup=reply_markup)
            else:
                await query.message.reply_text('Сегодня нет объемов для редактирования.')
        else:
            await query.message.reply_text('Ошибка при получении списка объемов. Попробуйте снова.')

    elif data == 'volume_delete':
        await query.message.delete()
        today = datetime.now().date().isoformat()
        response = requests.get(f'{DJANGO_API_URL}volumes/')

        if response.status_code == 200:
            volumes = response.json()
            user_data = requests.get(f'{DJANGO_API_URL}users/chat/{user_id}/').json()
            user_id = user_data.get('id')

            # Фильтруем записи только за сегодняшний день и для текущего объекта
            today_volumes = [
                vol for vol in volumes
                if parser.parse(vol['date']).date().isoformat() == today and vol['user_id'] == user_id
            ]
            if today_volumes:
                buttons = []
                for vol in today_volumes:
                    work_type_name = requests.get(f'{DJANGO_API_URL}worktypes/{vol["work_type_id"]}/').json()['name']
                    block_section_name = \
                        requests.get(f'{DJANGO_API_URL}blocksections/{vol["block_section_id"]}/').json()['name']
                    floor_text = f"— Этаж {vol['floor']}" if vol['floor'] and vol['floor'] != "None" else ""
                    button_text = f"{vol['volume']} - {work_type_name} - {block_section_name} {floor_text}"
                    buttons.append([InlineKeyboardButton(button_text, callback_data=f'delete_volume_{vol["id"]}')])

                buttons.append([InlineKeyboardButton("Назад", callback_data='volume_menu')])

                reply_markup = InlineKeyboardMarkup(buttons)
                await query.message.reply_text('Выберите запись объема для удаления:', reply_markup=reply_markup)
            else:
                await query.message.reply_text('Сегодня не было передано объемов.')
                user_id = update.effective_user.id
                user_response = requests.get(f'{DJANGO_API_URL}users/chat/{user_id}')
                if user_response.status_code == 200:
                    user_data = user_response.json()
                    await send_main_menu(query.message.chat.id, context, user_data['full_name'],
                                         user_data['organization_id'])
                else:
                    await query.message.reply_text('Ошибка при получении данных пользователя.')

        else:
            await query.message.reply_text('Ошибка при получении данных об объеме.')

    elif data.startswith('refactorvolume_'):
        await query.message.delete()
        volume_id = int(data.split('_')[1])
        context.user_data['volume_id_to_refactor'] = volume_id

        keyboard = []
        keyboard.append([InlineKeyboardButton("\U0000274C Отмена", callback_data='main_menu')])
        reply_markup = InlineKeyboardMarkup(keyboard)

        message = await query.message.reply_text('Введите новый объем:', reply_markup=reply_markup)
        context.user_data['expecting_new_volume_count'] = True
        context.user_data['refactor_message_id'] = message.message_id

    elif data == 'view_volume':
        await query.message.delete()
        keyboard = [
            [InlineKeyboardButton("\U0001F4C6 Сегодня", callback_data='view_today_volume')],
            [InlineKeyboardButton("\U0001F50D Определенный день", callback_data='choose_volumemonth')],
            [InlineKeyboardButton("Назад", callback_data='main_menu')],
        ]
        reply_markup = InlineKeyboardMarkup(keyboard)
        await query.message.reply_text('Выберите действие для объемов:', reply_markup=reply_markup)

    elif data == 'view_today_volume':
        await query.message.delete()
        await view_today_volume(query, context)

    elif data == 'choose_volumemonth':
        await query.message.delete()

        # Разбиваем месяцы на три строки по четыре месяца в строке
        months = [
            ["Январь", "Февраль", "Март", "Апрель"],
            ["Май", "Июнь", "Июль", "Август"],
            ["Сентябрь", "Октябрь", "Ноябрь", "Декабрь"]
        ]

        keyboard = [
            [InlineKeyboardButton(month, callback_data=f'volumemonth_{i * 4 + j + 1}') for j, month in enumerate(row)]
            for i, row in enumerate(months)
        ]
        keyboard.append([InlineKeyboardButton("Назад", callback_data='view_volume')])
        reply_markup = InlineKeyboardMarkup(keyboard)
        await query.message.reply_text('Выберите месяц:', reply_markup=reply_markup)

    elif data.startswith('volumemonth_'):
        await query.message.delete()
        month = int(data.split('_')[1])
        context.user_data['selected_volumemonth'] = month
        days_in_month = (datetime(2024, month % 12 + 1, 1) - timedelta(days=1)).day

        # Разбиваем дни на строки по 7 дней в каждой строке
        keyboard = [
            [InlineKeyboardButton(str(day), callback_data=f'volumeday_{day}') for day in
             range(i, min(i + 7, days_in_month + 1))]
            for i in range(1, days_in_month + 1, 7)
        ]
        keyboard.append([InlineKeyboardButton("Назад", callback_data='choose_volumemonth')])
        reply_markup = InlineKeyboardMarkup(keyboard)

        await query.message.reply_text('Выберите день:', reply_markup=reply_markup)

    elif data.startswith('volumeday_'):
        await query.message.delete()
        day = int(data.split('_')[1])
        month = context.user_data['selected_volumemonth']
        await view_specific_day_volume(query, context, day, month)

    elif data == 'changeobject':
        await query.message.delete()
        response = requests.get(f'{DJANGO_API_URL}objects/')
        if response.status_code == 200:
            objects = response.json()
            if objects:
                # Сортируем объекты по имени в алфавитном порядке
                objects.sort(key=lambda obj: obj['name'])
                # Создаем клавиатуру с кнопками в две колонки
                keyboard = []
                for i in range(0, len(objects), 2):
                    row = [
                        InlineKeyboardButton(objects[i]['name'], callback_data=f'select_object_{objects[i]["id"]}')
                    ]
                    if i + 1 < len(objects):
                        row.append(InlineKeyboardButton(objects[i + 1]['name'],
                                                        callback_data=f'select_object_{objects[i + 1]["id"]}'))
                    keyboard.append(row)
                keyboard.append([InlineKeyboardButton("Назад", callback_data='main_menu')])
                reply_markup = InlineKeyboardMarkup(keyboard)
                await query.message.reply_text('Выберите новый объект:', reply_markup=reply_markup)
                context.user_data['stage'] = 'choose_new_object'
            else:
                await query.message.reply_text('Нет доступных объектов.')
        else:
            await query.message.reply_text('Ошибка при получении списка объектов. Попробуйте снова.')



    elif data.startswith('select_object_'):
        await query.message.delete()
        new_object_id = int(data.split('_')[2])

        # Получаем user_id по chat_id
        chat_id = query.message.chat_id
        user_response = requests.get(f'{DJANGO_API_URL}users/chat/{chat_id}/')
        if user_response.status_code == 200:
            user_data = user_response.json()
            user_id = user_data['id']

            # Обновляем object_id пользователя с использованием метода PUT
            user_data['object_id'] = new_object_id
            response = requests.put(f'{DJANGO_API_URL}users/{user_id}/', json=user_data)
            if response.status_code == 200:
                await query.message.reply_text('Ваш объект был успешно изменен.')

                # Возвращаем пользователя в главное меню
                full_name = user_data['full_name']
                organization_id = user_data['organization_id']
                await send_main_menu(chat_id, context, full_name, organization_id)
            else:
                await query.message.reply_text('Ошибка при обновлении объекта. Попробуйте снова.')
        else:
            await query.message.reply_text('Ошибка при получении данных пользователя. Попробуйте снова.')

    elif query.data == "fact_production":
        await query.message.delete()
        await send_prefab_types(query.message.chat_id, context)

    elif query.data.startswith("prefab_type_"):
        await query.message.delete()
        prefab_type_id = int(query.data.split("_")[2])
        context.user_data['selected_prefab_type_id'] = prefab_type_id
        await send_prefab_subtypes(query.message.chat_id, context, prefab_type_id)

    elif query.data.startswith("prefab_subtype_"):
        await query.message.delete()
        prefab_subtype_id = int(query.data.split("_")[2])
        context.user_data['selected_prefab_subtype_id'] = prefab_subtype_id
        await send_prefabs(query.message.chat_id, context, prefab_subtype_id)

    elif data.startswith('prefab_stage_'):
        await query.message.delete()
        parts = data.split('_')
        if len(parts) == 4:
            _, _, stage, prefab_id = parts
            context.user_data['selected_prefab_id'] = int(prefab_id)
            context.user_data['selected_stage'] = stage
            context.user_data['expecting_new_status_quantity'] = True
            await query.message.reply_text("Введите количество для префаба:")
        else:
            await query.message.reply_text("Некорректный формат данных.")

    elif query.data.startswith("prefab_"):
        await query.message.delete()
        prefab_id = int(query.data.split("_")[1])
        context.user_data['selected_prefab_id'] = prefab_id
        context.user_data['expecting_prefab_quantity'] = True
        await query.message.reply_text("Введите количество:")

    elif data == 'sgp':
        await query.message.delete()
        await send_prefabs_in_production(query.message.chat_id, context)

    elif query.data.startswith('sgp_prefab_'):
        await query.message.delete()
        prefabs_in_work_id = int(query.data.split('_')[2])
        context.user_data['selected_prefab_in_work_id'] = prefabs_in_work_id
        context.user_data['expecting_sgp_quantity'] = True
        await query.message.reply_text(
            "Введите количество префабов для передачи на СГП:"
        )


    elif data == 'shipment':
        await query.message.delete()
        await send_prefabs_for_shipment(query.message.chat.id, context)


    elif query.data.startswith('shipment_prefab_'):
        await query.message.delete()
        prefabs_in_work_id = int(query.data.split('_')[2])
        context.user_data['selected_prefab_in_work_id'] = prefabs_in_work_id
        context.user_data['expecting_shipment_quantity'] = True
        await query.message.reply_text(
            "Введите количество префабов для отправки на отгрузку:"
        )


    elif data == 'prefabsoptionlist':
        await query.message.delete()
        keyboard = [
            [InlineKeyboardButton("📄 Отчеты", callback_data='reportmenuprefab')],
            [InlineKeyboardButton("🏭 Просмотр завода", callback_data='view_prefabs')],
            [InlineKeyboardButton("📦 Площадка", callback_data='placespace')],
            [InlineKeyboardButton("🔩 Монтаж", callback_data='montage')],
            [InlineKeyboardButton("Назад", callback_data='main_menu')]
        ]
        reply_markup = InlineKeyboardMarkup(keyboard)
        await query.message.reply_text('Выберите действие с префабами:', reply_markup=reply_markup)


    elif data == 'reportmenuprefab':
        await query.message.delete()
        keyboard = [
            [InlineKeyboardButton("📄 Отчет за сегодня", callback_data='generate_report_today')],
            [InlineKeyboardButton("📅 Отчет на определенный день", callback_data='generate_report_specific_day')],
            [InlineKeyboardButton("Назад", callback_data='prefabsoptionlist')]

        ]
        reply_markup = InlineKeyboardMarkup(keyboard)
        await query.message.reply_text('Выберите вид отчета:', reply_markup=reply_markup)


    elif data == 'generate_report_specific_day':
        await query.message.delete()

        # Разбиваем месяцы на три строки по четыре месяца в строке
        months = [
            ["Январь", "Февраль", "Март", "Апрель"],
            ["Май", "Июнь", "Июль", "Август"],
            ["Сентябрь", "Октябрь", "Ноябрь", "Декабрь"]
        ]

        keyboard = [
            [InlineKeyboardButton(month, callback_data=f'reportmonth_{i * 4 + j + 1}') for j, month in enumerate(row)]
            for i, row in enumerate(months)
        ]
        keyboard.append([InlineKeyboardButton("Назад", callback_data='prefabsoptionlist')])
        reply_markup = InlineKeyboardMarkup(keyboard)
        await query.message.reply_text('Выберите месяц:', reply_markup=reply_markup)


    elif data.startswith('reportmonth_'):
        await query.message.delete()
        month = int(data.split('_')[1])
        context.user_data['selected_month'] = month
        days_in_month = (datetime(2024, month % 12 + 1, 1) - timedelta(days=1)).day

        # Разбиваем дни на строки по 7 дней в каждой строке
        keyboard = [
            [InlineKeyboardButton(str(day), callback_data=f'reportday_{day}') for day in
             range(i, min(i + 7, days_in_month + 1))]
            for i in range(1, days_in_month + 1, 7)
        ]
        keyboard.append([InlineKeyboardButton("Назад", callback_data='generate_report_specific_day')])
        reply_markup = InlineKeyboardMarkup(keyboard)

        await query.message.reply_text('Выберите день:', reply_markup=reply_markup)


    elif data.startswith('reportday_'):
        chat_id = query.message.chat.id
        # Удаление кнопок и изменение текста на "Создание отчета, подождите..."

        await query.message.edit_text("Создание отчета, подождите...", reply_markup=None)
        day = int(data.split('_')[1])
        month = context.user_data['selected_month']
        selected_date = datetime(2024, month, day).strftime('%Y-%m-%d')

        # Теперь можно использовать функцию для создания отчета, передав выбранную дату
        await report_specific_day_pdf(chat_id=query.message.chat_id, context=context, selected_date=selected_date)


    elif data == 'generate_report_today':
        chat_id = query.message.chat.id
        # Удаление кнопок и изменение текста на "Создание отчета, подождите..."

        await query.message.edit_text("Создание отчета, подождите...", reply_markup=None)
        # Запуск процесса создания отчета
        await report_today_pdf(chat_id, context)

    elif data == 'view_prefabs':
        await query.message.delete()
        keyboard = [
            [InlineKeyboardButton("🏭 В производстве", callback_data='view_prefabs_production')],
            [InlineKeyboardButton("📋 СГП", callback_data='view_prefabs_sgp')],
            [InlineKeyboardButton("🚚 Отгружены", callback_data='view_prefabs_shipped')],
            [InlineKeyboardButton("Назад", callback_data='prefabsoptionlist')]
        ]
        reply_markup = InlineKeyboardMarkup(keyboard)
        await query.message.reply_text('Выберите категорию префабов:', reply_markup=reply_markup)

    elif data == 'view_prefabs_production':
        await query.message.delete()
        await send_prefabs_list(query.message.chat.id, context, 'production')

    elif data == 'view_prefabs_sgp':
        await query.message.delete()
        await send_prefabs_list(query.message.chat.id, context, 'sgp')

    elif data == 'view_prefabs_shipped':
        await query.message.delete()
        await send_prefabs_list(query.message.chat.id, context, 'shipment')

    # Обработчик для выбора склада
    elif data == 'placespace':
        await query.message.delete()
        await send_warehouses_list(query.message.chat.id, context)

    elif query.data.startswith('select_warehouse_'):
        await query.message.delete()
        warehouse_id = int(query.data.split('_')[2])
        context.user_data['selected_warehouse_id'] = warehouse_id

        await send_prefabs_list_for_shipment(update.callback_query.message.chat_id, context, warehouse_id)

    elif query.data.startswith('prefabinstock_'):
        await query.message.delete()
        prefabs_in_work_id = int(query.data.split('_')[1])
        context.user_data['selected_prefab_in_work_id'] = prefabs_in_work_id
        context.user_data['expecting_stock_quantity'] = True
        await query.message.reply_text(
            "Введите количество префабов для отгрузки на склад:"
        )

    elif data == 'acceptstockquantity':
        await query.message.delete()
        await handle_accept_stock_quantity(update, context)


    elif data == 'acceptstockquantity_with_comments':
        await query.message.delete()
        await handle_accept_stock_quantity(update, context, with_comments=True)

    elif data == 'remark_stockquantity':
        await query.message.delete()
        await handle_accept_stock_quantity(update, context, remark=True)

    elif data.startswith('remark_'):
        await query.message.delete()
        remark_id = int(data.split('_')[1])
        await choose_stage_for_prefab(query, context, remark_id)

    elif data.startswith('selectwarehouse_for_montage_'):
        await query.message.delete()
        warehouse_id = int(data.split('_')[-1])
        context.user_data['selected_warehouse_id'] = warehouse_id
        await send_prefab_types_montage(update.callback_query.message.chat.id, context)

    elif data.startswith('select_prefab_type_for_montage_'):
        await query.message.delete()
        prefab_type_id = int(data.split('_')[-1])
        context.user_data['selected_prefab_type_id'] = prefab_type_id
        await send_prefab_subtypes_montage(update.callback_query.message.chat.id, context)

    elif data.startswith('select_prefab_subtype_for_montage_'):
        await query.message.delete()
        prefab_subtype_id = int(data.split('_')[-1])
        context.user_data['selected_prefab_subtype_id'] = prefab_subtype_id
        await send_prefabs_list_montage(update.callback_query.message.chat.id, context)

    elif data.startswith('prefabin_stock_'):
        await query.message.delete()
        prefab_id = int(data.split('_')[-1])
        context.user_data['selected_prefab_id'] = prefab_id
        context.user_data['expecting_montage_quantity'] = True
        await query.message.reply_text(
            "Введите количество префабов для отгрузки на монтаж:"
        )


    elif data.startswith('select_block_section_'):
        await handle_select_block_section(query, context)

    elif data.startswith('select_floor_'):
        await handle_select_floor(query, context)


    elif data.startswith('select_prefab_subtype_'):
        await query.message.delete()
        prefab_subtype_id = int(data.split('_')[-1])
        context.user_data['selected_prefab_subtype_id'] = prefab_subtype_id
        await send_prefabs_list_montage(update.callback_query.message.chat.id, context)

    elif data == 'montage':
        await query.message.delete()
        await send_warehouses_list_montage(query.message.chat.id, context)

    elif data == 'support':
        await query.message.delete()

        keyboard = []
        keyboard.append([InlineKeyboardButton("\U0000274C Отмена", callback_data='main_menu')])
        reply_markup = InlineKeyboardMarkup(keyboard)

        context.user_data['expecting_prefab_quantity'] = False
        context.user_data['stage'] = 'support_question'
        await query.message.reply_text('Пожалуйста, введите ваш вопрос для тех. поддержки:', reply_markup=reply_markup)

    elif data.startswith('ticket_'):
        ticket_id = int(data.split('_')[1])
        await handle_ticket_selection(update, context, ticket_id)

    elif data == 'remarks':
        await query.message.delete()
        user_response = requests.get(f'{DJANGO_API_URL}users/chat/{query.message.chat_id}/')
        if user_response.status_code == 200:
            user_data = user_response.json()
            organization_id = user_data.get('organization_id')
            await send_remarks(query.message.chat_id, context, organization_id)
        else:
            await query.message.reply_text("Ошибка при получении данных пользователя.")

    elif data.startswith('answer_'):
        await query.message.delete()
        ticket_id = int(data.split('_')[1])
        context.user_data['ticket_id'] = ticket_id
        context.user_data['stage'] = 'support_answer'
        await query.message.reply_text('Введите ваш ответ:')

    elif data == 'edit_prefab':
        await query.message.delete()
        keyboard = [
            [InlineKeyboardButton("Факт производство", callback_data='edit_production')],
            [InlineKeyboardButton("СГП", callback_data='edit_sgp')],
            [InlineKeyboardButton("Отгрузка", callback_data='edit_shipment')],
            [InlineKeyboardButton("Назад", callback_data='main_menu')]
        ]
        reply_markup = InlineKeyboardMarkup(keyboard)
        await query.message.reply_text("Выберите статус для редактирования префаба:", reply_markup=reply_markup)

    elif data.startswith('edit_'):
        await query.message.delete()
        _, stage = data.split('_')
        context.user_data['selected_stage'] = stage
        await send_prefabs_in_stage(query.message.chat.id, context, stage)

    elif data.startswith('refactorprefab_'):
        await query.message.delete()
        try:
            selected_prefab_id = int(data.split('_')[1])
            context.user_data['selected_prefab_id'] = selected_prefab_id
            context.user_data['refactor_prefab_count'] = True

            keyboard = []
            keyboard.append([InlineKeyboardButton("\U0000274C Отмена", callback_data='main_menu')])
            reply_markup = InlineKeyboardMarkup(keyboard)

            message = await query.message.reply_text("Введите количество префаба для изменения:", reply_markup=reply_markup)
            context.user_data['refactor_message_id'] = message.message_id
        except (IndexError, ValueError):
            await query.message.reply_text("Некорректный формат данных.")

    elif data == 'accept_refactor_quantity':
        await query.message.delete()
        await handle_refactor_prefab_quantity(update, context)

    elif data.startswith('new_status_'):
        await query.message.delete()
        await handle_new_status(update, context)

    elif data == 'summary_by_object':
        await query.message.delete()
        await send_objects_list(query.message.chat_id, context)

    elif data.startswith('selectobjectprefabs_'):
        print(data)
        object_id = int(data.split('_')[1])
        await query.message.delete()
        await send_prefab_summary(query.message.chat_id, context, object_id)

def main() -> None:
    # Вставьте свой токен


    #оригинальный
    application = Application.builder().token("7363654158:AAFfqLnieUtbqgpoKnTH0TAQajNRa4xjg-M").build()


    #тестовый
    # application = Application.builder().token("7313015944:AAGpc2o5qF3rnYph_xRKUWNKaSjedPog1bs").build()


    application.add_handler(CommandHandler("info", welcome_message))
    application.add_handler(CommandHandler("choice", choose_organization))
    application.add_handler(CommandHandler("start", start))
    application.add_handler(CommandHandler("done", handle_done_command))
    application.add_handler(CommandHandler("tech", handle_tech_command))
    application.add_handler(CallbackQueryHandler(button))
    application.add_handler(MessageHandler(filters.TEXT & ~filters.COMMAND, handle_workforce_count))
    application.add_handler(MessageHandler(filters.PHOTO & ~filters.COMMAND, handle_photo_upload))

    application.run_polling()


if __name__ == '__main__':
    main()